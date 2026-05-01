import json
import os
import base64
import uuid
import logging
import jwt
import psycopg2
import boto3
from contextlib import contextmanager
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)

ALLOWED_ORIGINS = [o.strip() for o in os.environ.get('ALLOWED_ORIGINS', '').split(',') if o.strip()]

def get_cors(event: dict) -> dict:
    origin = (event.get('headers') or {}).get('origin') or (event.get('headers') or {}).get('Origin') or ''
    allowed = origin if (origin and (any(origin == o for o in ALLOWED_ORIGINS) or not ALLOWED_ORIGINS)) else (ALLOWED_ORIGINS[0] if ALLOWED_ORIGINS else '*')
    return {
        'Access-Control-Allow-Origin': allowed,
        'Access-Control-Allow-Methods': 'GET, POST, PUT, PATCH, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization, X-Authorization',
        'Access-Control-Allow-Credentials': 'true',
    }
JWT_SECRET = os.environ['JWT_SECRET']
S3_KEY = os.environ.get('AWS_ACCESS_KEY_ID', '')
S3_SECRET = os.environ.get('AWS_SECRET_ACCESS_KEY', '')


@contextmanager
def get_db():
    conn = psycopg2.connect(os.environ['DATABASE_URL'])
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def verify_token(event: dict):
    headers = event.get('headers') or {}
    auth = headers.get('X-Authorization') or headers.get('Authorization') or ''
    token = auth[7:].strip() if auth.startswith('Bearer ') else ''
    if not token:
        return None
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
    except Exception:
        return None


def _make_ok(cors):
    def ok(data, status=200):
        return {'statusCode': status, 'headers': {**cors, 'Content-Type': 'application/json'}, 'body': json.dumps(data, default=str)}
    return ok

def _make_err(cors):
    def err(msg, status=400):
        return {'statusCode': status, 'headers': {**cors, 'Content-Type': 'application/json'}, 'body': json.dumps({'error': msg})}
    return err


def s3_client():
    return boto3.client(
        's3',
        endpoint_url='https://bucket.poehali.dev',
        aws_access_key_id=S3_KEY,
        aws_secret_access_key=S3_SECRET,
    )


def row_to_client(row, cur):
    cols = [d[0] for d in cur.description]
    return dict(zip(cols, row))


def _get_company(user_id) -> dict:
    """Читает реквизиты компании из app-state пользователя."""
    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT state FROM app_state WHERE user_id = %s', (user_id,))
            row = cur.fetchone()
            if not row:
                cur.execute('SELECT state FROM app_state WHERE id = 1 AND user_id IS NULL')
                row = cur.fetchone()
        if not row:
            return {}
        state = row[0] if isinstance(row[0], dict) else json.loads(row[0])
        return state.get('settings', {}).get('company', {}) or {}
    except Exception:
        return {}


def _co(company: dict, field: str, fallback: str = '___________') -> str:
    """Возвращает поле компании или заглушку."""
    return str(company.get(field) or '').strip() or fallback


def log_history(conn, client_id, payload, action, description, old_val=None, new_val=None):
    cur = conn.cursor()
    cur.execute(
        '''INSERT INTO client_history (client_id, user_id, user_name, action, description, old_value, new_value)
           VALUES (%s, %s, %s, %s, %s, %s, %s)''',
        (
            client_id,
            payload.get('sub'),
            payload.get('login', ''),
            action,
            description,
            json.dumps(old_val) if old_val else None,
            json.dumps(new_val) if new_val else None,
        )
    )


def handler(event: dict, context) -> dict:
    """
    CRUD клиентов + загрузка фото.
    GET    ?action=list               — список всех клиентов
    GET    ?action=get&id=UUID        — один клиент + фото + история
    POST   ?action=create             — создать клиента
    POST   ?action=update&id=UUID     — обновить клиента
    POST   ?action=status&id=UUID     — сменить статус
    POST   ?action=upload_photo&id=UUID — загрузить фото (base64)
    POST   ?action=delete_photo&photo_id=UUID — удалить фото
    """
    cors = get_cors(event)
    ok = _make_ok(cors)
    err = _make_err(cors)

    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': cors, 'body': ''}

    payload = verify_token(event)
    if not payload:
        return err('Не авторизован', 401)

    qs = event.get('queryStringParameters') or {}
    action = qs.get('action', 'list')
    method = event.get('httpMethod', 'GET')

    # ── LIST ──────────────────────────────────────────────────────
    if action == 'list':
        page = max(1, int(qs.get('page', 1)))
        per_page = min(500, max(1, int(qs.get('per_page', 50))))
        search_q = (qs.get('q') or '').strip()
        status_f = (qs.get('status') or '').strip()
        designer_f = (qs.get('designer') or '').strip()
        measurer_f = (qs.get('measurer') or '').strip()
        date_from = (qs.get('date_from') or '').strip()
        date_to = (qs.get('date_to') or '').strip()
        delivery_from = (qs.get('delivery_from') or '').strip()
        delivery_to = (qs.get('delivery_to') or '').strip()
        amount_min = qs.get('amount_min')
        amount_max = qs.get('amount_max')
        sort_field = qs.get('sort', 'created_at')
        sort_dir = 'ASC' if qs.get('sort_dir', 'desc').lower() == 'asc' else 'DESC'
        allowed_sorts = {'created_at', 'delivery_date', 'total_amount', 'last_name'}
        if sort_field not in allowed_sorts:
            sort_field = 'created_at'
        offset = (page - 1) * per_page

        conditions = []
        params = []

        if search_q:
            conditions.append(
                "(last_name ILIKE %s OR first_name ILIKE %s OR middle_name ILIKE %s "
                "OR phone ILIKE %s OR phone2 ILIKE %s OR contract_number ILIKE %s)"
            )
            like = f'%{search_q}%'
            params.extend([like, like, like, like, like, like])

        if status_f and status_f != 'all':
            conditions.append('status = %s')
            params.append(status_f)

        if designer_f:
            conditions.append('designer = %s')
            params.append(designer_f)

        if measurer_f:
            conditions.append('measurer = %s')
            params.append(measurer_f)

        if date_from:
            conditions.append('created_at::date >= %s')
            params.append(date_from)

        if date_to:
            conditions.append('created_at::date <= %s')
            params.append(date_to)

        if delivery_from:
            conditions.append('delivery_date >= %s')
            params.append(delivery_from)

        if delivery_to:
            conditions.append('delivery_date <= %s')
            params.append(delivery_to)

        if amount_min is not None:
            try:
                conditions.append('total_amount >= %s')
                params.append(float(amount_min))
            except ValueError:
                pass

        if amount_max is not None:
            try:
                conditions.append('total_amount <= %s')
                params.append(float(amount_max))
            except ValueError:
                pass

        where = ('WHERE ' + ' AND '.join(conditions)) if conditions else ''
        null_last = 'NULLS LAST' if sort_dir == 'ASC' else 'NULLS LAST'

        with get_db() as conn:
            cur = conn.cursor()
            cur.execute(f'SELECT COUNT(*) FROM clients {where}', params)
            total = cur.fetchone()[0]
            cur.execute(f'''
                SELECT id, status, last_name, first_name, middle_name, phone, phone2, messenger,
                       contract_number, contract_date, total_amount, payment_type,
                       delivery_date, designer, measurer, reminder_date, reminder_note,
                       comment, created_at, updated_at, project_ids,
                       source, tags, rating, property_type, property_area, has_children, has_pets
                FROM clients {where}
                ORDER BY {sort_field} {sort_dir} {null_last}
                LIMIT %s OFFSET %s
            ''', params + [per_page, offset])
            rows = cur.fetchall()
            cols = [d[0] for d in cur.description]
            clients = [dict(zip(cols, r)) for r in rows]
        return ok({
            'clients': clients,
            'total': total,
            'page': page,
            'per_page': per_page,
            'pages': (total + per_page - 1) // per_page,
        })

    # ── GET ONE ───────────────────────────────────────────────────
    if action == 'get':
        cid = qs.get('id')
        if not cid:
            return err('Нет id')
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT * FROM clients WHERE id = %s', (cid,))
            row = cur.fetchone()
            if not row:
                return err('Клиент не найден', 404)
            client = row_to_client(row, cur)
            cur.execute('SELECT * FROM client_photos WHERE client_id = %s ORDER BY uploaded_at', (cid,))
            photos_rows = cur.fetchall()
            pcols = [d[0] for d in cur.description]
            photos = [dict(zip(pcols, r)) for r in photos_rows]
            cur.execute('SELECT * FROM client_history WHERE client_id = %s ORDER BY created_at DESC LIMIT 50', (cid,))
            hist_rows = cur.fetchall()
            hcols = [d[0] for d in cur.description]
            history = [dict(zip(hcols, r)) for r in hist_rows]
        return ok({'client': client, 'photos': photos, 'history': history})

    # ── CREATE ────────────────────────────────────────────────────
    if action == 'create' and method == 'POST':
        body = json.loads(event.get('body') or '{}')
        c = body.get('client', {})
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('''
            INSERT INTO clients (
                status, last_name, first_name, middle_name, phone, phone2, messenger, email,
                passport_series, passport_number, passport_issued_by, passport_issued_date, passport_dept_code,
                reg_city, reg_street, reg_house, reg_apt,
                delivery_city, delivery_street, delivery_house, delivery_apt,
                delivery_entrance, delivery_floor, delivery_elevator, delivery_note,
                contract_number, contract_date, products, total_amount, payment_type,
                prepaid_amount, balance_due, custom_payment_scheme,
                delivery_date, production_days, assembly_days,
                delivery_cost, assembly_cost,
                designer, measurer, manager_name, project_ids, reminder_date, reminder_note, comment,
                source, tags, rating, property_type, property_area, has_children, has_pets,
                created_by, updated_by
            ) VALUES (
                %s,%s,%s,%s,%s,%s,%s,%s,
                %s,%s,%s,%s,%s,
                %s,%s,%s,%s,
                %s,%s,%s,%s,
                %s,%s,%s,%s,
                %s,%s,%s,%s,%s,
                %s,%s,%s,
                %s,%s,%s,
                %s,%s,
                %s,%s,%s,%s,%s,%s,%s,
                %s,%s,%s,%s,%s,%s,%s,
                %s,%s
            ) RETURNING id
        ''', (
            c.get('status', 'new'),
            c.get('last_name', ''), c.get('first_name', ''), c.get('middle_name', ''),
            c.get('phone', ''), c.get('phone2', ''), c.get('messenger', 'WhatsApp'), c.get('email', ''),
            c.get('passport_series', ''), c.get('passport_number', ''), c.get('passport_issued_by', ''),
            c.get('passport_issued_date', ''), c.get('passport_dept_code', ''),
            c.get('reg_city', ''), c.get('reg_street', ''), c.get('reg_house', ''), c.get('reg_apt', ''),
            c.get('delivery_city', ''), c.get('delivery_street', ''), c.get('delivery_house', ''),
            c.get('delivery_apt', ''), c.get('delivery_entrance', ''), c.get('delivery_floor', ''),
            c.get('delivery_elevator', 'нет'), c.get('delivery_note', ''),
            c.get('contract_number', ''), c.get('contract_date', ''),
            json.dumps(c.get('products', [])),
            c.get('total_amount', 0), c.get('payment_type', '100% предоплата'),
            c.get('prepaid_amount', 0), c.get('balance_due', 0), c.get('custom_payment_scheme', ''),
            c.get('delivery_date', ''), c.get('production_days', 0), c.get('assembly_days', 0),
            c.get('delivery_cost', 0), c.get('assembly_cost', 0),
            c.get('designer', ''), c.get('measurer', ''), c.get('manager_name', ''),
            json.dumps(c.get('project_ids', [])),
            c.get('reminder_date', ''), c.get('reminder_note', ''), c.get('comment', ''),
            c.get('source', ''), c.get('tags', []), c.get('rating') or None,
            c.get('property_type', ''), c.get('property_area', ''),
            bool(c.get('has_children', False)), bool(c.get('has_pets', False)),
            payload.get('sub'), payload.get('sub'),
        ))
            new_id = cur.fetchone()[0]
            log_history(conn, str(new_id), payload, 'created', 'Клиент создан')
        return ok({'id': str(new_id)}, 201)

    # ── UPDATE ────────────────────────────────────────────────────
    if action == 'update' and method == 'POST':
        cid = qs.get('id')
        if not cid:
            return err('Нет id')
        body = json.loads(event.get('body') or '{}')
        c = body.get('client', {})
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT last_name, first_name, status FROM clients WHERE id = %s', (cid,))
            old = cur.fetchone()
            if not old:
                return err('Клиент не найден', 404)
            cur.execute('''
                UPDATE clients SET
                    status=%s, last_name=%s, first_name=%s, middle_name=%s,
                    phone=%s, phone2=%s, messenger=%s, email=%s,
                    passport_series=%s, passport_number=%s, passport_issued_by=%s,
                    passport_issued_date=%s, passport_dept_code=%s,
                    reg_city=%s, reg_street=%s, reg_house=%s, reg_apt=%s,
                    delivery_city=%s, delivery_street=%s, delivery_house=%s, delivery_apt=%s,
                    delivery_entrance=%s, delivery_floor=%s, delivery_elevator=%s, delivery_note=%s,
                    contract_number=%s, contract_date=%s, products=%s,
                    total_amount=%s, payment_type=%s, prepaid_amount=%s, balance_due=%s,
                    custom_payment_scheme=%s, delivery_date=%s, production_days=%s, assembly_days=%s,
                    delivery_cost=%s, assembly_cost=%s,
                    designer=%s, measurer=%s, manager_name=%s, project_ids=%s,
                    reminder_date=%s, reminder_note=%s, comment=%s,
                    source=%s, tags=%s, rating=%s, property_type=%s, property_area=%s,
                    has_children=%s, has_pets=%s,
                    updated_at=NOW(), updated_by=%s
                WHERE id=%s
            ''', (
                c.get('status', 'new'),
                c.get('last_name', ''), c.get('first_name', ''), c.get('middle_name', ''),
                c.get('phone', ''), c.get('phone2', ''), c.get('messenger', 'WhatsApp'), c.get('email', ''),
                c.get('passport_series', ''), c.get('passport_number', ''), c.get('passport_issued_by', ''),
                c.get('passport_issued_date', ''), c.get('passport_dept_code', ''),
                c.get('reg_city', ''), c.get('reg_street', ''), c.get('reg_house', ''), c.get('reg_apt', ''),
                c.get('delivery_city', ''), c.get('delivery_street', ''), c.get('delivery_house', ''),
                c.get('delivery_apt', ''), c.get('delivery_entrance', ''), c.get('delivery_floor', ''),
                c.get('delivery_elevator', 'нет'), c.get('delivery_note', ''),
                c.get('contract_number', ''), c.get('contract_date', ''),
                json.dumps(c.get('products', [])),
                c.get('total_amount', 0), c.get('payment_type', '100% предоплата'),
                c.get('prepaid_amount', 0), c.get('balance_due', 0), c.get('custom_payment_scheme', ''),
                c.get('delivery_date', ''), c.get('production_days', 0), c.get('assembly_days', 0),
                c.get('delivery_cost', 0), c.get('assembly_cost', 0),
                c.get('designer', ''), c.get('measurer', ''), c.get('manager_name', ''),
                json.dumps(c.get('project_ids', [])),
                c.get('reminder_date', ''), c.get('reminder_note', ''), c.get('comment', ''),
                c.get('source', ''), c.get('tags', []), c.get('rating') or None,
                c.get('property_type', ''), c.get('property_area', ''),
                bool(c.get('has_children', False)), bool(c.get('has_pets', False)),
                payload.get('sub'), cid,
            ))
            log_history(conn, cid, payload, 'updated', 'Данные клиента обновлены')
        return ok({'ok': True})

    # ── STATUS ────────────────────────────────────────────────────
    if action == 'status' and method == 'POST':
        cid = qs.get('id')
        if not cid:
            return err('Нет id')
        body = json.loads(event.get('body') or '{}')
        new_status = body.get('status')
        if not new_status:
            return err('Нет status')
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT status FROM clients WHERE id = %s', (cid,))
            row = cur.fetchone()
            if not row:
                return err('Клиент не найден', 404)
            old_status = row[0]
            cur.execute('UPDATE clients SET status=%s, updated_at=NOW() WHERE id=%s', (new_status, cid))
            log_history(conn, cid, payload, 'status_changed', f'Статус: {old_status} → {new_status}',
                        {'status': old_status}, {'status': new_status})
        return ok({'ok': True})

    # ── UPLOAD PHOTO ──────────────────────────────────────────────
    if action == 'upload_photo' and method == 'POST':
        cid = qs.get('id')
        if not cid:
            return err('Нет id')
        body = json.loads(event.get('body') or '{}')
        data_b64 = body.get('data', '')
        category = body.get('category', 'measure')
        name = body.get('name', 'photo.jpg')
        content_type = body.get('content_type', 'image/jpeg')

        ALLOWED_TYPES = {'image/jpeg', 'image/png', 'image/gif', 'image/webp'}
        if content_type not in ALLOWED_TYPES:
            return err('Недопустимый тип файла. Разрешены: JPEG, PNG, GIF, WEBP')

        MAX_SIZE_BYTES = 10 * 1024 * 1024
        if len(data_b64) > MAX_SIZE_BYTES * 4 // 3:
            return err('Файл слишком большой. Максимум 10 МБ')

        img_data = base64.b64decode(data_b64)

        if len(img_data) > MAX_SIZE_BYTES:
            return err('Файл слишком большой. Максимум 10 МБ')
        photo_id = str(uuid.uuid4())
        ext = name.rsplit('.', 1)[-1] if '.' in name else 'jpg'
        key = f'clients/{cid}/{photo_id}.{ext}'

        s3 = s3_client()
        s3.put_object(Bucket='files', Key=key, Body=img_data, ContentType=content_type)
        cdn_url = f'https://cdn.poehali.dev/projects/{S3_KEY}/bucket/{key}'

        with get_db() as conn:
            cur = conn.cursor()
            cur.execute(
                'INSERT INTO client_photos (id, client_id, category, url, name, uploaded_by) VALUES (%s,%s,%s,%s,%s,%s) RETURNING id',
                (photo_id, cid, category, cdn_url, name, payload.get('sub'))
            )
            log_history(conn, cid, payload, 'photo_added', f'Добавлено фото: {name} ({category})')
        return ok({'id': photo_id, 'url': cdn_url}, 201)

    # ── DELETE PHOTO ──────────────────────────────────────────────
    if action == 'delete_photo' and method == 'POST':
        photo_id = qs.get('photo_id')
        if not photo_id:
            return err('Нет photo_id')
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT client_id, url, name FROM client_photos WHERE id = %s', (photo_id,))
            row = cur.fetchone()
            if not row:
                return err('Фото не найдено', 404)
            cid, url, name = row
            s3_key = url.split('/bucket/', 1)[-1] if '/bucket/' in url else None
            if s3_key:
                try:
                    s3_client().delete_object(Bucket='files', Key=s3_key)
                except Exception as e:
                    logger.error(f'S3 delete failed for {s3_key}: {e}')
            cur.execute('DELETE FROM client_photos WHERE id = %s', (photo_id,))
            log_history(conn, str(cid), payload, 'photo_deleted', f'Удалено фото: {name}')
        return ok({'ok': True})

    # ── UPLOAD COMPANY ASSET (печать / подпись) ───────────────────
    if action == 'upload_asset' and method == 'POST':
        body = json.loads(event.get('body') or '{}')
        data_b64 = body.get('data', '')
        asset_type = body.get('asset_type', 'stamp')  # 'stamp' | 'signature'
        name = body.get('name', 'asset.png')
        content_type = body.get('content_type', 'image/png')

        if content_type not in {'image/jpeg', 'image/png', 'image/webp'}:
            return err('Допустимы только PNG, JPEG, WEBP')
        if not data_b64:
            return err('Нет данных файла')

        img_data = base64.b64decode(data_b64)
        if len(img_data) > 5 * 1024 * 1024:
            return err('Файл слишком большой. Максимум 5 МБ')

        user_id = payload.get('sub') or payload.get('user_id') or payload.get('id')
        ext = name.rsplit('.', 1)[-1] if '.' in name else 'png'
        asset_id = str(uuid.uuid4())
        key = f'company/{user_id}/{asset_type}/{asset_id}.{ext}'

        s3 = s3_client()
        s3.put_object(Bucket='files', Key=key, Body=img_data, ContentType=content_type)
        cdn_url = f'https://cdn.poehali.dev/projects/{S3_KEY}/bucket/{key}'
        return ok({'url': cdn_url}, 201)

    # ── DOCUMENT: HTML preview ────────────────────────────────────
    if action == 'doc_html':
        cid = qs.get('client_id')
        doc_type = qs.get('doc', 'contract')
        if not cid:
            return err('Нет client_id')
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT * FROM clients WHERE id = %s', (cid,))
            row = cur.fetchone()
            if not row:
                return err('Клиент не найден', 404)
            cols = [d[0] for d in cur.description]
            client = dict(zip(cols, row))
        user_id = payload.get('sub') or payload.get('user_id') or payload.get('id')
        company = _get_company(user_id)
        html = _build_contract_html(client, doc_type, company)
        return {'statusCode': 200, 'headers': {**cors, 'Content-Type': 'text/html; charset=utf-8'}, 'body': html}

    # ── DOCUMENT: save HTML to S3, return link ────────────────────
    if action == 'doc_link':
        cid = qs.get('client_id')
        doc_type = qs.get('doc', 'contract')
        if not cid:
            return err('Нет client_id')
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT * FROM clients WHERE id = %s', (cid,))
            row = cur.fetchone()
            if not row:
                return err('Клиент не найден', 404)
            cols = [d[0] for d in cur.description]
            client = dict(zip(cols, row))
        user_id = payload.get('sub') or payload.get('user_id') or payload.get('id')
        company = _get_company(user_id)
        html = _build_contract_html(client, doc_type, company)
        doc_id = str(uuid.uuid4())
        key = f'documents/{doc_id}.html'
        s3_client().put_object(Bucket='files', Key=key, Body=html.encode('utf-8'), ContentType='text/html; charset=utf-8')
        cdn_url = f'https://cdn.poehali.dev/projects/{S3_KEY}/bucket/{key}'
        return ok({'url': cdn_url})

    # ── DOCUMENT: generate DOCX (возвращаем base64 напрямую, без S3) ─
    if action == 'doc_docx':
        cid = qs.get('client_id')
        doc_type = qs.get('doc', 'contract')
        if not cid:
            return err('Нет client_id')
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute('SELECT * FROM clients WHERE id = %s', (cid,))
            row = cur.fetchone()
            if not row:
                return err('Клиент не найден', 404)
            cols = [d[0] for d in cur.description]
            client = dict(zip(cols, row))
        user_id = payload.get('sub') or payload.get('user_id') or payload.get('id')
        company = _get_company(user_id)
        docx_bytes = _build_docx(client, doc_type, company)
        b64 = base64.b64encode(docx_bytes).decode('utf-8')
        return {
            'statusCode': 200,
            'headers': {**cors, 'Content-Type': 'application/json'},
            'body': json.dumps({'data': b64}),
        }

    return err('Неизвестное действие', 404)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT GENERATION HELPERS
# ═══════════════════════════════════════════════════════════════

def _num_to_words(n: float) -> str:
    n = int(round(n))
    if n == 0:
        return 'ноль рублей'
    ones = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять',
            'десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать',
            'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    ones_f = ['', 'одна', 'две', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять',
              'десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать',
              'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    tens = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    hundreds = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']

    def chunk(num, feminine=False):
        parts = []
        h = num // 100
        t = (num % 100) // 10
        o = num % 10
        if h:
            parts.append(hundreds[h])
        if t == 1:
            parts.append((ones_f if feminine else ones)[num % 100])
        else:
            if t:
                parts.append(tens[t])
            if o:
                parts.append((ones_f if feminine else ones)[o])
        return parts

    result = []
    millions = n // 1_000_000
    thousands = (n % 1_000_000) // 1_000
    remainder = n % 1_000

    if millions:
        parts = chunk(millions)
        o = millions % 10
        t2 = (millions % 100) // 10
        suffix = 'миллионов' if (t2 == 1 or o == 0 or o >= 5) else ('миллион' if o == 1 else 'миллиона')
        result.extend(parts); result.append(suffix)

    if thousands:
        parts = chunk(thousands, feminine=True)
        o = thousands % 10
        t2 = (thousands % 100) // 10
        suffix = 'тысяч' if (t2 == 1 or o == 0 or o >= 5) else ('тысяча' if o == 1 else 'тысячи')
        result.extend(parts); result.append(suffix)

    if remainder:
        result.extend(chunk(remainder))

    o = n % 10
    t2 = (n % 100) // 10
    rub = 'рублей' if (t2 == 1 or o == 0 or o >= 5) else ('рубль' if o == 1 else 'рубля')
    return ' '.join(result) + ' ' + rub


def _days_words(n: int) -> str:
    w = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять',
         'десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать',
         'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать', 'двадцать']
    tens = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    if n <= 20:
        return w[n] if n < len(w) else str(n)
    t = n // 10; o = n % 10
    return (tens[t] + (' ' + w[o] if o else '')).strip()


def _fmt_date(d):
    if not d: return '___________'
    try: return datetime.strptime(str(d)[:10], '%Y-%m-%d').strftime('%d.%m.%Y')
    except: return str(d)


def _fmt_date_full(d):
    months = ['января','февраля','марта','апреля','мая','июня','июля','августа','сентября','октября','ноября','декабря']
    if not d: return '«___» ___________ ______ г.'
    try:
        dt = datetime.strptime(str(d)[:10], '%Y-%m-%d')
        return f'«{dt.day:02d}» {months[dt.month-1]} {dt.year} г.'
    except: return str(d)


def _full_name(c):
    return ' '.join(filter(None,[c.get('last_name',''),c.get('first_name',''),c.get('middle_name','')])) or '___________'


def _passport_str(c):
    s = c.get('passport_series',''); n = c.get('passport_number','')
    return f'{s} {n}' if s and n else '___________'


def _reg_address(c):
    parts = []
    if c.get('reg_city'): parts.append(f"г. {c['reg_city']}")
    if c.get('reg_street'): parts.append(f"ул. {c['reg_street']}")
    if c.get('reg_house'): parts.append(f"д. {c['reg_house']}")
    if c.get('reg_apt'): parts.append(f"кв. {c['reg_apt']}")
    return ', '.join(parts) or '___________'


def _delivery_address(c):
    parts = []
    if c.get('delivery_city'): parts.append(f"г. {c['delivery_city']}")
    if c.get('delivery_street'): parts.append(f"ул. {c['delivery_street']}")
    if c.get('delivery_house'): parts.append(f"д. {c['delivery_house']}")
    if c.get('delivery_apt'): parts.append(f"кв. {c['delivery_apt']}")
    return ', '.join(parts) or '___________'


def _get_products(c):
    products = c.get('products') or []
    if isinstance(products, str):
        try: products = json.loads(products)
        except: products = []
    return products


def _doc_style(title='', contract_num=''):
    header_content = f'«{title}» · № {contract_num}' if contract_num else (title or '')
    return f'''<style>
@import url('https://fonts.googleapis.com/css2?family=PT+Serif:ital,wght@0,400;0,700;1,400&display=swap');
@page{{
  size:A4 portrait;
  margin:20mm 20mm 20mm 25mm;
  @top-center{{
    content:"{header_content}";
    font-family:'PT Serif',Georgia,serif;
    font-size:8pt;
    color:#666;
    padding-bottom:4mm;
  }}
  @bottom-center{{
    content:"Стр. " counter(page) " из " counter(pages);
    font-family:'PT Serif',Georgia,serif;
    font-size:8pt;
    color:#666;
    padding-top:4mm;
  }}
}}
*{{box-sizing:border-box;margin:0;padding:0}}
html{{background:#2d2d2d}}
body{{font-family:'PT Serif',Georgia,serif;font-size:11pt;line-height:1.6;color:#000;background:#fff}}
.page{{width:210mm;min-height:297mm;margin:10mm auto;padding:20mm 20mm 20mm 25mm;background:#fff;box-shadow:0 4px 24px rgba(0,0,0,.45)}}
h1{{font-size:13pt;text-align:center;font-weight:bold;margin:0 0 3px;text-transform:uppercase;letter-spacing:.1em}}
h2{{font-size:11pt;text-align:center;font-weight:normal;margin:0 0 14px}}
.city-date{{display:flex;justify-content:space-between;margin:10px 0 14px;font-size:11pt}}
p{{margin:0 0 6px;text-align:justify;text-indent:1.27cm;line-height:1.6;orphans:3;widows:3}}
p.no-indent{{text-indent:0}}
p.center{{text-align:center;text-indent:0}}
.sec{{font-weight:bold;margin:14px 0 4px;text-indent:0;font-size:11pt;page-break-after:avoid}}
.sec-block{{page-break-inside:avoid}}
table{{width:100%;border-collapse:collapse;margin:8px 0;font-size:10pt;page-break-inside:avoid}}
th,td{{border:1px solid #000;padding:5px 8px}}
th{{background:#f0f0f0;font-weight:bold;text-align:center;font-size:10pt;letter-spacing:.04em}}
td{{vertical-align:top}}
td.num{{text-align:center;width:5%}}
td.right{{text-align:right;font-variant-numeric:tabular-nums}}
tbody tr:nth-child(even){{background:#f8f8f8}}
.ul{{border-bottom:1px solid #000;display:inline-block;min-width:180px}}
.sig-wrap{{margin-top:24px;border-top:1px solid #000;padding-top:10px}}
.sig-table{{width:100%;border-collapse:collapse}}
.sig-table td{{border:none;padding:4px 8px;vertical-align:top;width:50%}}
.sig-table .sig-label{{font-weight:bold;font-size:10pt;text-transform:uppercase;letter-spacing:.08em;padding-bottom:6px}}
.sig-table .sig-line{{border-bottom:1px solid #000;min-width:180px;display:inline-block;margin-right:8px}}
a{{color:#000;text-decoration:none}}
@media print{{
  html{{background:#fff}}
  body{{margin:0;font-size:11pt}}
  .page{{width:auto;min-height:auto;margin:0;padding:0;box-shadow:none}}
  @page{{size:A4 portrait;margin:20mm 20mm 20mm 25mm}}
  a{{color:#000;text-decoration:none}}
}}
</style>'''


def _typo(text: str) -> str:
    import re
    text = re.sub(r'"([^"]*)"', r'«\1»', text)
    text = re.sub(r' - ', ' \u2014 ', text)
    text = re.sub(r' -- ', ' \u2014 ', text)
    return text


def _fmt_money(n: float) -> str:
    if n == int(n):
        return f'{int(n):,}'.replace(',', '\u202f')
    return f'{n:,.2f}'.replace(',', '\u202f')


def _build_contract_html(c: dict, doc_type: str, company: dict = None) -> str:
    if company is None:
        company = {}
    co_name     = _co(company, 'name',              'ООО «Интерьерные решения»')
    co_city     = _co(company, 'city',              'Саратов')
    co_inn      = _co(company, 'inn',               '6450106826')
    co_ogrn     = _co(company, 'ogrn',              '1196451012251')
    co_kpp      = _co(company, 'kpp',               '')
    co_address  = _co(company, 'address',           '410018, Саратовская обл., г. Саратов, ул. Усть-Курдюмская, д. 3, пом. 1')
    co_phone    = _co(company, 'phone',             '')
    co_director = _co(company, 'director',          '')
    co_dir_pos  = _co(company, 'directorPosition',  'менеджера')
    co_bank        = _co(company, 'bank',              '')
    co_bik         = _co(company, 'bik',               '')
    co_rs          = _co(company, 'rs',                '')
    co_ks          = _co(company, 'ks',                '')
    co_poa_num     = _co(company, 'poaNumber',         '')
    co_poa_date    = _co(company, 'poaDate',           '')
    co_stamp_url   = company.get('stampUrl', '') or ''
    co_sig_url     = company.get('signatureUrl', '') or ''
    use_stamp      = bool(company.get('useStamp', False)) and bool(co_stamp_url)
    use_sig        = bool(company.get('useSignature', False)) and bool(co_sig_url)

    # Строка доверенности для преамбулы
    if co_poa_num or co_poa_date:
        poa_num_str  = co_poa_num or '<span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>'
        poa_date_str = co_poa_date or '<span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>'
        poa_str = f'доверенности № {poa_num_str} от {poa_date_str}'
    else:
        poa_str = 'доверенности № <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>'

    bank_block = ''
    if co_bank or co_bik or co_rs or co_ks:
        bank_block = '<br>' + '<br>'.join(filter(None, [
            f'Банк: {co_bank}' if co_bank else '',
            f'БИК: {co_bik}' if co_bik else '',
            f'р/с: {co_rs}' if co_rs else '',
            f'к/с: {co_ks}' if co_ks else '',
        ]))

    # Блок печати/подписи для реквизитов
    stamp_sig_html = ''
    if use_sig or use_stamp:
        stamp_sig_html = '<div style="display:flex;gap:16px;margin-top:12px;align-items:flex-end">'
        if use_sig:
            stamp_sig_html += f'<img src="{co_sig_url}" alt="Подпись" style="height:60px;max-width:120px;object-fit:contain" />'
        if use_stamp:
            stamp_sig_html += f'<img src="{co_stamp_url}" alt="Печать" style="height:70px;max-width:70px;object-fit:contain" />'
        stamp_sig_html += '</div>'

    co_requisites = f'''{co_name}<br>
ОГРН:&nbsp;{co_ogrn}{f' &nbsp;|&nbsp; КПП:&nbsp;{co_kpp}' if co_kpp else ''} &nbsp;|&nbsp; ИНН:&nbsp;{co_inn}<br>
{co_address}{f'<br>Тел.: {co_phone}' if co_phone else ''}{bank_block}'''
    fname = _full_name(c)
    total = float(c.get('total_amount') or 0)
    total_words = _num_to_words(total)
    contract_num = c.get('contract_number') or '___'
    contract_date_full = _fmt_date_full(c.get('contract_date') or '')
    prod_days = int(c.get('production_days') or 45)
    prepaid = float(c.get('prepaid_amount') or 0)
    balance = float(c.get('balance_due') or 0)
    ptype = c.get('payment_type', '100% предоплата')
    custom = c.get('custom_payment_scheme', '') or ''
    products = _get_products(c)

    products_rows = ''
    for i, p in enumerate(products, 1):
        products_rows += f'<tr><td class="num">{i}</td><td>{p.get("name","Кухонный гарнитур")}</td><td class="num">шт.</td><td class="num">{p.get("qty",1)}</td><td class="right"></td></tr>'
    if not products_rows:
        products_rows = f'<tr><td class="num">1</td><td>Кухонный гарнитур</td><td class="num">шт.</td><td class="num">1</td><td class="right">{_fmt_money(total)}</td></tr>'

    if custom:
        pay_html = f'<p>{custom}</p>'
    elif ptype == '100% предоплата':
        pay_html = f'<p>3.2.1. Предварительная оплата производится при заключении Договора в размере {_fmt_money(prepaid)} ({_num_to_words(prepaid)}) рублей.</p><p>3.2.2. Окончательный платёж не предусмотрен. Стоимость работ оплачена полностью при заключении Договора.</p>'
    else:
        pay_html = f'<p>3.2.1. Предварительная оплата производится при заключении Договора в размере {_fmt_money(prepaid)} ({_num_to_words(prepaid)}) рублей.</p><p>3.2.2. Окончательный платёж за выполненные по Договору работы в размере {_fmt_money(balance)} ({_num_to_words(balance)}) рублей осуществляется в течение 3 (трёх) дней с момента получения Заказчиком уведомления о готовности мебели, но не позднее дня доставки.</p>'

    manager = c.get('manager_name') or ''
    manager_line = manager if manager else '&nbsp;' * 30

    if doc_type == 'contract':
        style = _doc_style('Договор бытового подряда', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Договор №{contract_num}</title>{style}</head><body><div class="page">
<h1>ДОГОВОР</h1>
<h2>бытового подряда на изготовление мебели</h2>
<div class="city-date"><span>г. {co_city}</span><span>№ {contract_num} от {contract_date_full}</span></div>
<p class="no-indent">{co_name}, в лице {co_dir_pos}а <strong>{manager_line}</strong>, действующего на основании {poa_str}, именуемый в дальнейшем «Подрядчик», и гр. <strong>{fname}</strong>, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», заключили настоящий Договор о нижеследующем:</p>

<p class="sec">1. ПРЕДМЕТ ДОГОВОРА</p>
<p>1.1. Подрядчик обязуется выполнить работу по изготовлению мебели и передать результат работы Заказчику (мебель передается в разобранном виде), а Заказчик обязуется принять и оплатить результат работ.</p>
<p>1.2. Наименование, качественные характеристики, количество, дизайн мебели указываются в Техническом проекте, который является Приложением № 1 к настоящему Договору.</p>
<p>1.3. В случае необходимости выполнения Подрядчиком дополнительных работ, влекущих изменение объема, цены работ, включая изменение срока выполнения работ, стороны согласовывают данные изменения путем заключения дополнительного соглашения.</p>
<p>1.4. Срок выполнения работ составляет {prod_days} ({_days_words(prod_days)}) рабочих дней, с момента согласования Технического проекта и получения Подрядчиком предварительной оплаты, в размере, указанном в разделе 3 Договора. Подрядчик вправе досрочно выполнить работу без получения предварительного согласия Заказчика.</p>
<p>1.4.1. В случае нарушения технологического процесса (поломка, остановка производственных линий, отсутствие энергоснабжения) по вине коммунальных и иных служб, нехватки сырья и (или) рабочей силы, если эти обстоятельства непосредственно повлияли на возможность надлежащего исполнения Подрядчиком своих обязательств по настоящему Договору, срок изготовления мебели переносится соразмерно времени, в течение которого действовали такие обстоятельства. В случае переноса сроков изготовления мебели по причине возникновения таких обстоятельств, Подрядчик уведомляет Заказчика о сроках переноса изготовления мебели в письменном виде по имеющимся средствам связи (социальные мессенджеры, электронная почта) с последующим заключением дополнительного соглашения об увеличении сроков по договору.</p>
<p>1.4.2. Если Заказчик не обеспечил возможность проезда Подрядчика к месту разгрузки мебели, Подрядчик не несет ответственности за сроки доставки и разгрузки мебели и оставляет за собой право выставить счет за компенсацию дополнительных затрат, понесенных в связи с доставкой мебели, который Заказчик обязуется оплатить в течение 3 (трех) рабочих дней.</p>
<p>1.5. Работы, выполняемые Подрядчиком, предназначены удовлетворять бытовые или другие личные потребности Заказчика.</p>
<p>1.6. Заказчик проинформирован о том, что допускаются различия оттенка цвета и текстуры покрытий мебели и (или) ее элементов, по сравнению с образцами, в пределах одного цветового тона, что не будет являться нарушением условий договора и основанием для предъявления соответствующих претензий Подрядчику.</p>
<p>1.7. Мебель изготавливается согласно индивидуальному заказу Заказчика (Технический проект). Подписанием данного договора Заказчик подтверждает, что он согласовал все характеристики мебели, включая, но не ограничиваясь: количество, размер, форма, габариты, материал, расцветка, комплектация, отделка, фурнитура, крепления и т.д. В связи с этим, на данный Договор не распространяются положения Закона РФ «О защите прав потребителей», в части возможности реализации потребителем права заменить товар или возвратить товар надлежащего качества.</p>

<p class="sec">2. РАЗРАБОТКА И СОГЛАСОВАНИЕ ТЕХНИЧЕСКОГО ПРОЕКТА</p>
<p>2.1. Технический проект разрабатывается Подрядчиком в течение 10 (десяти) рабочих дней с момента получения Подрядчиком предварительной оплаты, в размере, указанном в разделе 3 Договора.</p>
<p>2.2. Заказчик в течение 5 (пяти) рабочих дней с момента перечисления предварительной оплаты обеспечивает доступ Подрядчика в помещение, в котором будет установлена мебель для проведения замеров помещения.</p>
<p>2.3. Заказчик не позднее 5 (пяти) рабочих дней со дня получения от Подрядчика Технического проекта согласовывает, путем проставления личной подписи и даты согласования, или направляет мотивированный отказ от согласования Технического проекта в редакции Подрядчика с предложением своих замечаний/корректировок. В случае направления Заказчиком своих предложений/замечаний Подрядчик в течение 3 (трех) рабочих дней вносит изменения в Технический проект и направляет его на согласование.</p>
<p>Стороны согласовали допустимость не более 2 (двух) кругов правок в представленный Подрядчиком Технический проект.</p>
<p>2.4. Все изменения или дополнения после подписания Сторонами Технического проекта по инициативе Заказчика недопустимы, за исключением изменений и дополнений, которые Подрядчик признает существенными (изменение параметров помещения по обстоятельствам, не зависящим от Заказчика и т.п.).</p>
<p>2.5. Заказчик обязуется в течение 3 (трех) рабочих дней с момента наступления обстоятельств, в силу которых возникла необходимость внесения изменений или дополнений в Технический проект, направить Подрядчику письменное уведомление с указанием перечня таких изменений или дополнений, а также обосновать причины внесения изменений или дополнений в Технический проект.</p>
<p>2.6. Подрядчик в течение 5 (пяти) рабочих дней с момента получения уведомления о внесении изменений или дополнений в Технический проект сообщает Заказчику о возможности или невозможности таких изменений. В случае, если изменения или дополнения Технического проекта повлекут увеличение стоимости, сроков изготовления и иных условий, стороны внесут соответствующие изменения в договор путем заключения дополнительного соглашения.</p>

<p class="sec">3. СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЕТОВ</p>
<p>3.1. Общая стоимость работ, подлежащих выполнению по настоящему Договору складывается на основании Калькуляции (Приложение №2 к настоящему договору) и составляет <strong>{total:,.0f} ({total_words})</strong> рублей. НДС не облагается на основании ст. 346.11 НК РФ. В стоимость работ включается стоимость материалов Подрядчика, из которых производится работы.</p>
<p>3.2. Оплата работ осуществляется в следующем порядке:</p>
{pay_html}
<p>3.3. Оплата производится безналичным расчетом на счет Подрядчика либо наличными денежными средствами в кассу.</p>
<p>3.4. Обязательство Заказчика по безналичной оплате считается исполненным в момент зачисления денежных средств на счет Подрядчика, указанный в реквизитах.</p>

<p class="sec">4. ПРАВА И ОБЯЗАННОСТИ СТОРОН</p>
<p><strong>4.1. Подрядчик обязан:</strong></p>
<p>4.1.1. Выполнить работу по Договору согласно Техническому проекту и передать Заказчику мебель в установленный срок.</p>
<p>4.1.2. Уведомить Заказчика о готовности мебели по электронной почте, путем обмена сообщениями (Telegram, WhatsUp, СМС и т.д.) на телефонный номер Стороны согласно реквизитам, указанным в Договоре.</p>
<p>4.1.3. Устранить недостатки, выявленные Заказчиком по результатам приемки работ.</p>
<p>4.1.4. Предоставить Заказчику необходимую и достоверную информацию о предлагаемой работе, ее видах и особенностях, о цене и форме оплаты, а также сообщить Заказчику по его просьбе другие относящиеся к Договору и соответствующей работе сведения. Подтверждением предоставления Заказчику указанной информации и его ознакомление с ней будет являться подписание Заказчиком настоящего Договора.</p>
<p>4.1.5. Предупредить Заказчика о возможных неблагоприятных для него последствиях выполнения его указаний о способе исполнения работы, а также иных не зависящих от Подрядчика обстоятельств, которые грозят годности или прочности результатов выполняемой работы либо создают невозможность ее завершения в срок.</p>
<p><strong>4.2. Заказчик обязан:</strong></p>
<p>4.2.1. Согласовать и подписать Технический проект в срок, установленный Договором.</p>
<p>4.2.2. Оплатить стоимость работ в соответствии с условиями настоящего Договора.</p>
<p>4.2.3. Принять результат работ путем подписания Акта выполненных работ.</p>
<p>4.2.4. Обеспечить сохранность помещения, в котором будет установлена мебель в том виде, в котором помещение было на момент проведения замеров Подрядчиком.</p>
<p>4.2.5. Проводить ремонтно-отделочные работы в помещении, где будет размещена мебель, в строгом соответствии с Техническим проектом и только после получения письменного согласия Подрядчика на проведение таких работ.</p>
<p>4.2.6. Предоставить Подрядчику информацию необходимую для выполнения работ по Договору и отражении ее в Техническом проекте, в том числе, но не ограничиваясь:</p>
<p>- систему крепления мебели. При системе крепления мебели без врезки задней стенки ДВП между стеной и мебелью (например, при монтаже верхних шкафов при зазоре 10-20мм), что позволяет скрывать строительные дефекты стен. При системе крепления мебели с «утопленной» задней стенкой - зазор исключен, но для обеспечения такого крепления требуется наличие идеально ровных стен без строительных дефектов (впадин, бугров и т.д.).</p>
<p>- места при установке мебели, в которых зазоры не приемлемы для Заказчика.</p>
<p>- свес кухонной столешницы у лицевой стороны и у стены.</p>
<p>- месторасположение и размеры техники (холодильник, вытяжка, варочная панель или газовая плита и т.д.), а также указать свои предпочтения по выбору техники для возможности получения у Подрядчика консультации по подбору мебели, во избежание порчи мебели и (или) техники при эксплуатации техники и мебели.</p>
<p>- особенности отделки и (или) ремонта помещения (неровные стены, полы разного уровня, навесной потолок, карнизы, наличие осветительных приборов, наличие труб, стояков, радиаторов и т.д.</p>
<p>- элементы декора, фотопечать, цветовые решения.</p>
<p><strong>4.3. Подрядчик вправе:</strong></p>
<p>4.3.1. Требовать подписания Заказчиком Акта выполненных работ в течение 5 (пяти) календарных дней с даты передачи мебели.</p>
<p>4.3.2. Требовать своевременной оплаты работ в соответствии с п. 3.2 настоящего Договора.</p>
<p>4.3.3. Запрашивать у Заказчика информацию необходимую для надлежащего исполнения Договора.</p>
<p>4.3.4. Привлекать третьих лиц для исполнения обязательств по Договору. Подрядчик несет ответственность за действия/бездействие третьих лиц, выполняющих работу по настоящему Договору, как за свои собственные.</p>
<p>4.3.5. В случаях, когда исполнение работы по Договору стало невозможным вследствие виновных действий или упущений Заказчика, Подрядчик сохраняет право на уплату ему указанной в Договоре цены с учетом выполненной части работы.</p>
<p>4.3.6. Не приступать к работе, а начатую работу приостановить в случаях, когда нарушение Заказчиком своих обязанностей, установленных в п. 4.2 Договора препятствует исполнению Договора, а также при наличии иных обстоятельств, очевидно свидетельствующих о том, что исполнение указанных обязанностей не будет произведено в установленный срок.</p>
<p>4.3.7. Досрочно выполнить работы и требовать от Заказчика принять результат работ и произвести его оплату.</p>
<p>4.3.8. В рекламных целях сделать фотоснимки изготовленной и установленной на месте мебели.</p>
<p><strong>4.4. Заказчик вправе:</strong></p>
<p>4.4.1. Выбрать модель мебели, цвет, компоновку, дизайн, материалы, фурнитуру из которых будет выполнена работа.</p>
<p>4.4.2. Отозвать согласие на обработку своих персональных данных Подрядчиком.</p>

<p class="sec">5. ГАРАНТИЯ И КАЧЕСТВО ВЫПОЛНЕННЫХ РАБОТ</p>
<p>5.1. Подтверждением качества мебели со стороны Подрядчика являются сертификаты соответствия, паспорта на товар, инструкции по эксплуатации и иная документация, устанавливающая требования к качеству мебели на момент доставки в соответствии с условиями настоящего Договора.</p>
<p>Гарантийный срок на мебельную фурнитуру (петли, выдвижные механизмы, выдвижные корзины, подъемные механизмы, сушки для посуды и т.п.) составляет 6 месяцев, гарантийный срок на мебель (корпус, фасады, столешницы) составляет 24 месяца, которые исчисляются с даты подписания Сторонами акта сдачи-приемки выполненных Работ, при условии надлежащей эксплуатации мебели. Срок службы изделий – 5 лет. Гарантийный срок эксплуатации Мебели и срок службы Мебели не распространяются на светильники, таймеры и другие встроенные бытовые приборы. Гарантийный срок таких приборов и срок определяется в соответствии с гарантией производителя.</p>
<p>5.2. Мебель, поставляемая в соответствии с настоящим Договором, должна быть осмотрена Заказчиком (уполномоченным представителем Заказчика) на предмет внешних повреждений без нарушения упаковки непосредственно при получении мебели от Подрядчика на складе последнего, либо при доставке по адресу ее установки.</p>
<p>5.3. Приемка Мебели по количеству и качеству (внешние недостатки) осуществляется:</p>
<p>5.3.1. При самовывозе: Заказчиком в момент приемки мебели на складе Подрядчика и подписания акта выполненных работ;</p>
<p>5.3.2. При доставке: Подрядчиком (грузоперевозчиком Подрядчика): Заказчиком (уполномоченным им грузоперевозчиком или грузополучателем) в момент приемки мебели от Подрядчика уполномоченного на это акта оказания услуг.</p>
<p>5.3.3. Если в будущем сборка мебели будет осуществляться силами ООО «Интерьерные решения» на основании отдельно заключенного договора, то Заказчик обязуется вскрыть упаковку в присутствии уполномоченного лица ООО «Интерьерные решения», осуществляющего сборку мебели.</p>
<p>5.4. При обнаружении нарушения целостности упаковки, несоответствия мебели в момент приемки по качеству и/или количеству, Заказчик делает отметку в акте выполненных работ/оказания услуг. Мебель, в отношении которой у Заказчика имеются замечания, Заказчик принимает на ответственное хранение, о чем делается отметку в акте выполненных работ/оказания услуг. Плата за хранение в этом случае не взимается. Заказчик не вправе использовать (продавать, производить монтаж и т.д.) такую мебель.</p>
<p>5.4.1. Претензия должна содержать информацию о номере договора, дате приемки мебели, дате выявления недостатков, наименование и количество мебели с недостатками, описание недостатков, местонахождение мебели с недостатками, пожелания Заказчика по урегулированию претензии. К претензии Заказчик обязан приложить фотографии мебели с недостатками, на которых отчетливо различимы все выявленные недостатки. Несоблюдение Заказчиком вышеуказанных требований, является основанием для отклонения претензии Подрядчиком.</p>
<p>5.5. В случае несоответствия количества, ассортимента, качества (внешние недостатки) мебели, возникших по вине Подрядчика, Подрядчик обязуется доставить или обеспечить замену некачественной мебели в следующей поставке или в срок, согласованный Сторонами.</p>
<p>5.6. Заказчик лишается права ссылаться на недостатки мебели, если: мебель была принята без проверки на предмет внешних повреждений; без проверки на предмет соответствия ассортимента; мебель была принята без проверки на предмет соответствия количества.</p>
<p>5.7. В ходе рассмотрения претензии Подрядчик вправе запросить у Заказчика, а Заказчик обязан в течение 3 (трех) дней с момента получения запроса предоставить Подрядчику дополнительные фотографии либо видеоматериалы мебели с недостатками, образцы мебели с недостатками. Кроме того, Подрядчик вправе выехать к месту нахождения мебели с недостатками для проведения совместной приемки мебели.</p>
<p>5.8. Подрядчик должен направить письменный ответ на претензию в течение 10 (Десяти) календарных дней с момента получения претензии.</p>
<p>5.9. В случае если претензия признана обоснованной, Подрядчик обязуется соразмерно уменьшить цену некачественной мебели или заменить некачественную мебель или ее часть (комплектующие) на качественную в срок, согласованных с Заказчиком.</p>
<p>5.10. Гарантия не распространяется на недостатки мебели, которые возникли после передачи мебели Заказчику вследствие нарушения Заказчиком правил пользования мебелью, ее ненадлежащей эксплуатации, ненадлежащего или небрежного хранения, ненадлежащего и или небрежного обслуживания, чрезмерной нагрузки на мебель, повреждения, попадания на поверхность изделия едких веществ и/или жидкостей, использования мебели не по назначению, недостаточного и/или неправильного монтажа, произведенного третьими лицами, то есть Не Подрядчиком, а также на неисправности, возникшие из-за несоблюдения технических инструкций производителя, касающихся порядка и условий использования соответствующего вида мебели, на дефекты, возникшие вследствие естественного износа мебели. Гарантия также не распространяется на недостатки мебели, которые возникли после передачи мебели Заказчику вследствие действий третьих лиц, либо обстоятельств непреодолимой силы (форс-мажора).</p>

<p class="sec">6. ПОРЯДОК ПРИЕМКИ ВЫПОЛНЕННЫХ РАБОТ</p>
<p>6.1. По факту выполнения работ Подрядчик представляет Заказчику Акта выполненных работ в двух экземплярах по форме, согласованной в Приложении № 3 к Договору. При необоснованном отказе одной из сторон от подписания акта в нем делается отметка об этом и акт подписывается другой стороной в одностороннем порядке.</p>
<p>6.2. Передача мебели производится по Акту выполненных работ, который составляется в момент передачи мебели.</p>
<p>6.3. Приемка изготовленной мебели производится:</p>
<p>- в случае самовывоза Заказчиком по адресу склада Подрядчика;</p>
<p>- в случае доставки мебели Подрядчиком в рамках заключенного между Заказчиком и Подрядчиком дополнительного договора доставки мебели по адресу Заказчика.</p>
<p>6.4. В случае самовывоза Заказчиком мебели Подрядчик обеспечивает погрузку мебели на своем складе в транспортное средство Заказчика (перевозчика Заказчика), а Заказчик - их транспортировку и выгрузку.</p>
<p>6.5. Заказчик обязан в течение 3 (трех) календарных дней с даты получения уведомления от Подрядчика о готовности мебели к поставке, прибыть на склад Подрядчика для выборки мебели, в установленные в уведомлении дату и время.</p>
<p>6.6. После погрузки мебели на транспортное средство Заказчика Подрядчик не несет ответственности за повреждения мебели, произошедшие вследствие нарушения Заказчиком правил транспортировки, выгрузки, условий хранения и эксплуатации.</p>
<p>6.7. В случае доставки мебели Подрядчиком Заказчик обязуется подготовить помещение для приемки, подъезд и проход к нему, создать условия для сохранности мебели при приемке, обеспечить бесплатной парковкой автомобиль Подрядчика на период отгрузки мебели по адресу Заказчика.</p>
<p>6.8. При уклонении Заказчика от получения мебели, Подрядчик бесплатно хранит мебель в течение 7(семи) дней. С восьмого дня Подрядчик принимает мебель на ответственное хранение, о чем уведомляет Заказчика и вправе запросить оплату за хранение в сумме 500 рублей за каждый день. Момент окончания ответственного хранения - передача мебели по акту приема-передачи. Дополнительных документов о принятии на ответственное хранение, кроме одностороннего уведомления, между сторонами не требуется.</p>
<p>6.9. В случае наступления обстоятельств, объективно препятствующих поставке мебели в согласованные сторонами сроки, Подрядчик вправе в одностороннем порядке изменить срок отгрузки мебели. Указанными обстоятельствами могут быть действия третьих лиц, которые способствуют исполнению договора, погодные условия и иные обстоятельства.</p>
<p>6.10. При получении мебели Заказчик обязан осмотреть ее, проверить соответствие качества, количества и комплектности условиям Договора.</p>
<p>6.11. Риск случайной гибели или повреждения мебели, а также право собственности переходит от Подрядчика к Заказчику в момент подписания соответствующих документов о приемке (товарная накладная, универсальный передаточный документ и т.д.). В случае если доставку мебели осуществляет транспортная компания, то риск случайной гибели или повреждения мебели, а также право собственности на нее переходит к Заказчику с момента сдачи Подрядчиком мебели транспортной компании.</p>
<p>6.12. При наличии заказа на услуги монтажа мебели, который оформляется дополнительным договором между Подрядчиком и Заказчиком мебель монтируется Подрядчиком.</p>
<p>6.13. При отсутствии заказа на монтаж мебели, монтаж (установка) мебели по месту осуществляется Заказчиком самостоятельно и за свой счет. В этом случае Подрядчик не несёт ответственность за качество монтажных работ и возможные недостатки мебели, возникшие в результате самостоятельного монтажа Заказчиком.</p>

<p class="sec">7. ОТВЕТСТВЕННОСТЬ СТОРОН</p>
<p>7.1. Стороны несут ответственность за неисполнение или ненадлежащее исполнение своих обязательств по Договору в соответствии с Законом Российской Федерации от 7 февраля 1992 г. № 2300-1 «О защите прав потребителей» и иными правовыми актами, принятыми в соответствии с ним.</p>
<p>7.2. За нарушение сроков оплаты услуг Подрядчик вправе потребовать с Заказчика уплаты неустойки (пени) за каждый день просрочки в размере 0,1 % от суммы задолженности.</p>
<p>7.3. Подрядчик не несет ответственности за невыполнение обязательств по Договору, если оно вызвано неисполнением соответствующих обязанностей Заказчика.</p>
<p>7.4. В случае отказа Заказчика от мебели надлежащего качества, изготовленной по индивидуальному заказу (эскизу), которая имеет нестандартные размеры, цвет, форму, и ее дальнейшая реализация невозможна, Стороны признают, что размер убытков Подрядчика составляет стоимость соответствующей (не вывезенной) партии мебели.</p>
<p>7.5. При разгрузке мебели силами Заказчика материальную ответственность за сохранность груза и автотранспортного средства несет Заказчик. В случае нанесения ущерба автотранспортному средству Подрядчика (Перевозчика Подрядчика), представитель Заказчика и Подрядчика (или представитель транспортной организации) составляют акт о причиненных повреждениях и причинах возникновения повреждений. При отказе представителя Заказчика от подписи в акте, акт составляется в одностороннем порядке Подрядчиком (представителем транспортной организации) и является основанием для возмещения убытков.</p>

<p class="sec">8. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ</p>
<p>8.1. Все споры или разногласия, возникающие между Сторонами по настоящему Договору или в связи с ним, разрешаются путем переговоров между ними.</p>
<p>8.2. До передачи спора на рассмотрение суда Стороны предусматривают обязательный претензионный порядок урегулирования разногласий. Направление претензии возможно по электронной почте, путем обмена сообщениями (Telegram, WhatsUp, СМС и т.д.) на телефонный номер Стороны согласно реквизитам, указанным в Договоре. Срок ответа на претензию Сторон устанавливают в 10 (Десять) рабочих дней с момента получения претензии заинтересованной Стороны.</p>
<p>8.3. В случае невозможности разрешения разногласий путем переговоров споры или разногласия, возникающие между Сторонами, решаются в судебном порядке по месту нахождения Подрядчика.</p>

<p class="sec">9. СРОК ДЕЙСТВИЯ ДОГОВОРА</p>
<p>9.1. Договор вступает в силу с момента его подписания обеими Сторонами и действует до полного исполнения Сторонами принятых на себя обязательств.</p>
<p>9.2. Изменения и дополнения к Договору принимаются по обоюдному соглашению Сторон, путем подписания Дополнительного соглашения к Договору.</p>
<p>9.3. Договор может быть расторгнут досрочно по письменному соглашению Сторон, в одностороннем порядке в случаях, предусмотренных действующим законодательством Российской Федерации.</p>
<p>9.4. В случае одностороннего отказа Заказчика от Договора Заказчик обязуется оплатить Подрядчику фактически понесенных им расходов, связанных с исполнением обязательств по данному Договору.</p>

<p class="sec">10. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ</p>
<p>10.1. При заключении Договора и подписания других документов в рамках его исполнения, Стороны согласны на факсимильное воспроизведение подписи («факсимиле»), уполномоченных ими лиц, а также на обмен копиями документов. Стороны признают юридически значимой переписку между сторонами по электронной почте, а также переписку по номеру телефона, указанному в настоящем договоре в любых приложениях к нему (для обмена сообщениями: Telegram, WhatsUp, СМС и т.д.). Направленные таким образом документы считаются подписанными простой электронной подписью и признаются Сторонами равнозначными бумажным, подписанным собственноручной подписью Сторон.</p>
<p>10.2. Стороны признают юридическую силу всех документов, требований, уведомлений, претензий, извещений, оформленных должным образом и направленных друг другу в электронном виде по указанным адресам электронной почты и телефонам во исполнение настоящего Договора. Стороны договорились, что сообщение при такой форме коммуникации считается доставленным тогда, когда приложение, через которое производится коммуникация подтвердит факт направления сообщения через свой внутренний интерфейс. При этом и Заказчик, и Подрядчик обязуются регулярно просматривать указанные источники коммуникации. Стороны договорились, что сообщение считается доставленным и в тех случаях, если оно поступило лицу, которому оно направлено, но по обстоятельствам, зависящим от него, не было ему вручено или адресат не ознакомился с ним.</p>
<p>10.3. Каждая Сторона обязуется обеспечивать конфиденциальность полученной ею в связи с заключением или исполнением Договора от другой Стороны информации ограниченного доступа.</p>
<p>К информации ограниченного доступа относится информация в любой форме, содержащая коммерческую, служебную или иную охраняемую законом тайну, секрет производства, персональные данные. При передаче информации Сторона обязана явно обозначить статус информации как информации ограниченного доступа, сделав отметку об этом на материальных носителях информации, в акте приема-передачи документации или иным образом до передачи информации.</p>
<p>10.4 Заказчик дает свое согласие на обработку своих персональных данных, а именно: на действия, совершаемые с использованием средств автоматизации или без использования таких средств, включая сбор, запись, систематизацию, накопление, хранение, уточнение (обновление, изменение), извлечение, использование, передачу (распространение, предоставление, доступ), обезличивание, блокирование, удаление, уничтожение его персональных данных, {co_name} (ОГРН: {co_ogrn}, ИНН: {co_inn}), расположенному по адресу: {co_address}.</p>
<p>Цель обработки персональных данных: исполнение настоящего договора. Заказчик дает свое согласие на использование следующих персональных данных: фамилия, имя, отчество, паспортные данные, адрес места жительства, фотографии изделий на объекте Заказчика; номер телефона, адрес электронной почты. Согласие предоставляется на срок действия настоящего договора, а после прекращения договора - в течение 12 месяцев с даты подписания Сторонами акта сдачи-приемки выполненных Работ.</p>
<p>Настоящее условие договора может быть изменено Заказчиком - субъектом персональных данных, в любой момент в одностороннем порядке путем отзыва согласия на обработку персональных данных. Отзыв согласия на обработку персональных данных осуществляется посредством составления письменного документа, который может быть направлен в адрес Подрядчика почтовым отправлением с уведомлением о вручении, либо вручен лично под расписку представителю Подрядчика.</p>
<div style="display:flex;justify-content:space-between;margin:20px 0;border-top:1px solid #000;padding-top:8px">
<span>______________________ /</span>
<span>______________________ /</span>
</div>
<p style="display:flex;justify-content:space-between;text-indent:0" class="no-indent"><span style="text-align:center;width:48%">(подпись)</span><span style="text-align:center;width:48%">(расшифровка подписи от руки)</span></p>
<p>10.5. К настоящему Договору прилагаются и являются неотъемлемой частью следующие приложения:</p>
<p class="no-indent">1. Технический проект.</p>
<p class="no-indent">2. Калькуляция работ.</p>
<p class="no-indent">3. Правила эксплуатации корпусной мебели.</p>
<p class="no-indent">4. Образец Акта выполненных работ</p>

<p class="sec">11. РЕКВИЗИТЫ СТОРОН</p>
<table>
<tr><th style="width:50%">ПОДРЯДЧИК</th><th style="width:50%">ЗАКАЗЧИК</th></tr>
<tr>
<td style="vertical-align:top;padding:10px 12px">
{co_requisites}<br>
<br>
Менеджер: <strong>{manager_line}</strong><br>
<br>
Подпись: <span class="ul" style="min-width:140px">&nbsp;</span>&nbsp;/&nbsp;<span class="ul" style="min-width:120px">&nbsp;</span><br>
<span style="font-size:9pt;color:#555">(подпись) &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; (расшифровка)</span><br>
М.П.{stamp_sig_html}
</td>
<td style="vertical-align:top;padding:10px 12px">
<strong>{fname}</strong><br>
Паспорт: {_passport_str(c)}<br>
Выдан: {c.get("passport_issued_by") or "___________"}<br>
{_fmt_date(c.get("passport_issued_date") or "")}, код {c.get("passport_dept_code") or "___"}<br>
Адрес: {_reg_address(c)}<br>
Тел.: {c.get("phone") or "___________"}<br>
Email: {c.get("email") or "___________"}<br>
<br>
Подпись: <span class="ul" style="min-width:140px">&nbsp;</span>&nbsp;/&nbsp;<span class="ul" style="min-width:120px">&nbsp;</span><br>
<span style="font-size:9pt;color:#555">(подпись) &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; (расшифровка)</span>
</td>
</tr>
</table>
</div></body></html>'''

    elif doc_type == 'act':
        total = float(c.get('total_amount') or 0)
        total_words = _num_to_words(total) if total else '___________'
        style = _doc_style('Акт выполненных работ', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Акт выполненных работ к договору №{contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 4 к договору бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Акт выполненных работ»</h1>
<div class="city-date"><span></span><span>от «____» ______________ 20____ г.</span></div>
<p style="text-align:center">(ФОРМА)</p>
<p class="no-indent">{co_name}, в лице {co_dir_pos}а <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, действующего на основании {poa_str}, именуемый в дальнейшем «Подрядчик» и гр. <strong>{fname}</strong>, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», подписали настоящий Акт выполненных работ о нижеследующем:</p>
<p class="no-indent">1. Подрядчик изготовил для Заказчика мебель по договору бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>:</p>
<table><tr><th style="width:5%">№</th><th>Наименование мебели, включая ее элементы</th><th style="width:13%">Ед. измерения</th><th style="width:12%">Кол-во изделий</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td style="text-align:center">1</td><td>Кухонный гарнитур</td><td style="text-align:center">Шт.</td><td style="text-align:center">1</td><td style="text-align:center">(сумма прописью)</td></tr>
</table>
<p style="margin:20px 0 10px 0">&nbsp;</p>
<p class="no-indent">2. Комплектность, количество, вид, характеристики мебели соответствуют условиям договора. Визуальный осмотр мебели на предмет повреждений, царапин, сколов, трещин и других недостатков произведен Заказчиком. Фурнитура (петли, выдвижные механизмы, подъемные механизмы и т.д.) работает исправно. Заказчик претензий по объему, качеству, результату и срокам выполнения работ: <strong>не имеет / имеет</strong> (ненужное зачеркнуть).</p>
<p style="margin:20px 0">&nbsp;</p>
<p style="margin:20px 0">&nbsp;</p>
<p style="margin:20px 0">&nbsp;</p>
<p class="no-indent">3. В случае наличия замечаний Заказчик, после подписания акта, в праве требовать устранения замечаний отраженных в данном акте.</p>
<p class="no-indent">4. Настоящий акт подписан в 2 (двух) экземплярах по одному для каждой из Сторон.</p>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="height:60px"></td><td>{fname}</td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'tech':
        style = _doc_style('Технический проект', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Технический проект к договору №{contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 1 к договору бытового подряда на изготовление мебели № {contract_num} от {contract_date_full}</p>
<h1>«Технический проект»</h1>
<table style="margin-top:16px;font-size:10pt">
<tr>
  <td style="width:15%;font-weight:bold;border:1px solid #000;padding:4px 6px">Корпус:</td>
  <td style="width:35%;border:1px solid #000;padding:4px 6px">&nbsp;</td>
  <td style="width:15%;font-weight:bold;border:1px solid #000;padding:4px 6px">Столешница:</td>
  <td style="width:35%;border:1px solid #000;padding:4px 6px">&nbsp;</td>
</tr>
<tr>
  <td style="font-weight:bold;border:1px solid #000;padding:4px 6px">Фасад 1:</td>
  <td style="border:1px solid #000;padding:4px 6px">&nbsp;</td>
  <td style="font-weight:bold;border:1px solid #000;padding:4px 6px">Стеновая панель:</td>
  <td style="border:1px solid #000;padding:4px 6px">&nbsp;</td>
</tr>
<tr>
  <td style="font-weight:bold;border:1px solid #000;padding:4px 6px">Фасад 2:</td>
  <td style="border:1px solid #000;padding:4px 6px">&nbsp;</td>
  <td style="font-weight:bold;border:1px solid #000;padding:4px 6px">Подсветка &nbsp; Тип:</td>
  <td style="border:1px solid #000;padding:4px 6px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; Свет: &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</td>
</tr>
<tr>
  <td style="font-weight:bold;border:1px solid #000;padding:4px 6px">Фрезеровка:</td>
  <td colspan="3" style="border:1px solid #000;padding:4px 6px">&nbsp;</td>
</tr>
</table>
<div style="border:1px solid #000;min-height:150mm;margin-top:10px"></div>
<p style="margin-top:20px;font-style:italic;text-indent:0">Подписывая Технический проект, Заказчик подтверждает, что ознакомлен с наименованием, качественными характеристиками, количеством, дизайном мебели и ему полностью понятны выполняемые Подрядчиком работы. Стороны согласовали, что мебель изготовлена специально для Заказчика по его индивидуальным параметрам. Приложение: бланк замера.</p>
<table style="margin-top:16px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding:6px">Менеджер</td><td>&nbsp;</td></tr>
<tr><td style="height:50px"><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
</table>
</div></body></html>'''

    elif doc_type == 'delivery':
        daddr = _delivery_address(c)
        floor_ = c.get('delivery_floor') or '___'
        elevator = c.get('delivery_elevator') or 'нет'
        entrance = c.get('delivery_entrance') or '___'
        delivery_date_str = _fmt_date(c.get('delivery_date') or '')
        dcost = float(c.get('delivery_cost') or 0)
        dcost_str = f'{dcost:,.0f} ({_num_to_words(dcost)})' if dcost else '<span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>'
        style = _doc_style('Договор доставки', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Договор доставки — {contract_num}</title>{style}</head><body><div class="page">
<h1>ДОГОВОР</h1>
<h2>на оказание услуг по доставке мебели</h2>
<div class="city-date"><span>г. {co_city}</span><span>{contract_date_full}</span></div>
<p class="no-indent">{co_name}, в лице {co_dir_pos}а <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, действующего на основании {poa_str}, именуемый в дальнейшем «Исполнитель», и <strong>{fname}</strong>, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», заключили настоящий Договор о нижеследующем:</p>
<p class="sec">1. ПРЕДМЕТ ДОГОВОРА</p>
<p>1.1. В соответствии с настоящим Договором Исполнитель обязуется доставить мебель в разобранном виде (далее – мебель) в пункт назначения и выдать его Заказчику или иному лицу, указанному в п. 1.4. договора, на условиях, указанных в настоящем Договоре, а Заказчик обязуется оплатить за доставку установленную плату в порядке и на условиях, указанных в настоящем Договоре.</p>
<p>1.2. Исполнитель при оказании услуг по настоящему Договору обязуется соблюдать нормы действующего законодательства Российской Федерации.</p>
<p>1.3. Адрес доставки: <strong>{daddr}</strong><br>Этаж: {floor_}&nbsp;&nbsp;&nbsp;&nbsp;Лифт: {elevator}</p>
<p>1.4. Получателем мебели является: <strong>{fname}</strong></p>
<p>1.5. Срок оказания услуг – с даты подписания настоящего договора до <strong>{delivery_date_str}</strong></p>
<p>1.6. Исполнитель вправе привлечь к исполнению своих обязательств третьих лиц.</p>
<p>1.7. Обязательства Исполнителя считаются выполненными со дня доставки мебели по адресу, указанному в Договоре, в день доставки.<br>В случае невозможности доставки мебели в срок, предусмотренный Договором, по вине Заказчика, мебель остается у Исполнителя на ответственном хранении.</p>
<p class="sec">2. ПРАВА И ОБЯЗАННОСТИ СТОРОН</p>
<p><strong>2.1. Исполнитель обязан:</strong></p>
<p>2.1.1. Отгрузить мебель Заказчику в день доставки.</p>
<p>2.1.2. Предоставить под погрузку автомобили в технически исправном состоянии, обеспечивающем сохранность мебели. В случае несвоевременной подачи транспортного средства, наступившей вследствие действия непреодолимых обстоятельств, по возможности, проинформировать об этом Заказчика, а при невозможности подачи транспортного средства в силу вышеуказанных причин, предоставить его в другое, согласованное с Заказчиком время без дополнительной оплаты.</p>
<p>2.1.3. Оказать при необходимости услуги по упаковке, маркировке, хранению, погрузке и выгрузке мебели.</p>
<p>2.1.4. Организовать оформление всех необходимых для осуществления доставки мебели документов.</p>
<p>2.1.5. Обеспечить доставку мебели строго по адресу, указанному в Договоре.</p>
<p>2.1.6. Проинформировать Заказчика о задержке доставки мебели при наличии соответствующих обстоятельств.</p>
<p><strong>2.2. Заказчик обязан:</strong></p>
<p>2.2.1. Оплатить Исполнителю стоимость услуг по доставке мебели.</p>
<p>2.2.2. Предоставить достоверные сведения об адресе доставки и иной информации, необходимой для оказания услуги (размеры дверных (оконных) проемов, предметы интерьера, техника в помещении и т.д.), что может повлечь невозможность разгрузки мебели в помещении и т.д.).</p>
<p>2.2.3. При получении мебели подписать Акт об оказанных услугах и документы о получении мебели.</p>
<p>2.2.4. Обеспечить места для парковки автотранспортных средств Исполнителя, для проведения погрузо-разгрузочных работ, не дальше, чем 10 метров от входа (в дом, подъезд и т.д.). В случае невозможности обеспечить места для парковки у входа, в том числе по причине плохих условий (снег, гололед, размытая грунтовая дорога и т.д.), затрудняющих подъезд к подъезду, стороны согласовывают дополнительную оплату, в зависимости от расстояния от места возможной остановки транспортного средства до входа в дом и (или) подъезда.</p>
<p>2.2.5. Освободить все проходы и проемы от предметов, препятствующих разгрузке мебели по адресу доставки.</p>
<p>2.2.6. Не удалять с упаковки мебели никакую информацию (этикетки, ярлыки, инструкции по сборке, если они наклеены снаружи и т.д.), необходимую для дальнейшей сборки мебели.</p>
<p>2.2.7. Если в будущем сборка мебели будет осуществляться силами {co_name}: на основании отдельно заключенного договора об оказании услуг, Заказчик обязан вскрыть упаковку в присутствии уполномоченного лица {co_name}, осуществляющего сборку мебели.<br>Если сборка мебели будет осуществляться Заказчиком самостоятельно, в том числе с привлечением третьих лиц, не имеющим отношения к {co_name}, Заказчик осуществляет осмотр содержимого упаковки на предмет видимых дефектов мебели в присутствии представителя Исполнителя.<br>Если Заказчик пожелает не производить распаковку мебели до момента окончания оказания услуг по настоящему Договору, он лишается права ссылаться впоследствии на недостатки, обнаруженные им до вскрытия упаковки.</p>
<p><strong>2.3. Заказчик вправе:</strong></p>
<p>2.3.1. В любое время проверять ход и качество оказания услуг, выполняемых Исполнителем, не вмешиваясь в его деятельность.</p>
<p>2.3.2. Проверить упаковку и её содержимое, с учетом соблюдения п. 2.2.7 договора, на предмет целостности, отсутствия повреждений различного характера.</p>
<p class="sec">3. СТОИМОСТЬ УСЛУГ И ПОРЯДОК РАСЧЕТОВ</p>
<p>3.1. Общая стоимость услуг, определяется на основании Приложения №1 и №2 к настоящему договору, НДС не облагается на основании ст. 346.11 НК РФ и указывается в Акте оказания услуг за фактически оказанные Исполнителем услуги.</p>
<p>3.2. Оплата услуг производится в следующем порядке: 100 % предоплата безналичным расчетом на счет Исполнителя либо наличными денежными средствами в кассу Исполнителя.</p>
<p>3.3. Обязательство Заказчика по безналичной оплате считается исполненным в момент зачисления денежных средств на счет Исполнителя, указанный в реквизитах.</p>
<p>3.4. Если Заказчик в согласованный день и время доставки не примет мебель, доставленную Исполнителем, то услуги считаются полностью оказанными. Стоимость и дата повторной доставки определяется Сторонами в дополнительном соглашении.</p>
<p>3.5. Стоимость ответственного хранения:<br>- первые 7 (семи) дней бесплатно,<br>- с 8 (восьмого) дня - 500 рублей за каждый день. Момент окончания ответственного хранения - передача мебели по акту приема-передачи. Дополнительных документов о принятии на ответственное хранение, кроме одностороннего уведомления, между сторонами не составляется.</p>
<p class="sec">4. УСЛОВИЯ ДОСТАВКИ</p>
<p>4.1. Доставка производится силами Исполнителя или привлекаемыми им третьими лицами в любой день по согласованию с Заказчиком. Точное время и дата доставки могут быть согласованы Сторонами посредством телефонной связи на телефонный номер Сторон согласно реквизитам, указанным в Договоре, посредством обмена сообщениями (WhatsApp, Telegram, СМС и т.д.), электронной почты, по усмотрению каждой из Сторон настоящего Договора.</p>
<p>4.2. Стоимость доставки и подъема устанавливается согласно Приложению № 1 и №2 к настоящему договору.</p>
<p>4.3. В случае, если по какой-либо причине Заказчик откажется и/или не сможет принять мебель в срок и/или по адресу, указанному в п. 1.3. настоящего Договора, в том числе по независящим от Заказчика данного Договора причинам или каким-либо иным основаниям и/или причинам, мебель может быть помещена исполнителем на хранение с возложением всех расходов, связанных с хранением мебели на Заказчика, при этом, срок передачи (доставки), установленный в п. 1.5. настоящего Договора, не будет считаться нарушенным Исполнителем и будет считаться автоматически измененным Сторонами настоящего Договора на дату фактической передачи мебели Заказчику.</p>
<p>4.4. При наступлении обстоятельств, указанных в п. 4.3. настоящего Договора, Заказчик обязуется своевременно уведомить Исполнителя о новой дате и времени возможной доставки не позднее чем за 3 (три) рабочих дня до предполагаемой даты доставки. Датой уведомления в данном случае будет считаться дата поступления указанного в настоящем пункте Договора письменного уведомления в адрес Исполнителя.</p>
<p>4.5. Обязательства Исполнителя по доставке считается выполненной, а мебель считается принятой с момента подписания документов о её передаче Заказчику.</p>
<p>4.6. Право собственности на мебель, а также риск случайной гибели или её порчи, переходят к Заказчику в момент получения мебели Заказчиком.</p>
<p>4.7. В случае, если доставка произведена в установленные настоящим Договором сроки, но мебель не была передана Заказчику по его вине, то обязательства Исполнителя относительно сроков передачи мебели считаются исполненным, а новая доставка производится в новые сроки, согласованные с Заказчиком после повторной оплаты Заказчиком стоимости услуг по доставке.</p>
<p>4.8. Заказчик обязуется письменно уведомить Исполнителя о наличии у Заказчика представителей, уполномоченных от его имени принять мебель в дату доставки и подписать документы об оказанной услуге и получении мебели. При этом действия представителя Заказчика по приёме и подписанию документов распространяются и подтверждаются Сторонами настоящего Договора как действия самого Заказчика.</p>
<p>4.9. Если по какой-либо причине, не зависящей от Исполнителя, Заказчик не может принять мебель, либо уклоняется от её принятия и не сообщил о невозможности принять её, то обязательства Исполнителя относительно сроков передачи мебели Заказчику и срока доставки считаются выполненными, что указывается в Акте оказания услуг, а дата повторной доставки.</p>
<p>4.10. Несоответствие размеров дверных проемов и площади помещений размерам мебели с учетом её упаковки, а также наличие предметов, препятствующих её транспортировке в помещение, не служит основанием для возврата мебели. Проведение работ по устранению данных недостатков, в том числе, снятие установка дверей, увеличение дверных проемов, устранение препятствующих вносу предметов, производится за отдельную плату, либо в разумный срок силами, средствами и за счет Заказчика.</p>
<p class="sec">5. ОТВЕТСТВЕННОСТЬ СТОРОН</p>
<p>5.1. Стороны несут ответственность за неисполнение или ненадлежащее исполнение обязательств по Договору, если не докажут, что надлежащее исполнение оказалось невозможным вследствие непреодолимой силы, то есть чрезвычайных и непредотвратимых при данных условиях обстоятельств, повлекших неисполнение или ненадлежащее исполнение Стороной настоящего Договора.</p>
<p>5.2. Заказчик несет полную ответственность за сведения об адресе доставки мебели. Риски доставки по неправильному адресу несет Заказчик.</p>
<p>5.3. В случае неисполнения либо ненадлежащего исполнения обязательств по доставке Стороны несут ответственность, установленную действующим законодательством, а также следующую ответственность:<br>- в случае неисполнения Заказчиком обязанности по оплате Исполнитель услуг, он вправе потребовать уплаты неустойки в размере 0,1 процента от неуплаченной суммы за каждый день просрочки;<br>- в случае неисполнения им обязанности по доставке мебели Исполнителем Заказчик вправе расторгнуть договор в одностороннем порядке и получить возврат аванса.</p>
<p>5.4. Стороны освобождаются от ответственности, если неисполнение ими своих обязательств произошло вследствие:<br>- обстоятельств непреодолимой силы;<br>- иных не зависящих от сторон причин.</p>
<p class="sec">6. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ</p>
<p>6.1. Все споры или разногласия, возникающие между Сторонами по настоящему договору или в связи с ним, разрешаются путем переговоров между ними.</p>
<p>6.2. При отказе одной из сторон от рассмотрения суда Стороны предусматривают обязательный претензионный порядок урегулирования разногласий. Направление претензии возможно по электронной почте, путем обмена сообщениями (Viber, WhatsApp, Telegram, СМС и т.д.) на телефонный номер Сторон согласно реквизитам, указанным в Договоре. Срок ответа на претензию Сторон устанавливается в 10 (Десять) рабочих дней с момента получения претензии заинтересованной Стороны.</p>
<p>6.3. В случае невозможности разрешения споров путем переговоров споры или разногласия, возникающие между Сторонами, решаются в судебном порядке по месту нахождения Исполнителя.</p>
<p class="sec">7. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ</p>
<p>7.1. Договор вступает в силу с момента его подписания Сторонами и действует до полного исполнения Сторонами всех принятых на себя обязательств.</p>
<p>7.2. При заключении Договора и подписания других документов в рамках его исполнения, Стороны согласны на факсимильное воспроизведение подписи («факсимиле»), уполномоченных лиц, а также на обмен электронными копиями документов. Стороны признают юридически значимой переписку Сторон по электронной почте, а также переписку по номеру телефона, указанному в настоящем договоре в любых приложениях для обмена сообщениями (Viber, WhatsApp, Telegram, СМС и т.д.). Направленные таким образом документы считаются подписанными простой электронной подписью и признаются Сторонами равнозначными бумажным, подписанным собственноручной подписью Сторон.</p>
<p>7.3. Стороны признают юридическую силу всех документов, требований, уведомлений, претензий, извещений, оформленных должным образом и направленных друг другу в электронном виде по указанным адресам электронной почты и телефонам по исполнению настоящего Договора. Стороны договорились, что сообщение при такой форме коммуникации считается доставленным тогда, когда приложение, через которое производится коммуникация подтверждает факт направления сообщения через свой внутренний интерфейс. При этом и Исполнитель и Заказчик обязуются регулярно просматривать указанные источники коммуникации. Стороны договорились, что сообщение считается доставленным и в тех случаях, если оно поступило лицу, которому оно направлено, но по обстоятельствам, зависимым от него, было ему вручено или адресат не ознакомился с ним.</p>
<p>7.4. Заказчик дает свое согласие на обработку своих персональных данных, а именно: на действия, совершаемые с использованием средств автоматизации или без использования таких средств, включая сбор, запись, систематизацию, накопление, хранение, уточнение (обновление, изменение), извлечение, использование, передачу (распространение, предоставление, доступ), обезличивание, блокирование, удаление, уничтожение его персональных данных, {co_name} (ОГРН: {co_ogrn}, ИНН: {co_inn}), расположенному по адресу: {co_address}.<br>Цель обработки персональных данных: исполнение настоящего договора. Заказчик дает свое согласие на использование следующих персональных данных: фамилия, имя, отчество, паспортные данные, адрес места жительства, фотографии изделий на объекте Заказчика; номер телефона, адрес электронной почты. Согласие предоставляется на срок действия настоящего договора, а после прекращения договора - в течение 12 месяцев с даты подписания Сторонами акта сдачи-приемки выполненных Работ.<br>Настоящее условие договора может быть изменено Заказчиком - субъектом персональных данных, в любой момент в одностороннем порядке путем отказа от согласия на обработку персональных данных. Отзыв согласия на обработку персональных данных осуществляется посредством составления письменного документа, который может быть направлен в адрес Исполнителя почтовым отправлением с уведомлением о вручении, либо вручен лично под расписку представителю Исполнителя.</p>
<div style="margin:20px 0;display:flex;gap:40px"><div style="flex:1;border-top:1px solid #000;padding-top:4px;text-align:center">(подпись)</div><div style="flex:1"></div><div style="flex:1;border-top:1px solid #000;padding-top:4px;text-align:center">(расшифровка подписи от руки)</div></div>
<p>7.5. К настоящему Договору прилагаются и являются неотъемлемой частью следующие приложения:<br>1. Калькуляция на выполнение услуг по доставке мебели.<br>2. Прайс на выполнение услуг по подъему мебели.<br>3. Образец Акта оказанных услуг.</p>
<p class="sec">8. РЕКВИЗИТЫ СТОРОН</p>
<table><tr><th style="width:50%">Исполнитель:</th><th style="width:50%">Заказчик:</th></tr>
<tr><td>{co_requisites}{stamp_sig_html}</td>
<td>Паспорт. Серия, номер: {_passport_str(c)}<br>Кем выдан:<br><br>Адрес прописки:<br><br>Телефон: {c.get("phone") or "___________"}<br><br>Предпочитаемый канал обмена сообщениями:</td></tr>
<tr><td style="padding-top:20px">Менеджер: <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td></td></tr>
<tr><td style="height:60px"></td><td></td></tr>
<tr><td style="height:60px"></td><td>{fname}</td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
<div style="height:40px"></div>
<p class="no-indent" style="text-align:right">Приложение № 1 к договору на оказание услуг по доставке мебели от {contract_date_full}</p>
<h2>«Калькуляция на выполнение услуг по доставке мебели»</h2>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Доставка мебели в территориальных границах г. Саратова и г. Энгельса</td><td style="text-align:center">1 услуга</td><td style="text-align:center">1</td><td style="text-align:right">0</td><td style="text-align:right">0</td></tr>
<tr><td>Доставка мебели за пределы территориальных границ г. Саратова и г. Энгельса</td><td style="text-align:center">1 км</td><td style="text-align:center">70</td><td style="text-align:right"></td><td style="text-align:right">0</td></tr></table>
<div style="height:30px"></div>
<p class="no-indent" style="text-align:right">Приложение № 2 к договору на оказание услуг по доставке мебели от {contract_date_full}</p>
<h2>«Прайс на выполнение услуг по подъёму мебели»</h2>
<p><em>Дополнительные услуги, по факту оказания</em></p>
<p><strong>Подъем мебели (полный комплект), при отсутствии лифта и занос мебели, при невозможности парковки автомашины вплотную к подъезду</strong></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Подъем. Квадратура корпуса до 20 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">550</td><td></td></tr>
<tr><td>Подъем. Квадратура корпуса от 20 до 25 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">650</td><td></td></tr>
<tr><td>Подъем. Квадратура корпуса более 25 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">750</td><td></td></tr>
<tr><td>Перемещение мебели вручную в случае невозможности подъезда автомашины к месту разгрузки за каждые</td><td style="text-align:center">1 м</td><td></td><td style="text-align:right">30</td><td></td></tr></table>
<p><em>Подъем комплектующих отдельно</em></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:12%">Ед. измерения</th><th style="width:12%">Количество упаковок</th><th style="width:12%">Количество этажей</th><th style="width:16%">Цена за ед. измерения в руб.</th><th style="width:16%">Стоимость в руб.</th></tr>
<tr><td>Подъем столешницы</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">350</td><td></td></tr>
<tr><td>Подъем стеновой панели</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">250</td><td></td></tr>
<tr><td>Подъем крупных частей корпуса. (боковины кухонного пенала, боковины шкафов-купе и т.д.). Упаковка по 2 шт.</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">250</td><td></td></tr>
<tr><td>Подъем длинных комплектующих (цоколь, профильная ручка, труба – целиком)</td><td style="text-align:center">1 комплект / 1 этаж</td><td></td><td></td><td style="text-align:right">60</td><td></td></tr>
<tr><td>Подъем дверей-купе</td><td style="text-align:center">1 дверь / 1 этаж</td><td></td><td></td><td style="text-align:right">150</td><td></td></tr>
<tr><td colspan="5" style="text-align:right;font-weight:bold">Цена услуг составляет:</td><td></td></tr></table>
</div></body></html>'''

    elif doc_type == 'assembly':
        daddr = _delivery_address(c)
        assembly_days = int(c.get('assembly_days') or 1)
        delivery_date_str = _fmt_date(c.get('delivery_date') or '')
        acost = float(c.get('assembly_cost') or 0)
        acost_str = f'{acost:,.0f} ({_num_to_words(acost)})' if acost else '<span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>'
        style = _doc_style('Договор монтажа', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Договор монтажа — {contract_num}</title>{style}</head><body><div class="page">
<h1>ДОГОВОР</h1>
<h2>на выполнение работ по монтажу и сборке мебели</h2>
<div class="city-date"><span>г. {co_city}</span><span>{contract_date_full}</span></div>
<p class="no-indent">{co_name}, в лице {co_dir_pos}а <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, действующего на основании доверенности № <span class="ul" style="min-width:40px">&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, именуемый в дальнейшем «Подрядчик» и гр. <strong>{fname}</strong>, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», заключили настоящий Договор о нижеследующем:</p>
<p class="sec">1. ПРЕДМЕТ ДОГОВОРА</p>
<p>1.1. Подрядчик обязуется выполнить работу по монтажу и сборке мебели (далее – Работы), приобретенной у {co_name} согласно договору бытового подряда на изготовление мебели № {contract_num} от {contract_date_full} года, а Заказчик обязуется принять и оплатить результат Работ.<br>В перечень Работ входит: расстановка, расстановка в порядке и месте, определенном Заказчиком, и сборка мебели в соответствии с инструкциями, схемами и иной технической документацией, в помещении.<br>В перечень Работ не включаются сантехнические (установка смесителей, сифонов, кранов и т.д.), электромонтажные работы, а также работы по подключению плит, духовок, стиральных и посудомоечных машин, вытяжек и другой бытовой техники.<br>Подключение бытовой техники, приобретенной Заказчиком вместе с мебелью, осуществляют сервисные центры на основании своих расценок.</p>
<p>1.2. Работы выполняются по адресу: <strong>{daddr}</strong></p>
<p class="sec">2. СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЕТОВ</p>
<p>2.1. Общая стоимость работ, определяется на основании Приложения №1 и №2 к настоящему договору, НДС не облагается на основании ст. 346.11 НК РФ и указывается в Акте выполненных работ за фактически выполненные Подрядчиком работы.</p>
<p>2.2. Оплата работ осуществляется в следующем порядке: 100 % предоплата безналичным расчетом на счет Подрядчика либо наличными денежными средствами в кассу Подрядчика.</p>
<p>2.3. Обязательство Заказчика по безналичной оплате считается исполненным в момент зачисления денежных средств на счет Подрядчика, указанный в реквизитах.</p>
<p>2.4. В случае некомплектации дополнительной фурнитуры, материалов и т.д., приобретенных Заказчиком самостоятельно, препятствующего окончанию выполнения Работ, все последующие выезды Подрядчика оплачиваются в размере 2000,00 (две тысячи) рублей. Дополнительная оплата осуществляется Заказчиком в срок не позднее 1 (одного) рабочего дня до предполагаемой Сторонами даты выполнения Работ.</p>
<p class="sec">3. СРОКИ И ПОРЯДОК ВЫПОЛНЕНИЯ РАБОТ</p>
<p>3.1. Работы по договору выполняются в течение 30 (тридцати) рабочих дней со дня доставки мебели Заказчику, при условии получения от Заказчика предварительной оплаты по Договору. Подрядчик вправе досрочно выполнить работу без получения предварительного согласия Заказчика.</p>
<p>3.2. Стороны согласовывают дату и время проведения Работ путем обмена сообщениями по электронной почте или обмена сообщениями (WhatsApp, Telegram, СМС и т.д.) на телефонный номер Сторон согласно реквизитам, указанным в Договоре.</p>
<p>3.3. В случае, если в период выполнения Работ будут выявлены недостатки элементов мебели и иных изделий, которые необходимо устранять, то срок выполнения Работ будет увеличен на срок устранения недостатков. Заказчик и Подрядчик пришли к соглашению, что если выявленные недостатки возникли в результате действий Заказчика и/или третьих лиц, не имеющим отношения к {co_name}, то устранение недостатков осуществляется за счет Заказчика.</p>
<p>3.4. В случае, если Заказчик переносит дату выполнения Работ, то срок выполнения Работ Подрядчика будет увеличен.</p>
<p>3.5. В случае проведения Заказчиком и/или иными лицами строительных, ремонтных и иных работ в помещении, в котором необходимо выполнить Работы, а также, в случае не освобождения данного помещения от строительного и иного мусора к согласованной Сторонами дате и времени, срок выполнения Работ будет увеличен, при этом Подрядчик будет считаться приостановившим свои обязательства по срокам выполнения Работ и не будет нести никакой ответственности за нарушение условий о сроке по настоящему Договору.</p>
<p>3.6. Подрядчик выполняет Работы в рабочие дни с 09:00 до 21:00 часов, в выходные дни с 10:00 до 19:00 часов. Перерыв с 13:00 до 15:00 часов.</p>
<p>3.7. Окончанием Работ является день, когда Сторонами подписан Акт выполненных Работ.</p>
<p>3.8. Подрядчик оставляет за собой право не приступать к исполнению обязательств по настоящему Договору или приостановить исполнение обязательств до получения им суммы предварительной оплаты. При этом срок выполнения Работ автоматически продлевается на период просрочки оплаты, допущенной Заказчиком.</p>
<p>3.9. При выполнении Работ в помещении Заказчика Подрядчик имеет право проводить работы с использованием электроинструмента. Проводимые в процессе монтажа работы неизбежно связаны с шумом, образованием пыли и стружки материала, а также выделением характерных запахов.</p>
<p>3.10. В случае невозможности надлежащей установки мебели в связи с неровностью поверхности (пола, стен, потолков) в помещении Заказчика, Заказчик самостоятельно устраняет эти недостатки.</p>
<p>3.11. При установке мебели в неподготовленном помещении допускается отклонение от первоначальной конструкции, то есть появление зазоров и т.д. Подрядчик не несет ответственности за геометрию стен, пола и потолка Заказчика.</p>
<p class="sec">4. ОБЯЗАННОСТИ СТОРОН</p>
<p><strong>4.1. Подрядчик обязан:</strong></p>
<p>4.1.1. Приступить к выполнению Работ в срок установленные в п. 3.1. Договора.</p>
<p>4.1.3. Предупредить Заказчика о возможных неблагоприятных для него последствиях выполнения его указаний о способе исполнения работы, а также иных не зависящих от подрядчика обстоятельств, которые грозят годности или прочности результатов выполняемой работы либо создают невозможность ее завершения в срок.</p>
<p><strong>4.2. Заказчик обязан:</strong></p>
<p>4.2.1. Принять и оплатить выполненную Подрядчиком Работу на условиях и в порядке, предусмотренных Договором.</p>
<p>4.2.2. Присутствовать при выполнении Работ, либо обеспечить присутствие своего представителя, который уполномочен подписывать Акт выполненных работ. Заказчик предварительно уведомляет Подрядчика об уполномоченном представителе.</p>
<p>4.2.3. Обеспечить бесплатную парковку автомобиля Подрядчика на период выполнения Работ по адресу, указанному в п. 1.2. Договора. В случае, если парковка по адресу платная, то парковка оплачивается Заказчиком.</p>
<p>4.2.4. Закончить все ремонтные и иные работы в помещении, в котором будут выполнены Работы, а также освободить данное помещение от строительного и иного мусора.</p>
<p>4.2.5. Подготовить помещение для выполнения работ по монтажу:<br>- убрать по возможности мебель: полы, пол, бытовые приборы;<br>- убрать лишние предметы, на время проведения работ должны быть удалены все посторонние лица, а также домашние животные;<br>- обеспечить в помещении наличии освещения и электросети 220В;<br>- температура воздуха в помещении должна быть в пределах от +10 до +25 град. С;<br>- относительная влажность воздуха в помещении должна быть в пределах от 45 до 75 %;<br>- полы должны быть выполнены с уклоном не более 3 мм уклона на 1 метр пола;<br>- стены, полы и потолки должны быть выровнены и составлять между собой угол 90 градусов;<br>- стены должны быть приспособлены для крепления мебели.</p>
<p>4.2.6. Предоставлять информацию, необходимую для выполнения Работ по Договору.</p>
<p><strong>4.3. Подрядчик вправе:</strong></p>
<p>4.3.1. Требовать своевременного подписания Заказчиком Акта выполненных работ.</p>
<p>4.3.2. Требовать своевременной оплаты Работ в соответствии с пунктом 2.2. Договора.</p>
<p>4.3.3. Запрашивать у Заказчика разъяснения и уточнения относительно порядка выполнения Работ.</p>
<p>4.3.4. Привлекать третьих лиц для исполнения обязательств по настоящему Договору. Подрядчик несет ответственность за действия/бездействие третьих лиц, выполняющих Работу по Договору, как за свои собственные.</p>
<p>4.3.5. Не приступать к начатой Работе приостановить в случаях, когда нарушение Заказчиком своих обязанностей, установленных в п. 4.2 Договора препятствует исполнению Договора Подрядчиком, а также при наличии обстоятельств, очевидно свидетельствующих о том, что исполнение указанных обязанностей не будет произведено в установленный срок.</p>
<p>4.3.6. Самостоятельно определить количество специалистов, необходимых для выполнения работ, график работы, а также сроки в зависимости от сложности Изделия.</p>
<p><strong>4.4. Заказчик вправе:</strong></p>
<p>4.4.1. Требовать от Подрядчика надлежащего исполнения обязательств в соответствии с Договором, а также требовать своевременного устранения выявленных недостатков при приемке Работ.</p>
<p>4.4.2. Отозвать согласие на обработку своих персональных данных Подрядчиком.</p>
<p class="sec">5. ПОРЯДОК СДАЧИ-ПРИЕМКИ РАБОТ</p>
<p>5.1. По завершении выполнения работ Подрядчик предоставляет Заказчику результаты выполненных работ и Акт выполненных работ, который Заказчик утверждает в день выставления Акта либо дает мотивированный отказ от приемки работ в В день окончания выполнения Подрядчиком Сторонами (их уполномоченными представителями) подписывают Акт выполненных работ, в соответствии с которым Подрядчик передает, а Заказчик принимает Работы и их результат – партию собранной и смонтированной мебели. В случае наличия у Заказчика замечаний к качеству Работ, Заказчик обязан подписать указанный Акт с изложением замечаний. После подписания Акта выполненных работ Подрядчик не принимает претензий по указанному Акту с изложением замечаний.</p>
<p>5.2. В день окончания выполнения Подрядчиком Сторонами (их уполномоченными представителями) подписывают Акт выполненных работ.</p>
<p>5.3. При необоснованном отказе одной из сторон от подписания акта в нем делается отметка об этом и акт подписывается другой стороной в одностороннем порядке, работы считаются выполненными на день подписания такого Акта одной из сторон.</p>
<p class="sec">6. ОТВЕТСТВЕННОСТЬ СТОРОН</p>
<p>6.1. Стороны несут ответственность за неисполнение или ненадлежащее исполнение своих обязательств по Договору в соответствии с Законом Российской Федерации от 7 февраля 1992 г. № 2300-1 «О защите прав потребителей» и иными правовыми актами, принятыми в соответствии с ним.</p>
<p>6.2. За нарушение сроков оплаты Подрядчик вправе потребовать с Заказчика уплаты неустойки (пени) за каждый день просрочки в размере 0,1 % от суммы задолженности.</p>
<p>6.3. Подрядчик не несет ответственности за невыполнение обязательств по Договору, если оно вызвано неисполнением соответствующих обязательств Заказчика.</p>
<p class="sec">7. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ</p>
<p>7.1. Все споры или разногласия, возникающие между Сторонами по настоящему договору или в связи с ним, разрешаются путем переговоров между ними.</p>
<p>7.2. До рассмотрения спора на рассмотрение суда Стороны предусматривают обязательный претензионный порядок урегулирования разногласий. Направление претензии возможно по электронной почте, путем обмена сообщениями (WhatsApp, Telegram, СМС и т.д.) на телефонный номер Сторон согласно реквизитам, указанным в Договоре. Срок ответа на претензию Сторон устанавливается в 10 (Десять) рабочих дней с момента получения претензии заинтересованной Стороны.</p>
<p>7.3. В случае невозможности разрешения споров путем переговоров споры или разногласия, возникающие между Сторонами, решаются в судебном порядке по месту нахождения Подрядчика.</p>
<p class="sec">8. СРОК ДЕЙСТВИЯ ДОГОВОРА</p>
<p>8.1. Договор вступает в силу с момента его подписания обеими Сторонами и действует до полного исполнения Сторонами принятых на себя обязательств.</p>
<p>8.2. Изменения и дополнения к Договору принимаются по обоюдному соглашению Сторон, путем подписания Дополнительного соглашения к Договору.</p>
<p>8.3. Договор может быть расторгнут досрочно по письменному соглашению Сторон, в одностороннем порядке в случаях, предусмотренных действующим законодательством Российской Федерации.</p>
<p>8.4. В случае одностороннего отказа от Договора Заказчик обязуется, оплатить Подрядчику фактически понесенных им расходов, связанных с исполнением обязательств по данному Договору.</p>
<p class="sec">9. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ</p>
<p>9.1. Договор вступает в силу с момента его подписания Сторонами и действует до полного исполнения Сторонами всех принятых на себя обязательств.</p>
<p>9.2. При заключении Договора и подписания других документов в рамках его исполнения, Стороны согласны на факсимильное воспроизведение подписи («факсимиле»), уполномоченных лиц, а также на обмен электронными копиями документов. Стороны признают юридически значимой переписку между сторонами по электронной почте, а также переписку по номеру телефона, указанному в настоящем договоре в любых приложениях для обмена сообщениями (WhatsApp, Telegram, СМС и т.д.). Направленные таким образом документы считаются подписанными простой электронной подписью и признаются Сторонами равнозначными бумажным, подписанным собственноручной подписью Сторон.</p>
<p>9.3. Стороны признают юридическую силу всех документов, требований, уведомлений, претензий, извещений, оформленных должным образом и направленных друг другу в электронном виде по указанным адресам электронной почты и телефонам по исполнению настоящего Договора. Стороны договорились, что сообщение при такой форме коммуникации считается доставленным тогда, когда приложение, через которое производится коммуникация подтверждает факт направления сообщения через свой внутренний интерфейс. При этом и Заказчик, и Подрядчик обязуются регулярно просматривать указанные источники коммуникации. Стороны договорились, что сообщение считается доставленным и в тех случаях, если оно поступило лицу, которому оно направлено, но по обстоятельствам, зависимым от него, но не было ему вручено или адресат не ознакомился с ним.</p>
<p>9.4. Заказчик дает свое согласие на обработку своих персональных данных, а именно: на действия, совершаемые с использованием средств автоматизации или без использования таких средств, включая сбор, запись, систематизацию, накопление, хранение, уточнение (обновление, изменение), извлечение, использование, передачу (распространение, предоставление, доступ), обезличивание, блокирование, удаление, уничтожение его персональных данных, {co_name} (ОГРН: {co_ogrn}, ИНН: {co_inn}), расположенному по адресу: {co_address}.<br>Цель обработки персональных данных: исполнение настоящего договора. Заказчик дает свое согласие на использование следующих персональных данных: фамилия, имя, отчество, паспортные данные, адрес места жительства, фотографии изделий на объекте Заказчика; номер телефона, адрес электронной почты. Согласие предоставляется на срок действия настоящего договора, а после прекращения договора - в течение 12 месяцев с даты подписания Сторонами акта сдачи-приемки выполненных Работ.<br>Настоящее условие договора может быть изменено Заказчиком - субъектом персональных данных, в любой момент в одностороннем порядке путем отзыва согласия на обработку персональных данных. Отзыв согласия на обработку персональных данных осуществляется посредством составления письменного документа, который может быть направлен в адрес Подрядчика почтовым отправлением с уведомлением о вручении, либо вручен лично под расписку представителю Подрядчика.</p>
<div style="margin:20px 0;display:flex;gap:40px"><div style="flex:1;border-top:1px solid #000;padding-top:4px;text-align:center">(подпись)</div><div style="flex:1"></div><div style="flex:1;border-top:1px solid #000;padding-top:4px;text-align:center">(расшифровка подписи от руки)</div></div>
<p>9.5. К настоящему Договору прилагаются и являются неотъемлемой частью следующие приложения:<br>1. Калькуляция на выполнение работ по сборке мебели.<br>2. Прайс на дополнительные работы.<br>3. Образец Акта выполненных работ.</p>
<p class="sec">10. РЕКВИЗИТЫ СТОРОН</p>
<table><tr><th style="width:50%">Подрядчик:</th><th style="width:50%">Заказчик:</th></tr>
<tr><td>{co_requisites}{stamp_sig_html}</td>
<td>Паспорт. Серия, номер: {_passport_str(c)}<br>Кем выдан:<br><br>Адрес прописки:<br><br>Телефон: {c.get("phone") or "___________"}<br><br>Предпочитаемый канал обмена сообщениями:</td></tr>
<tr><td style="padding-top:20px">Менеджер: <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td></td></tr>
<tr><td style="height:60px"></td><td></td></tr>
<tr><td style="height:60px"></td><td>{fname}</td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'rules':
        style = _doc_style('Правила эксплуатации', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Правила эксплуатации мебели</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 3 к договору бытового подряда на изготовление мебели № {contract_num} от {contract_date_full}</p>
<h1>«Правила эксплуатации корпусной мебели»</h1>

<p class="sec">1. ОБЩИЕ РЕКОМЕНДАЦИИ</p>
<p>1.1. Срок службы Мебели и сохранение его потребительских свойств напрямую зависят от соблюдения Заказчиком правил, изложенных в настоящем приложении.</p>
<p>1.2. Климатические условия и воздействия окружающей среды (свет, влажность, температура) напрямую влияют на состояние мебели.</p>
<p>1.3. Подрядчик рекомендует соблюдать оптимальные климатические условия в помещении, где установлена мебель: температура воздуха от +18°C до +25°C, относительная влажность воздуха 45% - 70%.</p>
<p>1.4. Следует оберегать мебель от длительного воздействия прямых солнечных лучей, источников тепла (батареи, обогреватели, духовые шкафы, плиты на расстоянии менее 0,5 м), а также от резких перепадов температуры и влажности.</p>
<p>1.5. Запрещается воздействие на Мебель агрессивных жидкостей (кислот, щелочей, растворителей), абразивных чистящих средств и материалов, способных повредить покрытие.</p>
<p>1.6. Мебель предназначена для эксплуатации в жилых или общественных помещениях в соответствии с ее функциональным назначением. Подрядчик не несет ответственности за повреждения, вызванные несоблюдением рекомендуемых условий.</p>
<p>1.7. Подрядчик гарантирует соответствие Мебели обязательным требованиям нормативных документов, действующих на территории РФ, в том числе:<br>
* ГОСТ 19917-2014 «Мебель для сидения и лежания. Общие технические условия»;<br>
* ГОСТ 32289-2013 «Плиты древесно-стружечные, облицованные пленками на основе термореактивных полимеров. Технические условия»;<br>
* ГОСТ 16371-2014 «Мебель. Общие технические условия»;<br>
* ТР ТС 025/2012 «О безопасности мебельной продукции».</p>

<p class="sec">2. УСЛОВИЯ ХРАНЕНИЯ</p>
<p>2.1. Хранение Мебели до ее установки должно осуществляться в сухих, проветриваемых, отапливаемых помещениях, защищенных от атмосферных осадков, прямых солнечных лучей и источников тепла. Запрещается хранить изделия в помещениях с повышенной влажностью (санузлы, бани, подвалы без отопления).</p>
<p>2.2. Рекомендуемый диапазон температуры воздуха от +10°C до +25°C. Запрещена эксплуатация при температурах ниже -20°C и выше +40°C, а также резкие перепады температур.</p>
<p>2.3. Рекомендуемая относительная влажность воздуха 45% - 70%. Длительное воздействие повышенной влажности (свыше 80%) или чрезмерной сухости (ниже 40%) недопустимо и приводит к деформации элементов (разбуханию, рассыханию, расслоению).</p>
<p>2.4. Мебель должна храниться в заводской упаковке на ровной поверхности. Запрещается хранить изделия в вертикальном положении, прислонив к стене.</p>
<p>2.5. Перед установкой Мебели, доставленной или хранившейся при отрицательных температурах, необходимо выдержать ее в оригинальной упаковке в условиях помещения не менее 72 часов (3 суток) для адаптации к комнатной температуре и влажности.</p>

<p class="sec">3. ПРАВИЛА ЭКСПЛУАТАЦИИ И УХОДА ЗА МЕБЕЛЬЮ</p>
<p><strong>3.1. Общие правила:</strong></p>
<p>3.1.1. Не превышайте максимально допустимые нагрузки на элементы Мебели: полки - до 15 кг, выдвижные ящики из ЛДСП - до 5 кг, ящики на системе «метабокс» - до 18 кг, ящики на системе «тандембокс» - до 35 кг.</p>
<p>3.1.2. Общая нагрузка на подвесной шкаф не должна превышать 70 кг. Высокие конструкции (пеналы, стеллажи) должны быть больше нагружены в нижней части.</p>
<p>3.1.3. Равномерно распределяйте нагрузку внутри шкафов и ящиков: тяжелые предметы размещайте ближе к краям и опорам, легкие - в центре.</p>
<p>3.1.4. Исключите попадание воды на незащищенные торцы деталей, места стыков и торцов изделия.</p>
<p>3.1.5. Запрещается воздействовать на поверхности абразивными, кислотными, щелочными средствами, ацетоном, растворителями.</p>
<p>3.1.6. Не допускается заслонять вентиляционные решетки и воздухозаборные отверстия бытовых приборов, встроенных в мебель.</p>
<p>3.1.7. При установке Мебели в бревенчатых домах, вследствие усадки таких домов, возможна деформация: перекашивание фасадов, опускание как верхних, так и нижних модулей, это не является признаком некачественной работы Подрядчика.</p>
<p><strong>3.2. Корпус ЛДСП:</strong></p>
<p>3.2.1. Главное правило! ЛДСП боится длительного контакта с водой. Попадание жидкости на кромки и тем более в места стыков недопустимо – плита разбухнет и деформируется. Все пролитые жидкости следует немедленно вытирать насухо.</p>
<p>3.2.2. Запрещается мыть ЛДСП большим количеством воды, использовать пароочистители. Не оставляйте не отжатые тряпки, мокрые полотенца на торцах деталей.</p>
<p>3.2.3. Запрещается применять абразивные и едкие химические средства.</p>
<p>3.2.4. Протирайте сухой или слегка влажной хорошо отжатой тканью. Для удаления загрязнений использовать мягкие мыльные растворы.</p>
<p>3.2.5. Берегите кромки и торцы от сколов и ударов. Именно через поврежденные кромки влага легче всего проникает внутрь плиты.</p>
<p>3.2.6. Не превышайте максимально допустимые нагрузки на полки и перегородки. Равномерно распределяйте вес.</p>
<p>3.2.7. Избегайте расположения Мебели вплотную к отопительным приборам (батареям, обогревателям), это может привести к расслоению ламинированного слоя и деформации плиты.</p>
<p><strong>3.3. Фасады МДФ:</strong></p>
<p>3.3.1. Избегайте резких перепадов температуры и попадания прямых солнечных лучей.</p>
<p>3.3.2. Не допускайте длительного контакта поверхности с сильно нагретыми предметами (сковородки, кастрюли, утюги).</p>
<p>3.3.3. Избегайте ударов, царапин острыми предметами, давления на выступающие элементы.</p>
<p>3.3.4. Исключите контакт с агрессивными химическими веществами: растворителями, ацетоном, сильными чистящими порошками, средствами для мытья стекол. Регулярно удаляйте пыль мягкой тканью. Используйте губки с мягкой стороной (без абразивного слоя) для удаления более сложных загрязнений.</p>
<p>3.3.5. Разрешается использовать нейтральные (pH-нейтральные) средства, а также мягкие мыльные растворы.</p>
<p>3.3.6. Запрещается мыть МДФ большим количеством воды, использовать пароочистители. Не оставляйте не отжатые тряпки, мокрые полотенца на торцах деталей.</p>
<p>3.3.7. Не рекомендуется снимать защитную пленку с фасадов до окончания процесса установки мебели.</p>
<p>3.3.8. При изготовлении дверных полотен высотой более 1600 мм, рекомендовано использовать систему выпрямления дверей, чтобы избежать возможной деформации изделия.</p>
<p><strong>3.4. Столешницы и стеновые панели из пластика и искусственного (акрилового) камня:</strong></p>
<p>3.4.1. Для очистки используйте мягкую ткань, мягкую губку, салфетки из микрофибры и средства для ухода за глянцевыми поверхностями. Для удаления сложных загрязнений со столешниц из искусственного (акрилового) камня используются специальные средства.</p>
<p>3.4.2. Запрещается использовать жесткие и металлические губки, щетки и абразивные чистящие средства. Использовать при очистке кислоты, щелочи, соли, растворители.</p>
<p>3.4.5. Запрещается использовать поверхность в качестве разделочной доски. Воздействовать на поверхность столешницы острыми предметами. Передвигать по поверхности посуду с металлическим дном.</p>
<p>3.4.6. На пластиковой столешнице не оставляйте лужи воды, не отжатые тряпки, мокрые полотенца, особенно на стыках и возле мойки. Сразу удаляйте воду с поверхностей столешницы и со стыковочных швов. Еврозапил (зона стыка без соединительной планки) особенно уязвим для влаги и высоких температур. Он несет декоративную функцию.</p>
<p>3.4.7. Не допускается ставить на поверхность столешницы горячие (&gt;60°С) предметы (чайники, кастрюли, сковороды, утюги и т.п.). Используйте специальные термоизоляционные подставки под горячее, разделочные доски, салфетки, коврики и т.п. Не допускается резкий перепад температур.</p>
<p>3.4.8. Не допускается размораживать продукты, или оставлять на длительное время на поверхности столешницы сильно охлаждённые (&lt;0°C) предметы.</p>
<p><strong>3.6. Стеклянные и зеркальные поверхности:</strong></p>
<p>3.6.1. Очищайте специальными средствами для стекол и зеркал, нанося состав на мягкую ткань, а не прямо на поверхность.</p>
<p>3.6.2. Избегайте абразивных средств и жестких губок. Не допускайте ударных и чрезмерных нагрузок на стеклянные полки.</p>
<p><strong>3.7. Фурнитура и механизмы:</strong></p>
<p>3.7.1. Протирание сухой мягкой тканью. Для удаления загрязнений допускается использование слабого мыльного раствора.</p>
<p>3.7.2. Запрещается прилагать чрезмерные усилия для открывания/закрывания дверей и ящиков. Угол открывания распашных дверей, как правило, составляет 90–110°. Превышение угла открывания может привести к поломке петли.</p>
<p>3.7.3. Запрещается открывать фасады с системой «Push to Open» любым способом, кроме нажатия на фасад.</p>
<p>3.7.4. Выдвигайте ящики полностью только при необходимости, держась за ручки или фасад.</p>
<p>3.7.5. Регулярно удаляйте пыль, крошки и грязь с направляющих и других движущихся частей. Следите, чтобы в зону работы механизма (направляющие, петли) не попадали посторонние предметы, которые могут помешать движению.</p>
<p>3.7.6. Регулярная регулировка и смазка механизмов (парафином или специальными средствами) являются обязанностью Заказчика и не покрываются гарантией.</p>
<p>3.7.7. Лицевую фурнитуру следует чистить мягкими тканями с применением хозяйственного мыла, после чего вытирать насухо. Не использовать средства, содержащие абразивные материалы (наждачную бумагу, соду и др.).</p>

<p style="margin-top:16px;font-style:italic">Соблюдая эти несложные правила, вы сохраните безупречный вид и функциональность вашей мебели на долгие годы. Помните: что несоблюдение правил эксплуатации может привести к сокращению сроков службы и преждевременному выходу из строя элементов кухонной мебели.</p>
<p style="margin-top:12px;font-weight:bold">ВНИМАНИЕ! Подрядчик не несет ответственность за последствия от несоблюдения установленных норм и правил по уходу и эксплуатации корпусной мебели.</p>

<table style="margin-top:30px"><tr><th style="width:50%">Подрядчик: ООО «ИНТЕРЬЕРНЫЕ РЕШЕНИЯ»</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="height:50px">&nbsp;</td><td>&nbsp;</td></tr>
<tr><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
<tr><td style="height:50px">&nbsp;</td><td>&nbsp;</td></tr>
</table>
<p class="center" style="margin-top:30px">М.П.</p>
</div></body></html>'''

    elif doc_type == 'act_delivery':
        daddr = _delivery_address(c)
        floor_ = c.get('delivery_floor') or '___'
        elevator = c.get('delivery_elevator') or 'нет'
        delivery_date_str = _fmt_date(c.get('delivery_date') or '')
        dcost = float(c.get('delivery_cost') or 0)
        dcost_words = _num_to_words(dcost) if dcost else '___________'
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Акт об оказании услуг — {contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 3 к договору на оказание услуг по доставке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Акт об оказании услуг»</h1>
<div class="city-date"><span></span><span>от «____» ______________ 20____ г.</span></div>
<p style="text-align:center">(ФОРМА)</p>
<p class="no-indent">ООО «Интерьерные решения», в лице менеджера <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, действующего на основании доверенности № <span class="ul" style="min-width:40px">&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», подписали настоящий Акт об оказании услуг о нижеследующем:</p>
<p>1. В соответствии с договором на оказание услуг по доставке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> года, Исполнитель оказал, а Заказчик принял следующие услуги:<br>Адрес доставки: <strong>{daddr}</strong><br>Этаж: {floor_}&nbsp;&nbsp;&nbsp;&nbsp;Лифт: {elevator}</p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Доставка мебели в территориальных границах г. Саратова и г. Энгельса</td><td style="text-align:center">1 услуга</td><td style="text-align:center">1</td><td style="text-align:right">0</td><td style="text-align:right">0</td></tr>
<tr><td>Доставка мебели за пределы территориальных границ г. Саратова и г. Энгельса</td><td style="text-align:center">1 км</td><td style="text-align:center">70</td><td style="text-align:right"></td><td style="text-align:right">0</td></tr></table>
<p><em>Дополнительные услуги, по факту оказания</em></p>
<p><strong>Подъем мебели (полный комплект), при отсутствии лифта и занос мебели, при невозможности парковки автомашины вплотную к подъезду</strong></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Подъем. Квадратура корпуса до 20 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">550</td><td></td></tr>
<tr><td>Подъем. Квадратура корпуса от 20 до 25 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">650</td><td></td></tr>
<tr><td>Подъем. Квадратура корпуса более 25 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">750</td><td></td></tr>
<tr><td>Перемещение мебели вручную в случае невозможности подъезда автомашины к месту разгрузки за каждые</td><td style="text-align:center">1 м</td><td></td><td style="text-align:right">30</td><td></td></tr></table>
<p><em>Подъем комплектующих отдельно</em></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:12%">Ед. измерения</th><th style="width:12%">Количество упаковок</th><th style="width:12%">Количество этажей</th><th style="width:16%">Цена за ед. измерения в руб.</th><th style="width:16%">Стоимость в руб.</th></tr>
<tr><td>Подъем столешницы</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">350</td><td></td></tr>
<tr><td>Подъем стеновой панели</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">250</td><td></td></tr>
<tr><td>Подъем крупных частей корпуса. (боковины кухонного пенала, боковины шкафов-купе и т.д.). Упаковка по 2 шт.</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">250</td><td></td></tr>
<tr><td>Подъем длинных комплектующих (цоколь, профильная ручка, труба – целиком)</td><td style="text-align:center">1 комплект / 1 этаж</td><td></td><td></td><td style="text-align:right">60</td><td></td></tr>
<tr><td>Подъем дверей-купе</td><td style="text-align:center">1 дверь / 1 этаж</td><td></td><td></td><td style="text-align:right">150</td><td></td></tr>
<tr><td colspan="5" style="text-align:right;font-weight:bold">Цена услуг составляет:</td><td></td></tr></table>
<p>2. Дополнительные услуги, оказанные исполнителем: <strong>оплачены / не оплачены</strong> (ненужное зачеркнуть).</p>
<p>3. Услуги по договору оказаны Исполнителем полностью, своевременно и надлежащим образом. Заказчик претензий к Исполнителю: <strong>не имеет / имеет</strong> (ненужное зачеркнуть).</p>
<div style="height:20px"></div>
<p>4. В случае наличия замечаний Заказчик, после подписания акта, в праве требовать устранения замечаний отраженных в данном акте.</p>
<p>5. Настоящий акт подписан в 2 (двух) экземплярах по одному для каждой из Сторон.</p>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: ООО «ИНТЕРЬЕРНЫЕ РЕШЕНИЯ»</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="height:60px"></td><td>{fname}</td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'act_assembly':
        daddr = _delivery_address(c)
        acost = float(c.get('assembly_cost') or 0)
        acost_words = _num_to_words(acost) if acost else '___________'
        style = _doc_style('Акт выполненных работ (монтаж)', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Акт выполненных работ — {contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 3 к договору на выполнение работ по монтажу и сборке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Акт выполненных работ»</h1>
<div class="city-date"><span></span><span>от «____» ______________ 20____ г.</span></div>
<p style="text-align:center">(ФОРМА)</p>
<p class="no-indent">{co_name}, в лице {co_dir_pos}а <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, действующего на основании доверенности № <span class="ul" style="min-width:40px">&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, именуемый в дальнейшем «Подрядчик» и гр. <strong>{fname}</strong>, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», подписали настоящий Акт об оказании услуг о нижеследующем:</p>
<p>1. Подрядчик в соответствии с договором на выполнение работ по монтажу и сборке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> года выполнил для Заказчика следующие работы:<br>Адрес: <strong>{daddr}</strong></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Сборка и монтаж мебели согласно договору бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td style="text-align:center">1 услуга</td><td style="text-align:center">1</td><td style="text-align:right">0</td><td style="text-align:right">0</td></tr>
<tr><td>Выезд сборщика за пределы территориальных границ г. Саратова и г. Энгельса</td><td style="text-align:center">1 км</td><td style="text-align:center">40</td><td style="text-align:right"></td><td style="text-align:right">0</td></tr></table>
<p><em>Дополнительные работы, по факту выполнения</em></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Врезка мойки с герметизацией. Без подвода и пуска воды</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1250</td><td></td></tr>
<tr><td>Врезка варочной поверхности с герметизацией. Без пуска газа</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1250</td><td></td></tr>
<tr><td>Установка вытяжки без подключения принудительного рукава воздуховода.</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1500</td><td></td></tr>
<tr><td>Установка духового шкафа</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">750</td><td></td></tr>
<tr><td>Установка встраиваемой СВЧ-печи</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">650</td><td></td></tr>
<tr><td>Установка встраиваемого холодильника</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">2400</td><td></td></tr>
<tr><td>Установка посудомоечной машины. Без подвода и пуска воды</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1500</td><td></td></tr>
<tr><td>Установка светодиодной ленты (1 погонный метр)</td><td style="text-align:center">п.м.</td><td></td><td style="text-align:right">1250</td><td></td></tr>
<tr><td>Установка столешницы в подоконник</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">3000</td><td></td></tr>
<tr><td>Установка ручек (если ручки были куплены клиентом самостоятельно)</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">70</td><td></td></tr>
<tr><td colspan="4" style="text-align:right;font-weight:bold">Цена работ составляет:</td><td></td></tr></table>
<p>2. Дополнительные услуги, оказанные подрядчиком: <strong>оплачены / не оплачены</strong> (ненужное зачеркнуть).</p>
<p>3. Комплектность, количество, вид, характеристики мебели соответствуют условиям договора. Визуальный осмотр мебели на предмет повреждений, царапин, сколов, трещин и других недостатков произведен Заказчиком. Фурнитура (петли, выдвижные механизмы, подъемные механизмы и т.д.) работает исправно. Заказчик претензий по объему, качеству, результату и срокам выполнения работ: <strong>не имеет / имеет</strong> (ненужное зачеркнуть).</p>
<div style="height:20px"></div>
<p>4. В случае наличия замечаний Заказчик, после подписания акта, в праве требовать устранения замечаний отраженных в данном акте.</p>
<p>5. Настоящий акт подписан в 2 (двух) экземплярах по одному для каждой из Сторон.</p>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="height:60px"></td><td>{fname}</td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'measure':
        measurer = c.get('measurer') or ''
        measurer_line = measurer if measurer else ''
        style = _doc_style('Бланк замера', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Бланк замера к договору №{contract_num}</title>{style}<style>
.measure-area{{border:1px solid #000;min-height:180mm;margin:10px 0;position:relative;background:#fff}}
</style></head><body><div class="page">
<p class="no-indent" style="text-align:right">Бланк замера помещения к договору бытового подряда на изготовление мебели № {contract_num} от {contract_date_full}</p>
<div class="measure-area"></div>
<table style="margin-top:16px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding:6px">Замерщик</td><td>&nbsp;</td></tr>
<tr><td style="height:50px"><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
</table>
</div></body></html>'''

    elif doc_type == 'delivery_calc':
        dcost = float(c.get('delivery_cost') or 0)
        style = _doc_style('Калькуляция доставки', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Калькуляция доставки — {contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 1 к договору на оказание услуг по доставке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Калькуляция на выполнение услуг по доставке мебели»</h1>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Доставка мебели в территориальных границах г. Саратова и г. Энгельса *</td><td style="text-align:center">1 услуга</td><td style="text-align:center">1</td><td style="text-align:right">8000</td><td style="text-align:right">8000</td></tr>
<tr><td>Доставка мебели за пределы территориальных границ г. Саратова и г. Энгельса **</td><td style="text-align:center">1 км</td><td style="text-align:center"></td><td style="text-align:right">70</td><td style="text-align:right">0</td></tr>
<tr><td colspan="3" style="text-align:right">Сумма:</td><td colspan="2" style="text-align:right">8000</td></tr>
<tr><td colspan="3" style="text-align:right">Скидка ***:</td><td colspan="2" style="text-align:right">8000</td></tr>
<tr><td colspan="3" style="text-align:right;font-weight:bold">Сумма со скидкой:</td><td colspan="2" style="text-align:right;font-weight:bold">0</td></tr></table>
<p style="margin-top:16px">*Исполнитель вправе однократно предоставить Заказчику скидку в размере стоимости доставки, равной 8000 руб., которая осуществляется в пределах территориальных границ г. Саратова и г. Энгельса. В случае повторной доставки, по причинам не зависящим от Исполнителя, но зависящим от Заказчика, ее стоимость будет определена согласно приложению №1.</p>
<p>**Расчет стоимости доставки учитывает расстояние по километражу от склада Исполнителя по адресу: г. Саратов, ул. Усть-Курдюмская д. 3, до дома (подъезда) заказчика согласно п. 1.3 договора. Километраж определяется на основании сервиса Яндекс.Карты https://yandex.ru/maps/</p>
<p>*** Размер скидки определяется Исполнителем индивидуально.</p>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: ООО «ИНТЕРЬЕРНЫЕ РЕШЕНИЯ»</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding:6px">Менеджер</td><td>&nbsp;</td></tr>
<tr><td style="height:50px"><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'delivery_lift':
        style = _doc_style('Прайс — подъём мебели', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Прайс подъём мебели — {contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 2 к договору на оказание услуг по доставке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Прайс на выполнение услуг по подъему и заносу мебели»*</h1>
<p><strong>Подъем мебели (полный комплект), при отсутствии лифта и занос мебели, при невозможности парковки автомашины вплотную к подъезду</strong></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:15%">Количество (Этаж/метр)</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:15%">Стоимость в руб.</th></tr>
<tr><td>Подъем. Квадратура корпуса до 20 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">550</td><td style="text-align:right">0</td></tr>
<tr><td>Подъем. Квадратура корпуса от 20 до 25 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">650</td><td style="text-align:right">0</td></tr>
<tr><td>Подъем. Квадратура корпуса более 25 кв.м</td><td style="text-align:center">руб./этаж</td><td></td><td style="text-align:right">750</td><td style="text-align:right">0</td></tr>
<tr><td>Перемещение мебели вручную в случае невозможности подъезда автомашины к месту разгрузки за каждые</td><td style="text-align:center">1 м</td><td></td><td style="text-align:right">30</td><td style="text-align:right">0</td></tr></table>
<p style="margin-top:12px"><em>Подъем комплектующих отдельно</em></p>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:12%">Ед. измерения</th><th style="width:12%">Количество упаковок</th><th style="width:12%">Количество этажей</th><th style="width:16%">Цена за ед. измерения в руб.</th><th style="width:14%">Стоимость в руб.</th></tr>
<tr><td>Подъем столешницы</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">350</td><td style="text-align:right">0</td></tr>
<tr><td>Подъем стеновой панели</td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">250</td><td style="text-align:right">0</td></tr>
<tr><td>Подъем крупных частей корпуса. (боковины кухонного пенала, боковины шкафов-купе и т.д.). <strong>Упаковка по 2 шт.</strong></td><td style="text-align:center">1 упаковка / 1 этаж</td><td></td><td></td><td style="text-align:right">250</td><td style="text-align:right">0</td></tr>
<tr><td>Подъем длинных комплектующих (цоколь, профильная ручка, труба – целиком)</td><td style="text-align:center">1 комплект / 1 этаж</td><td></td><td></td><td style="text-align:right">60</td><td style="text-align:right">0</td></tr>
<tr><td>Подъем дверей-купе (<strong>Внимание! Обсудить размер лестничного пролета, относительно размеров двери</strong>)</td><td style="text-align:center">1 дверь / 1 этаж</td><td></td><td></td><td style="text-align:right">150</td><td style="text-align:right">0</td></tr>
<tr><td colspan="5" style="text-align:right;font-weight:bold">Сумма:</td><td style="text-align:right;font-weight:bold">0</td></tr></table>
<p style="margin-top:16px">Особые условия:</p>
<p>1. Если подъем мебели можно осуществить полностью на лифте - подъем бесплатный. Расчет подъема изделия или элементов начинается с первого этажа.</p>
<p>2. Занос мебели на первый этаж – бесплатный, при условии возможности парковки автомашины вплотную к подъезду.</p>
<p>3. При невозможности парковки автомашины в плотную (не далее чем на 10 метров) к подъезду (вход в частный дом и т.д.), сумма заноса рассчитывается по данному прайсу, в зависимости от расстояния на ситуации, если из-за погодных условий (снег, гололед, размытая грунтовая дорога и т.д.) или из-за заставленного автомобилями двора (включая аналогичные ситуации) нет возможности подъехать ко входу.</p>
<p>* Услуги могут быть рассчитаны по факту оказания, путем отражения в акте об оказании услуг.</p>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding:6px">Менеджер</td><td>&nbsp;</td></tr>
<tr><td style="height:50px"><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'assembly_calc':
        style = _doc_style('Калькуляция монтажа', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Калькуляция сборки — {contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 1 к договору на выполнение работ по монтажу и сборке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Калькуляция на выполнение работ по сборке мебели»</h1>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Сборка и монтаж мебели согласно договору бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> *</td><td style="text-align:center">работа</td><td style="text-align:center">1</td><td style="text-align:right">17500</td><td style="text-align:right">17500</td></tr>
<tr><td>Выезд сборщика за пределы территориальных границ г. Саратова и г. Энгельса **</td><td style="text-align:center">1 км</td><td style="text-align:center">40</td><td style="text-align:right"></td><td style="text-align:right">0</td></tr>
<tr><td colspan="3" style="text-align:right">Сумма:</td><td colspan="2" style="text-align:right">17500</td></tr>
<tr><td colspan="3" style="text-align:right">Скидка ***:</td><td colspan="2" style="text-align:right">17500</td></tr>
<tr><td colspan="3" style="text-align:right;font-weight:bold">Сумма со скидкой:</td><td colspan="2" style="text-align:right;font-weight:bold">0</td></tr></table>
<p style="margin-top:16px">*Подрядчик вправе однократно предоставить Заказчику скидку в размере стоимости сборки, равной 17500 руб., которая осуществляется в пределах территориальных границ г. Саратова и г. Энгельса. В случае повторного выезда к месту сборки, по причинам не зависящим от Подрядчика, но зависящим от Заказчика, ее стоимость будет определена согласно приложению №1.</p>
<p>**Расчет стоимости выезда сборщика учитывает расстояние по километражу от склада Подрядчика по адресу: г. Саратов, ул. Усть-Курдюмская д. 3, до дома (подъезда) заказчика согласно п. 1.2 договора. Километраж определяется на основании сервиса Яндекс.Карты https://yandex.ru/maps/</p>
<p>*** Размер скидки определяется Подрядчиком индивидуально.</p>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding:6px">Менеджер</td><td>&nbsp;</td></tr>
<tr><td style="height:50px"><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'assembly_extra':
        style = _doc_style('Прайс — дополнительные работы', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Прайс доп. работы сборка — {contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение № 2 к договору на выполнение работ по монтажу и сборке мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Прайс на дополнительные работы»*</h1>
<table><tr><th>Наименование (перечень) работ и услуг</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td>Врезка мойки с герметизацией. Без подвода и пуска воды</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1250</td><td style="text-align:right">0</td></tr>
<tr><td>Врезка варочной поверхности с герметизацией. Без пуска газа</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1250</td><td style="text-align:right">0</td></tr>
<tr><td>Установка вытяжки без подключения принудительного рукава воздуховода.</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1500</td><td style="text-align:right">0</td></tr>
<tr><td>Установка духового шкафа</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">750</td><td style="text-align:right">0</td></tr>
<tr><td>Установка встраиваемой СВЧ-печи</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">650</td><td style="text-align:right">0</td></tr>
<tr><td>Установка встраиваемого холодильника</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">2400</td><td style="text-align:right">0</td></tr>
<tr><td>Установка посудомоечной машины. Без подвода и пуска воды</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">1500</td><td style="text-align:right">0</td></tr>
<tr><td>Установка светодиодной ленты (1 погонный метр)</td><td style="text-align:center">п.м.</td><td></td><td style="text-align:right">1250</td><td style="text-align:right">0</td></tr>
<tr><td>Установка столешницы в подоконник</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">3000</td><td style="text-align:right">0</td></tr>
<tr><td>Установка ручек (если ручки были куплены клиентом самостоятельно)</td><td style="text-align:center">шт.</td><td></td><td style="text-align:right">70</td><td style="text-align:right">0</td></tr>
<tr><td colspan="4" style="text-align:right;font-weight:bold">Сумма:</td><td style="text-align:right;font-weight:bold">0</td></tr></table>
<p style="margin-top:16px">Стоимость и порядок выполнения дополнительных работ не включенных в данный прайс - обсуждаются Заказчиком и Подрядчиком индивидуально.</p>
<p>* Работы могут быть рассчитаны по факту выполнения, путем отражения в акте выполненных работ.</p>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding:6px">Менеджер</td><td>&nbsp;</td></tr>
<tr><td style="height:50px"><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'tech_spec':
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Спецификация на технику — {contract_num}</title>{style}</head><body><div class="page">
<p class="no-indent" style="text-align:right">Приложение к договору бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<h1>«Спецификация на поставку техники»</h1>
<table><tr><th>Наименование</th><th style="width:15%">Ед. измерения</th><th style="width:12%">Количество</th><th style="width:18%">Цена за ед. измерения в руб.</th><th style="width:18%">Стоимость в руб.</th></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td></td><td style="text-align:center">шт.</td><td></td><td></td><td style="text-align:right">0</td></tr>
<tr><td colspan="4" style="text-align:right;font-weight:bold">Сумма:</td><td style="text-align:right;font-weight:bold">0</td></tr></table>
<table style="margin-top:20px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding:6px">Менеджер</td><td>&nbsp;</td></tr>
<tr><td style="height:50px"><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td><td><span class="ul" style="min-width:200px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    elif doc_type == 'addendum':
        style = _doc_style('Дополнительное соглашение', contract_num)
        return f'''<!DOCTYPE html><html lang="ru"><head><meta charset="UTF-8"><title>Дополнительное соглашение к договору №{contract_num}</title>{style}</head><body><div class="page">
<h1>ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ №<span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></h1>
<h2>К договору бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></h2>
<div class="city-date"><span>г. {co_city}</span><span>«____» ______________ 20____ г.</span></div>
<p class="no-indent">{co_name}, в лице {co_dir_pos}а <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, действующего на основании доверенности № <span class="ul" style="min-width:40px">&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul" style="min-width:80px">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>, именуемый в дальнейшем «Подрядчик» и гр. <strong>{fname}</strong>, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», заключили настоящее дополнительное соглашение к договору бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<p>Стороны пришли к соглашению внести в Договор следующие изменения:</p>
<p>1. <strong>(ВПИСАТЬ ИЗМЕНЕНИЯ)</strong></p>
<p>2. Все другие условия договора, прямо не поименованные в настоящем дополнительном соглашении, считать неизменными и обязательными для исполнения Сторонами.</p>
<p>3. Настоящее Дополнительное соглашение составлено в двух экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из Сторон.</p>
<p>4. Настоящее Дополнительное соглашение вступает в силу с даты его подписания, распространяется на отношения, возникшие с даты заключения договора и является неотъемлемой частью Договора бытового подряда на изготовление мебели <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span> от <span class="ul">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></p>
<table style="margin-top:30px"><tr><th style="width:50%">Подрядчик: {co_name.upper()}</th><th style="width:50%">Заказчик:</th></tr>
<tr><td style="padding-top:8px">Менеджер</td><td>{fname}</td></tr>
<tr><td style="height:60px"></td><td></td></tr>
<tr><td colspan="2" style="text-align:center;padding-top:10px">М.П.</td></tr></table>
</div></body></html>'''

    return '<html><body><p>Неизвестный тип документа</p></body></html>'


def _build_docx(c: dict, doc_type: str, company: dict = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if company is None:
        company = {}
    co_name     = _co(company, 'name',              'ООО «Интерьерные решения»')
    co_city     = _co(company, 'city',              'Саратов')
    co_inn      = _co(company, 'inn',               '6450106826')
    co_ogrn     = _co(company, 'ogrn',              '1196451012251')
    co_kpp      = _co(company, 'kpp',               '')
    co_address  = _co(company, 'address',           '410018, Саратовская обл., г. Саратов, ул. Усть-Курдюмская, д. 3, пом. 1')
    co_phone    = _co(company, 'phone',             '')
    co_dir_pos  = _co(company, 'directorPosition',  'менеджера')
    co_bank     = _co(company, 'bank',              '')
    co_bik      = _co(company, 'bik',               '')
    co_rs       = _co(company, 'rs',                '')
    co_ks       = _co(company, 'ks',                '')

    fname = _full_name(c)
    total = float(c.get('total_amount') or 0)
    total_words = _num_to_words(total)
    contract_num = c.get('contract_number') or '___'
    contract_date_full = _fmt_date_full(c.get('contract_date') or '')
    prod_days = int(c.get('production_days') or 45)
    prepaid = float(c.get('prepaid_amount') or 0)
    balance = float(c.get('balance_due') or 0)
    ptype = c.get('payment_type', '100% предоплата')
    custom = c.get('custom_payment_scheme', '') or ''
    products = _get_products(c)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.left_margin = Mm(20); sec.right_margin = Mm(15)
    sec.top_margin = Mm(15); sec.bottom_margin = Mm(15)

    # Устанавливаем вид «одна страница» и масштаб 100% чтобы Word не открывал в режиме двух страниц
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    settings = doc.settings.element
    zoom_el = _OxmlElement('w:zoom')
    zoom_el.set(_qn('w:percent'), '100')
    zoom_el.set(_qn('w:val'), 'bestFit')
    settings.append(zoom_el)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = Pt(13)
    style.paragraph_format.space_after = Pt(0)

    def _set_font(run, size=11, bold=False, name='Times New Roman'):
        run.font.name = name; run.font.size = Pt(size); run.bold = bold

    def h(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text.upper()); _set_font(r, 12, bold=True)
        return p

    def h2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text); _set_font(r, 11)
        return p

    def sec_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text); _set_font(r, 11, bold=True)
        return p

    def para(text, indent=True, bold_parts=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(13)
        if indent: p.paragraph_format.first_line_indent = Mm(10)
        r = p.add_run(text); _set_font(r, 11)
        return p

    def sig_table(left, right):
        t = doc.add_table(rows=2, cols=2)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, txt in enumerate([left[0], right[0]]):
            c2 = t.cell(0, i); c2.text = txt
            run = c2.paragraphs[0].runs[0]
            _set_font(run, 11, bold=True)
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, txt in enumerate([left[1], right[1]]):
            c2 = t.cell(1, i); c2.text = txt
            _set_font(c2.paragraphs[0].runs[0] if c2.paragraphs[0].runs else c2.paragraphs[0].add_run(txt), 11)
        return t

    if doc_type == 'contract':
        manager = c.get('manager_name') or ''
        manager_line = manager if manager else '_' * 30

        h('ДОГОВОР')
        h2('бытового подряда на изготовление мебели')
        # Город и дата
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_city = p.add_run(f'г. {co_city}'); _set_font(run_city, 11)
        p.add_run('\t\t\t\t')
        run_date = p.add_run(f'№ {contract_num} от {contract_date_full}'); _set_font(run_date, 11)

        para(f'{co_name}, в лице {co_dir_pos}а {manager_line}, действующего на основании доверенности № ____ от ____________, именуемый в дальнейшем «Подрядчик», и гр. {fname}, именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», заключили настоящий Договор о нижеследующем:', indent=False)

        sec_title('1. ПРЕДМЕТ ДОГОВОРА')
        para('1.1. Подрядчик обязуется выполнить работу по изготовлению мебели и передать результат работы Заказчику (мебель передается в разобранном виде), а Заказчик обязуется принять и оплатить результат работ.')
        para('1.2. Наименование, качественные характеристики, количество, дизайн мебели указываются в Техническом проекте, который является Приложением № 1 к настоящему Договору.')
        para('1.3. В случае необходимости выполнения Подрядчиком дополнительных работ, влекущих изменение объема, цены работ, включая изменение срока выполнения работ, стороны согласовывают данные изменения путем заключения дополнительного соглашения.')
        para(f'1.4. Срок выполнения работ составляет {prod_days} ({_days_words(prod_days)}) рабочих дней, с момента согласования Технического проекта и получения Подрядчиком предварительной оплаты, в размере, указанном в разделе 3 Договора. Подрядчик вправе досрочно выполнить работу без получения предварительного согласия Заказчика.')
        para('1.4.1. В случае нарушения технологического процесса (поломка, остановка производственных линий, отсутствие энергоснабжения) по вине коммунальных и иных служб, нехватки сырья и (или) рабочей силы, если эти обстоятельства непосредственно повлияли на возможность надлежащего исполнения Подрядчиком своих обязательств по настоящему Договору, срок изготовления мебели переносится соразмерно времени, в течение которого действовали такие обстоятельства.')
        para('1.4.2. Если Заказчик не обеспечил возможность проезда Подрядчика к месту разгрузки мебели, Подрядчик не несет ответственности за сроки доставки и разгрузки мебели и оставляет за собой право выставить счет за компенсацию дополнительных затрат, понесенных в связи с доставкой мебели, который Заказчик обязуется оплатить в течение 3 (трех) рабочих дней.')
        para('1.5. Работы, выполняемые Подрядчиком, предназначены удовлетворять бытовые или другие личные потребности Заказчика.')
        para('1.6. Заказчик проинформирован о том, что допускаются различия оттенка цвета и текстуры покрытий мебели и (или) ее элементов, по сравнению с образцами, в пределах одного цветового тона, что не будет являться нарушением условий договора.')
        para('1.7. Мебель изготавливается согласно индивидуальному заказу Заказчика (Технический проект). Подписанием данного договора Заказчик подтверждает, что он согласовал все характеристики мебели, включая, но не ограничиваясь: количество, размер, форма, габариты, материал, расцветка, комплектация, отделка, фурнитура, крепления и т.д.')

        sec_title('2. РАЗРАБОТКА И СОГЛАСОВАНИЕ ТЕХНИЧЕСКОГО ПРОЕКТА')
        para('2.1. Технический проект разрабатывается Подрядчиком в течение 10 (десяти) рабочих дней с момента получения Подрядчиком предварительной оплаты, в размере, указанном в разделе 3 Договора.')
        para('2.2. Заказчик в течение 5 (пяти) рабочих дней с момента перечисления предварительной оплаты обеспечивает доступ Подрядчика в помещение, в котором будет установлена мебель для проведения замеров помещения.')
        para('2.3. Заказчик не позднее 5 (пяти) рабочих дней со дня получения от Подрядчика Технического проекта согласовывает, путем проставления личной подписи и даты согласования, или направляет мотивированный отказ. Стороны согласовали допустимость не более 2 (двух) кругов правок.')
        para('2.4. Все изменения или дополнения после подписания Сторонами Технического проекта по инициативе Заказчика недопустимы, за исключением изменений, которые Подрядчик признает существенными.')

        sec_title('3. СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЕТОВ')
        para(f'3.1. Общая стоимость работ, подлежащих выполнению по настоящему Договору складывается на основании Калькуляции (Приложение №2 к настоящему договору) и составляет {total:,.0f} ({total_words}) рублей. НДС не облагается на основании ст. 346.11 НК РФ. В стоимость работ включается стоимость материалов Подрядчика, из которых производится работы.')
        para('3.2. Оплата работ осуществляется в следующем порядке:')
        if custom:
            para(custom)
        elif ptype == '100% предоплата':
            para(f'3.2.1. Предварительная оплата производится при заключении Договора в размере {prepaid:,.0f} ({_num_to_words(prepaid)}) рублей.')
            para('3.2.2. Окончательный платёж не предусмотрен. Стоимость работ оплачена полностью при заключении Договора.')
        else:
            para(f'3.2.1. Предварительная оплата производится при заключении Договора в размере {prepaid:,.0f} ({_num_to_words(prepaid)}) рублей.')
            para(f'3.2.2. Окончательный платёж за выполненные по Договору работы в размере {balance:,.0f} ({_num_to_words(balance)}) рублей осуществляется в течение 3 (трёх) дней с момента получения Заказчиком уведомления о готовности мебели, но не позднее дня доставки.')
        para('3.3. Оплата производится безналичным расчетом на счет Подрядчика либо наличными денежными средствами в кассу.')
        para('3.4. Обязательство Заказчика по безналичной оплате считается исполненным в момент зачисления денежных средств на счет Подрядчика, указанный в реквизитах.')

        sec_title('4. ПРАВА И ОБЯЗАННОСТИ СТОРОН')
        para('4.1. Подрядчик обязан:', indent=False); run = doc.paragraphs[-1].runs[0]; run.bold = True
        para('4.1.1. Выполнить работу по Договору согласно Техническому проекту и передать Заказчику мебель в установленный срок.')
        para('4.1.2. Уведомить Заказчика о готовности мебели по электронной почте, путем обмена сообщениями (Telegram, WhatsApp, СМС и т.д.) на телефонный номер Стороны согласно реквизитам, указанным в Договоре.')
        para('4.1.3. Устранить недостатки, выявленные Заказчиком по результатам приемки работ.')
        para('4.1.4. Предоставить Заказчику необходимую и достоверную информацию о предлагаемой работе, ее видах и особенностях, о цене и форме оплаты.')
        para('4.1.5. Предупредить Заказчика о возможных неблагоприятных последствиях выполнения его указаний о способе исполнения работы.')
        para('4.2. Заказчик обязан:', indent=False); run = doc.paragraphs[-1].runs[0]; run.bold = True
        para('4.2.1. Согласовать и подписать Технический проект в срок, установленный Договором.')
        para('4.2.2. Оплатить стоимость работ в соответствии с условиями настоящего Договора.')
        para('4.2.3. Принять результат работ путем подписания Акта выполненных работ.')
        para('4.2.4. Обеспечить сохранность помещения, в котором будет установлена мебель в том виде, в котором помещение было на момент проведения замеров Подрядчиком.')
        para('4.2.5. Проводить ремонтно-отделочные работы в помещении только после получения письменного согласия Подрядчика.')
        para('4.3. Подрядчик вправе:', indent=False); run = doc.paragraphs[-1].runs[0]; run.bold = True
        para('4.3.1. Требовать подписания Заказчиком Акта выполненных работ в течение 5 (пяти) календарных дней с даты передачи мебели.')
        para('4.3.2. Требовать своевременной оплаты работ в соответствии с п. 3.2 настоящего Договора.')
        para('4.3.3. Привлекать третьих лиц для исполнения обязательств по Договору.')
        para('4.3.4. Досрочно выполнить работы и требовать от Заказчика принять результат работ и произвести его оплату.')
        para('4.4. Заказчик вправе:', indent=False); run = doc.paragraphs[-1].runs[0]; run.bold = True
        para('4.4.1. Выбрать модель мебели, цвет, компоновку, дизайн, материалы, фурнитуру из которых будет выполнена работа.')
        para('4.4.2. Отозвать согласие на обработку своих персональных данных Подрядчиком.')

        sec_title('5. ГАРАНТИЯ И КАЧЕСТВО ВЫПОЛНЕННЫХ РАБОТ')
        para('5.1. Гарантийный срок на мебельную фурнитуру составляет 6 месяцев, на мебель (корпус, фасады, столешницы) — 24 месяца с даты подписания Акта выполненных работ при условии надлежащей эксплуатации. Срок службы изделий — 5 лет.')
        para('5.2. Мебель должна быть осмотрена Заказчиком на предмет внешних повреждений непосредственно при получении от Подрядчика.')
        para('5.3. При обнаружении несоответствия мебели по качеству и/или количеству, Заказчик делает отметку в Акте выполненных работ.')
        para('5.4. Гарантия не распространяется на недостатки, возникшие вследствие нарушения Заказчиком правил эксплуатации, ненадлежащего хранения, чрезмерной нагрузки, попадания едких веществ, использования не по назначению, а также естественного износа мебели.')

        sec_title('6. ПОРЯДОК ПРИЕМКИ ВЫПОЛНЕННЫХ РАБОТ')
        para('6.1. По факту выполнения работ Подрядчик представляет Заказчику Акт выполненных работ в двух экземплярах.')
        para('6.2. Передача мебели производится по Акту выполненных работ, который составляется в момент передачи мебели.')
        para('6.3. При уклонении Заказчика от получения мебели, Подрядчик бесплатно хранит мебель в течение 7 (семи) дней. С восьмого дня вправе запросить оплату за хранение в сумме 500 рублей за каждый день.')
        para('6.4. Риск случайной гибели или повреждения мебели, а также право собственности переходит от Подрядчика к Заказчику в момент подписания соответствующих документов о приемке.')

        sec_title('7. ОТВЕТСТВЕННОСТЬ СТОРОН')
        para('7.1. Стороны несут ответственность за неисполнение или ненадлежащее исполнение своих обязательств по Договору в соответствии с Законом РФ от 7 февраля 1992 г. № 2300-1 «О защите прав потребителей» и иными правовыми актами.')
        para('7.2. За нарушение сроков оплаты услуг Подрядчик вправе потребовать с Заказчика уплаты неустойки (пени) за каждый день просрочки в размере 0,1% от суммы задолженности.')
        para('7.3. Подрядчик не несет ответственности за невыполнение обязательств по Договору, если оно вызвано неисполнением соответствующих обязанностей Заказчика.')
        para('7.4. В случае отказа Заказчика от мебели надлежащего качества, изготовленной по индивидуальному заказу, Стороны признают, что размер убытков Подрядчика составляет стоимость соответствующей партии мебели.')

        sec_title('8. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ')
        para('8.1. Все споры или разногласия, возникающие между Сторонами, разрешаются путем переговоров между ними.')
        para('8.2. До передачи спора на рассмотрение суда Стороны предусматривают обязательный претензионный порядок урегулирования разногласий. Срок ответа на претензию — 10 (десять) рабочих дней с момента получения.')
        para('8.3. В случае невозможности разрешения разногласий путем переговоров, споры решаются в судебном порядке по месту нахождения Подрядчика.')

        sec_title('9. СРОК ДЕЙСТВИЯ ДОГОВОРА')
        para('9.1. Договор вступает в силу с момента его подписания обеими Сторонами и действует до полного исполнения Сторонами принятых на себя обязательств.')
        para('9.2. Изменения и дополнения к Договору принимаются по обоюдному соглашению Сторон, путем подписания Дополнительного соглашения к Договору.')
        para('9.3. Договор может быть расторгнут досрочно по письменному соглашению Сторон или в одностороннем порядке в случаях, предусмотренных действующим законодательством РФ.')
        para('9.4. В случае одностороннего отказа Заказчика от Договора, Заказчик обязуется оплатить Подрядчику фактически понесенные им расходы, связанные с исполнением обязательств по данному Договору.')

        sec_title('10. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ')
        para('10.1. Стороны согласны на факсимильное воспроизведение подписи уполномоченных ими лиц, а также на обмен копиями документов. Стороны признают юридически значимой переписку по электронной почте и мессенджерам (Telegram, WhatsApp, СМС и т.д.).')
        para('10.2. Каждая Сторона обязуется обеспечивать конфиденциальность полученной ею в связи с заключением или исполнением Договора от другой Стороны информации ограниченного доступа.')
        para(f'10.3. Заказчик дает свое согласие на обработку своих персональных данных (ФИО, паспортные данные, адрес места жительства, фотографии изделий, номер телефона, адрес электронной почты) {co_name} (ОГРН: {co_ogrn}, ИНН: {co_inn}). Цель обработки: исполнение настоящего договора. Срок согласия — в течение срока действия договора и 12 месяцев после его окончания.')
        para('10.4. К настоящему Договору прилагаются и являются неотъемлемой частью следующие приложения: 1. Технический проект. 2. Калькуляция работ. 3. Правила эксплуатации корпусной мебели. 4. Образец Акта выполненных работ.')

        sec_title('11. РЕКВИЗИТЫ СТОРОН')
        sig_table(
            ['Подрядчик', f'{co_name}\nОГРН: {co_ogrn}, ИНН: {co_inn}\n{co_address}\n\nМенеджер: {manager_line}\n\nПодпись: ______________________________\nМ.П.'],
            ['Заказчик', f'{fname}\nПаспорт: {_passport_str(c)}\nВыдан: {c.get("passport_issued_by") or "___________"}\nДата выдачи: {_fmt_date(c.get("passport_issued_date") or "")}\nКод подразделения: {c.get("passport_dept_code") or "___________"}\nАдрес регистрации: {_reg_address(c)}\nТел.: {c.get("phone") or "___________"}\nEmail: {c.get("email") or "___________"}\n\nПодпись: ______________________________']
        )

    elif doc_type == 'act':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f'Приложение № 4 к договору бытового подряда\nна изготовление мебели от {contract_date_full}')
        h('«Акт выполненных работ»')
        doc.add_paragraph(f'г. {co_city}')
        para(f'{co_name}, в лице {co_dir_pos}а ______________________________, действующего на основании доверенности № ____ от ____________, именуемый в дальнейшем «Подрядчик», и гр. {fname}, именуемый (ая) в дальнейшем «Заказчик», подписали настоящий Акт выполненных работ о нижеследующем:', indent=False)
        para(f'1. Подрядчик изготовил для Заказчика мебель по договору бытового подряда на изготовление мебели от {contract_date_full}:')

        t = doc.add_table(rows=1, cols=5)
        t.style = 'Table Grid'
        for i, txt in enumerate(['№', 'Наименование мебели', 'Ед. изм.', 'Кол-во', 'Стоимость, руб.']):
            t.rows[0].cells[i].text = txt
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        if products:
            for i, pi in enumerate(products, 1):
                r = t.add_row().cells
                r[0].text = str(i); r[1].text = pi.get('name','Кухонный гарнитур')
                r[2].text = 'шт.'; r[3].text = str(pi.get('qty',1)); r[4].text = ''
        else:
            r = t.add_row().cells
            r[0].text = '1'; r[1].text = 'Кухонный гарнитур'; r[2].text = 'шт.'; r[3].text = '1'; r[4].text = f'{total:,.0f}'
        tr = t.add_row().cells
        tr[3].text = 'ИТОГО:'; tr[3].paragraphs[0].runs[0].bold = True
        tr[4].text = f'{total:,.0f}'; tr[4].paragraphs[0].runs[0].bold = True

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(f'({total_words})')
        para('2. Комплектность, количество, вид, характеристики мебели соответствуют условиям договора. Визуальный осмотр мебели на предмет повреждений, царапин, сколов, трещин и других недостатков произведён Заказчиком. Фурнитура (петли, выдвижные механизмы, подъёмники) проверена на предмет работоспособности.')
        para('3. Заказчик принял результат выполненных работ. Претензий к качеству и комплектности мебели не имеет.')
        para('4. Обязательства Подрядчика по Договору выполнены в полном объёме. Оплата по Договору произведена Заказчиком в полном объёме.')
        sig_table(
            ['Подрядчик', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.'],
            ['Заказчик', f'{fname}\nПаспорт: {_passport_str(c)}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________']
        )

    elif doc_type == 'tech':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f'Приложение № 1 к договору бытового подряда\nна изготовление мебели от {contract_date_full}')
        h('«Технический проект»')
        t = doc.add_table(rows=9, cols=2); t.style = 'Table Grid'
        fields = ['Корпус:','Фасад 1:','Фасад 2:','Столешница:','Стеновая панель:','Подсветка:','Фрезеровка:','Примечание:']
        for i, f in enumerate(fields):
            t.cell(i, 0).text = f; t.cell(i, 0).paragraphs[0].runs[0].bold = True
            t.cell(i, 1).text = ' '
        doc.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\nМесто для схемы / эскиза:')
        sig_table(
            ['Подрядчик', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.'],
            ['Заказчик', f'{fname}\n\nПодпись: ______________________________']
        )

    elif doc_type == 'delivery':
        daddr = _delivery_address(c)
        floor_ = c.get('delivery_floor') or '___'
        elevator = c.get('delivery_elevator') or 'нет'
        entrance = c.get('delivery_entrance') or '___'
        delivery_date_str = _fmt_date(c.get('delivery_date') or '')

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('ДОГОВОР').bold = True
        h2('на оказание услуг по доставке мебели')
        doc.add_paragraph(f'г. {co_city}                                                          {contract_date_full}')
        para(f'{co_name}, в лице {co_dir_pos}а ______________________________, действующего на основании доверенности № ____ от ____________, именуемый в дальнейшем «Исполнитель», и гр. {fname}, именуемый (ая) в дальнейшем «Заказчик», заключили настоящий Договор о нижеследующем:', indent=False)
        sec_title('1. ПРЕДМЕТ ДОГОВОРА')
        para(f'1.1. Исполнитель обязуется оказать услуги по доставке мебели по договору бытового подряда № {contract_num} от {contract_date_full}, а Заказчик обязуется принять и оплатить указанные услуги.')
        para(f'1.2. Мебель доставляется по адресу: {daddr}, подъезд: {entrance}, этаж: {floor_}, лифт: {elevator}.')
        para(f'1.3. Дата доставки: {delivery_date_str}. Конкретное время доставки согласовывается Сторонами дополнительно.')
        para('1.4. Доставка осуществляется в разобранном виде (без сборки и монтажа), если иное не предусмотрено отдельным договором.')
        sec_title('2. ПРАВА И ОБЯЗАННОСТИ СТОРОН')
        para('2.1. Исполнитель обязан:')
        para('2.1.1. Доставить мебель в согласованный срок по указанному адресу.')
        para('2.1.2. Обеспечить сохранность мебели при транспортировке.')
        para('2.1.3. Уведомить Заказчика об изменении времени доставки не менее чем за 2 часа.')
        para('2.2. Заказчик обязан:')
        para('2.2.1. Обеспечить свободный доступ к месту доставки и наличие ответственного лица для приёмки.')
        para('2.2.2. Оплатить услуги по доставке в соответствии с Приложением № 1.')
        para('2.2.3. При обнаружении внешних повреждений мебели в ходе доставки — незамедлительно сообщить Исполнителю до подписания акта приёма-передачи.')
        para('2.3. Исполнитель вправе однократно предоставить Заказчику скидку в размере стоимости доставки (8 000 руб.) при доставке в пределах г. Саратова и г. Энгельса.')
        sec_title('3. СТОИМОСТЬ УСЛУГ И ПОРЯДОК ОПЛАТЫ')
        para('3.1. Стоимость услуг определяется в соответствии с Приложением № 1 к настоящему Договору.')
        para('3.2. Оплата производится в день доставки до начала разгрузки мебели, если иное не согласовано Сторонами.')
        para('3.3. Расчёт стоимости доставки за пределы г. Саратова и г. Энгельса производится исходя из километража от склада Исполнителя (г. Саратов, ул. Усть-Курдюмская, д. 3) до адреса Заказчика согласно Яндекс Картам.')
        sec_title('4. ОТВЕТСТВЕННОСТЬ СТОРОН')
        para('4.1. В случае повреждения мебели по вине Исполнителя при транспортировке Исполнитель обязан возместить причинённый ущерб.')
        para('4.2. Исполнитель не несёт ответственности за повреждения, произошедшие по вине Заказчика.')
        para('4.3. В случае отказа Заказчика от доставки менее чем за 24 часа Исполнитель вправе удержать фактически понесённые расходы.')
        sec_title('5. ПРОЧИЕ УСЛОВИЯ')
        para('5.1. Настоящий Договор вступает в силу с момента подписания и действует до полного исполнения Сторонами своих обязательств.')
        para('5.2. Договор составлен в двух экземплярах, имеющих равную юридическую силу.')
        sig_table(
            ['Исполнитель', f'{co_name}\n{co_address}\nИНН/КПП: {co_inn}/{co_kpp}' + (f'\nр/с: {co_rs}' if co_rs else '') + (f'\nБИК: {co_bik}' if co_bik else '') + '\n\nМенеджер: ______________________________\nМ.П.'],
            ['Заказчик', f'{fname}\nПаспорт: {_passport_str(c)}\nАдрес доставки: {daddr}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________']
        )
        doc.add_paragraph()
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run(f'Приложение № 1 к договору на оказание услуг по доставке мебели от {contract_date_full}')
        h2('«Калькуляция на выполнение услуг по доставке мебели»')
        t2 = doc.add_table(rows=1, cols=5); t2.style = 'Table Grid'
        for i, txt in enumerate(['Наименование работ и услуг', 'Ед. изм.', 'Кол-во', 'Цена, руб.', 'Стоимость, руб.']):
            t2.rows[0].cells[i].text = txt; t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        r1 = t2.add_row().cells; r1[0].text = 'Доставка мебели в пределах г. Саратова и г. Энгельса *'; r1[1].text = '1 услуга'; r1[2].text = '1'; r1[3].text = '8 000'; r1[4].text = '8 000'
        r2 = t2.add_row().cells; r2[0].text = 'Доставка мебели за пределы г. Саратова и г. Энгельса **'; r2[1].text = '1 км'; r2[2].text = '___'; r2[3].text = '70'; r2[4].text = '___'
        rt = t2.add_row().cells; rt[3].text = 'Итого:'; rt[3].paragraphs[0].runs[0].bold = True; rt[4].text = '___________'
        rs = t2.add_row().cells; rs[3].text = 'Скидка ***:'; rs[4].text = '___________'
        rts = t2.add_row().cells; rts[3].text = 'Итого со скидкой:'; rts[3].paragraphs[0].runs[0].bold = True; rts[4].text = '___________'

    elif doc_type == 'assembly':
        daddr = _delivery_address(c)
        assembly_days_n = int(c.get('assembly_days') or 1)
        delivery_date_str = _fmt_date(c.get('delivery_date') or '')

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('ДОГОВОР').bold = True
        h2('на оказание услуг по сборке и монтажу мебели')
        doc.add_paragraph(f'г. {co_city}                                                          {contract_date_full}')
        para(f'{co_name}, в лице {co_dir_pos}а ______________________________, действующего на основании доверенности № ____ от ____________, именуемый в дальнейшем «Исполнитель», и гр. {fname}, именуемый (ая) в дальнейшем «Заказчик», заключили настоящий Договор о нижеследующем:', indent=False)
        sec_title('1. ПРЕДМЕТ ДОГОВОРА')
        para(f'1.1. Исполнитель обязуется выполнить работы по сборке и монтажу мебели, изготовленной по договору бытового подряда № {contract_num} от {contract_date_full}, а Заказчик обязуется принять и оплатить указанные работы.')
        para(f'1.2. Адрес выполнения работ: {daddr}.')
        para(f'1.3. Ориентировочная дата начала работ: {delivery_date_str}. Срок выполнения работ: {assembly_days_n} ({_days_words(assembly_days_n)}) рабочих дней.')
        para('1.4. В объём работ входит: сборка всех корпусных элементов, установка фасадов и фурнитуры, регулировка петель и выдвижных систем, монтаж столешницы и стеновых панелей (при наличии), подключение подсветки (при наличии).')
        sec_title('2. ПРАВА И ОБЯЗАННОСТИ СТОРОН')
        para('2.1. Исполнитель обязан:')
        para('2.1.1. Выполнить монтажные работы качественно, в соответствии с техническим проектом.')
        para('2.1.2. Соблюдать чистоту на месте производства работ, убрать строительный мусор по окончании монтажа.')
        para('2.1.3. Уведомить Заказчика об обнаруженных в ходе монтажа дефектах мебели или несоответствии помещения.')
        para('2.2. Заказчик обязан:')
        para('2.2.1. Обеспечить свободный доступ к месту монтажа и освободить помещение от посторонних предметов.')
        para('2.2.2. Обеспечить наличие электроснабжения и освещения в зоне монтажа.')
        para('2.2.3. Принять выполненные работы и подписать акт приёмки.')
        para('2.2.4. Оплатить работы по монтажу в соответствии с условиями настоящего Договора.')
        sec_title('3. СТОИМОСТЬ РАБОТ И ПОРЯДОК ОПЛАТЫ')
        para('3.1. Стоимость работ по сборке и монтажу мебели составляет ________________________________ рублей.')
        para('3.2. Оплата производится в день завершения монтажных работ до подписания акта приёмки, если иное не согласовано Сторонами.')
        sec_title('4. ОТВЕТСТВЕННОСТЬ СТОРОН')
        para('4.1. Исполнитель несёт ответственность за качество выполненных монтажных работ в течение гарантийного срока — 12 (двенадцати) месяцев с момента подписания акта приёмки.')
        para('4.2. Гарантия Исполнителя не распространяется на дефекты, возникшие вследствие нарушения Заказчиком Правил эксплуатации мебели, механических повреждений или воздействия влаги.')
        para('4.3. Исполнитель не несёт ответственности за невозможность выполнения монтажа по причинам, зависящим от Заказчика (неготовность помещения, отсутствие электроснабжения и т.п.).')
        sec_title('5. ПОРЯДОК СДАЧИ И ПРИЁМКИ РАБОТ')
        para('5.1. Приёмка выполненных работ оформляется подписанием Сторонами акта приёмки смонтированной мебели.')
        para('5.2. При наличии замечаний Заказчик фиксирует их в акте. Исполнитель устраняет обоснованные замечания в согласованные сроки.')
        sec_title('6. ПРОЧИЕ УСЛОВИЯ')
        para('6.1. Настоящий Договор вступает в силу с момента подписания и действует до полного исполнения Сторонами своих обязательств.')
        para('6.2. Договор составлен в двух экземплярах, имеющих равную юридическую силу.')
        sig_table(
            ['Исполнитель', f'{co_name}\n{co_address}\nИНН/КПП: {co_inn}/{co_kpp}' + (f'\nр/с: {co_rs}' if co_rs else '') + (f'\nБИК: {co_bik}' if co_bik else '') + '\n\nМенеджер: ______________________________\nМ.П.'],
            ['Заказчик', f'{fname}\nПаспорт: {_passport_str(c)}\nАдрес монтажа: {daddr}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________']
        )

    elif doc_type == 'rules':
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f'Приложение № 3 к договору бытового подряда на изготовление мебели от {contract_date_full}')
        h('«Правила эксплуатации корпусной мебели»')

        sec_title('1. ОБЩИЕ РЕКОМЕНДАЦИИ')
        para('1.1. Срок службы Мебели и сохранение его потребительских свойств напрямую зависят от соблюдения Заказчиком правил, изложенных в настоящем приложении.')
        para('1.2. Климатические условия и воздействия окружающей среды (свет, влажность, температура) напрямую влияют на состояние мебели.')
        para('1.3. Подрядчик рекомендует соблюдать оптимальные климатические условия в помещении, где установлена мебель: температура воздуха от +18°C до +25°C, относительная влажность воздуха 45% - 70%.')
        para('1.4. Следует оберегать мебель от длительного воздействия прямых солнечных лучей, источников тепла (батареи, обогреватели, духовые шкафы, плиты на расстоянии менее 0,5 м), а также от резких перепадов температуры и влажности.')
        para('1.5. Запрещается воздействие на Мебель агрессивных жидкостей (кислот, щелочей, растворителей), абразивных чистящих средств и материалов, способных повредить покрытие.')
        para('1.6. Мебель предназначена для эксплуатации в жилых или общественных помещениях в соответствии с ее функциональным назначением. Подрядчик не несет ответственности за повреждения, вызванные несоблюдением рекомендуемых условий.')
        para('1.7. Подрядчик гарантирует соответствие Мебели обязательным требованиям нормативных документов, действующих на территории РФ, в том числе:\n* ГОСТ 19917-2014 «Мебель для сидения и лежания. Общие технические условия»;\n* ГОСТ 32289-2013 «Плиты древесно-стружечные, облицованные пленками на основе термореактивных полимеров. Технические условия»;\n* ГОСТ 16371-2014 «Мебель. Общие технические условия»;\n* ТР ТС 025/2012 «О безопасности мебельной продукции».')

        sec_title('2. УСЛОВИЯ ХРАНЕНИЯ')
        para('2.1. Хранение Мебели до ее установки должно осуществляться в сухих, проветриваемых, отапливаемых помещениях, защищенных от атмосферных осадков, прямых солнечных лучей и источников тепла. Запрещается хранить изделия в помещениях с повышенной влажностью (санузлы, бани, подвалы без отопления).')
        para('2.2. Рекомендуемый диапазон температуры воздуха от +10°C до +25°C. Запрещена эксплуатация при температурах ниже -20°C и выше +40°C, а также резкие перепады температур.')
        para('2.3. Рекомендуемая относительная влажность воздуха 45% - 70%. Длительное воздействие повышенной влажности (свыше 80%) или чрезмерной сухости (ниже 40%) недопустимо и приводит к деформации элементов (разбуханию, рассыханию, расслоению).')
        para('2.4. Мебель должна храниться в заводской упаковке на ровной поверхности. Запрещается хранить изделия в вертикальном положении, прислонив к стене.')
        para('2.5. Перед установкой Мебели, доставленной или хранившейся при отрицательных температурах, необходимо выдержать ее в оригинальной упаковке в условиях помещения не менее 72 часов (3 суток) для адаптации к комнатной температуре и влажности.')

        sec_title('3. ПРАВИЛА ЭКСПЛУАТАЦИИ И УХОДА ЗА МЕБЕЛЬЮ')
        para('3.1. Общие правила:', indent=False)
        para('3.1.1. Не превышайте максимально допустимые нагрузки на элементы Мебели: полки - до 15 кг, выдвижные ящики из ЛДСП - до 5 кг, ящики на системе «метабокс» - до 18 кг, ящики на системе «тандембокс» - до 35 кг.')
        para('3.1.2. Общая нагрузка на подвесной шкаф не должна превышать 70 кг. Высокие конструкции (пеналы, стеллажи) должны быть больше нагружены в нижней части.')
        para('3.1.3. Равномерно распределяйте нагрузку внутри шкафов и ящиков: тяжелые предметы размещайте ближе к краям и опорам, легкие - в центре.')
        para('3.1.4. Исключите попадание воды на незащищенные торцы деталей, места стыков и торцов изделия.')
        para('3.1.5. Запрещается воздействовать на поверхности абразивными, кислотными, щелочными средствами, ацетоном, растворителями.')
        para('3.1.6. Не допускается заслонять вентиляционные решетки и воздухозаборные отверстия бытовых приборов, встроенных в мебель.')
        para('3.1.7. При установке Мебели в бревенчатых домах, вследствие усадки таких домов, возможна деформация: перекашивание фасадов, опускание как верхних, так и нижних модулей, это не является признаком некачественной работы Подрядчика.')
        para('3.2. Корпус ЛДСП:', indent=False)
        para('3.2.1. Главное правило! ЛДСП боится длительного контакта с водой. Попадание жидкости на кромки и тем более в места стыков недопустимо – плита разбухнет и деформируется. Все пролитые жидкости следует немедленно вытирать насухо.')
        para('3.2.2. Запрещается мыть ЛДСП большим количеством воды, использовать пароочистители. Не оставляйте не отжатые тряпки, мокрые полотенца на торцах деталей.')
        para('3.2.3. Запрещается применять абразивные и едкие химические средства.')
        para('3.2.4. Протирайте сухой или слегка влажной хорошо отжатой тканью. Для удаления загрязнений использовать мягкие мыльные растворы.')
        para('3.2.5. Берегите кромки и торцы от сколов и ударов. Именно через поврежденные кромки влага легче всего проникает внутрь плиты.')
        para('3.2.6. Не превышайте максимально допустимые нагрузки на полки и перегородки. Равномерно распределяйте вес.')
        para('3.2.7. Избегайте расположения Мебели вплотную к отопительным приборам (батареям, обогревателям), это может привести к расслоению ламинированного слоя и деформации плиты.')
        para('3.3. Фасады МДФ:', indent=False)
        para('3.3.1. Избегайте резких перепадов температуры и попадания прямых солнечных лучей.')
        para('3.3.2. Не допускайте длительного контакта поверхности с сильно нагретыми предметами (сковородки, кастрюли, утюги).')
        para('3.3.3. Избегайте ударов, царапин острыми предметами, давления на выступающие элементы.')
        para('3.3.4. Исключите контакт с агрессивными химическими веществами: растворителями, ацетоном, сильными чистящими порошками, средствами для мытья стекол. Регулярно удаляйте пыль мягкой тканью.')
        para('3.3.5. Разрешается использовать нейтральные (pH-нейтральные) средства, а также мягкие мыльные растворы.')
        para('3.3.6. Запрещается мыть МДФ большим количеством воды, использовать пароочистители. Не оставляйте не отжатые тряпки, мокрые полотенца на торцах деталей.')
        para('3.3.7. Не рекомендуется снимать защитную пленку с фасадов до окончания процесса установки мебели.')
        para('3.3.8. При изготовлении дверных полотен высотой более 1600 мм, рекомендовано использовать систему выпрямления дверей, чтобы избежать возможной деформации изделия.')
        para('3.4. Столешницы и стеновые панели из пластика и искусственного (акрилового) камня:', indent=False)
        para('3.4.1. Для очистки используйте мягкую ткань, мягкую губку, салфетки из микрофибры и средства для ухода за глянцевыми поверхностями.')
        para('3.4.2. Запрещается использовать жесткие и металлические губки, щетки и абразивные чистящие средства. Использовать при очистке кислоты, щелочи, соли, растворители.')
        para('3.4.5. Запрещается использовать поверхность в качестве разделочной доски. Воздействовать на поверхность столешницы острыми предметами. Передвигать по поверхности посуду с металлическим дном.')
        para('3.4.6. На пластиковой столешнице не оставляйте лужи воды, не отжатые тряпки, мокрые полотенца, особенно на стыках и возле мойки. Сразу удаляйте воду с поверхностей столешницы и со стыковочных швов. Еврозапил особенно уязвим для влаги и высоких температур.')
        para('3.4.7. Не допускается ставить на поверхность столешницы горячие (>60°С) предметы. Используйте специальные термоизоляционные подставки под горячее. Не допускается резкий перепад температур.')
        para('3.4.8. Не допускается размораживать продукты, или оставлять на длительное время на поверхности столешницы сильно охлаждённые (<0°C) предметы.')
        para('3.6. Стеклянные и зеркальные поверхности:', indent=False)
        para('3.6.1. Очищайте специальными средствами для стекол и зеркал, нанося состав на мягкую ткань, а не прямо на поверхность.')
        para('3.6.2. Избегайте абразивных средств и жестких губок. Не допускайте ударных и чрезмерных нагрузок на стеклянные полки.')
        para('3.7. Фурнитура и механизмы:', indent=False)
        para('3.7.1. Протирание сухой мягкой тканью. Для удаления загрязнений допускается использование слабого мыльного раствора.')
        para('3.7.2. Запрещается прилагать чрезмерные усилия для открывания/закрывания дверей и ящиков. Угол открывания распашных дверей, как правило, составляет 90–110°. Превышение угла открывания может привести к поломке петли.')
        para('3.7.3. Запрещается открывать фасады с системой «Push to Open» любым способом, кроме нажатия на фасад.')
        para('3.7.4. Выдвигайте ящики полностью только при необходимости, держась за ручки или фасад.')
        para('3.7.5. Регулярно удаляйте пыль, крошки и грязь с направляющих и других движущихся частей. Следите, чтобы в зону работы механизма не попадали посторонние предметы, которые могут помешать движению.')
        para('3.7.6. Регулярная регулировка и смазка механизмов (парафином или специальными средствами) являются обязанностью Заказчика и не покрываются гарантией.')
        para('3.7.7. Лицевую фурнитуру следует чистить мягкими тканями с применением хозяйственного мыла, после чего вытирать насухо. Не использовать средства, содержащие абразивные материалы (наждачную бумагу, соду и др.).')

        p_italic = doc.add_paragraph()
        p_italic.add_run('Соблюдая эти несложные правила, вы сохраните безупречный вид и функциональность вашей мебели на долгие годы. Помните: что несоблюдение правил эксплуатации может привести к сокращению сроков службы и преждевременному выходу из строя элементов кухонной мебели.').italic = True

        p_warn = doc.add_paragraph()
        r_warn = p_warn.add_run('ВНИМАНИЕ! Подрядчик не несет ответственность за последствия от несоблюдения установленных норм и правил по уходу и эксплуатации корпусной мебели.')
        r_warn.bold = True; r_warn.italic = True

        sig_table(
            ['Подрядчик', f'{co_name.upper()}\n\n\n\nМ.П.'],
            ['Заказчик', '\n\n\n\n']
        )

    elif doc_type == 'act_delivery':
        daddr = _delivery_address(c)
        delivery_date_str = _fmt_date(c.get('delivery_date') or '')
        dcost = float(c.get('delivery_cost') or 0)
        dcost_words = _num_to_words(dcost) if dcost else '___________'

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f'к договору на оказание услуг по доставке мебели от {contract_date_full}')
        h('«Акт приёма-передачи доставки мебели»')
        doc.add_paragraph(f'г. {co_city}')

        para(f'{co_name}, в лице {co_dir_pos}а ______________________________, действующего на основании доверенности № ____ от ____________, именуемый в дальнейшем «Исполнитель», и гр. {fname}, именуемый (ая) в дальнейшем «Заказчик», составили настоящий Акт о нижеследующем:', indent=False)
        para(f'1. Исполнитель доставил Заказчику мебель по адресу: {daddr}.')
        para(f'2. Дата доставки: {delivery_date_str}.')
        para('3. Мебель доставлена в полном объёме, внешних механических повреждений при доставке не выявлено.')
        para('4. Заказчик произвёл визуальный осмотр доставленной мебели в момент приёмки. Претензий к состоянию мебели на момент доставки не имеет.')
        para(f'5. Стоимость услуг по доставке составила {dcost:,.0f} ({dcost_words}) рублей. Оплата произведена в полном объёме.')
        para('6. Услуги по доставке выполнены Исполнителем в полном объёме.')
        sig_table(
            ['Исполнитель', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.'],
            ['Заказчик', f'{fname}\nПаспорт: {_passport_str(c)}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________']
        )

    elif doc_type == 'act_assembly':
        daddr = _delivery_address(c)
        acost = float(c.get('assembly_cost') or 0)
        acost_words = _num_to_words(acost) if acost else '___________'
        products = _get_products(c)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f'к договору на оказание услуг по сборке и монтажу мебели от {contract_date_full}')
        h('«Акт приёма-передачи выполненных работ по сборке и монтажу мебели»')
        doc.add_paragraph(f'г. {co_city}')

        para(f'{co_name}, в лице {co_dir_pos}а ______________________________, действующего на основании доверенности № ____ от ____________, именуемый в дальнейшем «Исполнитель», и гр. {fname}, именуемый (ая) в дальнейшем «Заказчик», составили настоящий Акт о нижеследующем:', indent=False)
        para(f'1. Исполнитель выполнил работы по сборке и монтажу мебели по адресу: {daddr}:')

        t_a = doc.add_table(rows=1, cols=4); t_a.style = 'Table Grid'
        for i, txt in enumerate(['№', 'Наименование мебели', 'Ед. изм.', 'Кол-во']):
            t_a.rows[0].cells[i].text = txt; t_a.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        if products:
            for i, pi in enumerate(products, 1):
                r = t_a.add_row().cells
                r[0].text = str(i); r[1].text = pi.get('name', 'Кухонный гарнитур'); r[2].text = 'шт.'; r[3].text = str(pi.get('qty', 1))
        else:
            r = t_a.add_row().cells; r[0].text = '1'; r[1].text = 'Кухонный гарнитур'; r[2].text = 'шт.'; r[3].text = '1'

        para('2. Объём выполненных работ: сборка корпусных элементов, установка фасадов и фурнитуры, регулировка петель и выдвижных систем, монтаж столешницы и стеновых панелей (при наличии), подключение подсветки (при наличии).')
        para('3. Заказчик произвёл проверку выполненных работ. Фурнитура проверена на предмет работоспособности. Регулировка произведена в присутствии Заказчика.')
        para('4. Работы по сборке и монтажу выполнены Исполнителем качественно, в соответствии с техническим проектом. Претензий к качеству и объёму выполненных работ Заказчик не имеет.')
        para(f'5. Стоимость работ по сборке и монтажу составила {acost:,.0f} ({acost_words}) рублей. Оплата произведена в полном объёме.')
        para('6. С момента подписания настоящего Акта гарантийный срок на работы по сборке и монтажу составляет 12 (двенадцать) месяцев.')
        sig_table(
            ['Исполнитель', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.'],
            ['Заказчик', f'{fname}\nПаспорт: {_passport_str(c)}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________']
        )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()