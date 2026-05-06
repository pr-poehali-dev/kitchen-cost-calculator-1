"""
Генератор документов DOCX.
Каждый тип документа — отдельная функция с чистой вёрсткой.
PDF — через HTML + window.print() на фронтенде.
"""
import json, os, base64, logging
import jwt, psycopg2, boto3
from contextlib import contextmanager
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)

JWT_SECRET = os.environ['JWT_SECRET']
S3_KEY     = os.environ.get('AWS_ACCESS_KEY_ID', '')
S3_SECRET  = os.environ.get('AWS_SECRET_ACCESS_KEY', '')

CORS = {
    'Access-Control-Allow-Origin': '*',
    'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
    'Access-Control-Allow-Headers': 'Content-Type, Authorization, X-Authorization',
}

# ── helpers ───────────────────────────────────────────────────────────────────

@contextmanager
def get_db():
    conn = psycopg2.connect(os.environ['DATABASE_URL'])
    try:
        yield conn; conn.commit()
    except Exception:
        conn.rollback(); raise
    finally:
        conn.close()

def verify_token(event):
    h = event.get('headers') or {}
    auth = h.get('X-Authorization') or h.get('Authorization') or ''
    token = auth[7:].strip() if auth.startswith('Bearer ') else ''
    if not token: return None
    try: return jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
    except: return None

def ok(data, status=200):
    return {'statusCode': status, 'headers': {**CORS, 'Content-Type': 'application/json'}, 'body': json.dumps(data, default=str)}

def err(msg, status=400):
    return {'statusCode': status, 'headers': {**CORS, 'Content-Type': 'application/json'}, 'body': json.dumps({'error': msg})}

def fmt_date(d):
    if not d: return '___________'
    try: return datetime.strptime(str(d)[:10], '%Y-%m-%d').strftime('%d.%m.%Y')
    except: return str(d)

def fmt_date_full(d):
    months = ['января','февраля','марта','апреля','мая','июня','июля','августа','сентября','октября','ноября','декабря']
    if not d: return '«___» ___________ ______ г.'
    try:
        dt = datetime.strptime(str(d)[:10], '%Y-%m-%d')
        return f'«{dt.day:02d}» {months[dt.month-1]} {dt.year} г.'
    except: return str(d)

def full_name(c):
    return ' '.join(filter(None, [c.get('last_name',''), c.get('first_name',''), c.get('middle_name','')])).strip() or '___________'

def genitive_name(full_name_str: str) -> str:
    """Склоняет ФИО (Фамилия Имя Отчество) в родительный падеж."""
    if not full_name_str or not full_name_str.strip():
        return full_name_str
    parts = full_name_str.strip().split()
    if not parts:
        return full_name_str
    consonants = set('бвгджзйклмнпрстфхцчшщБВГДЖЗЙКЛМНПРСТФХЦЧШЩ')

    def _word_gen(word):
        if not word or len(word) < 2:
            return word
        low = word.lower()
        if low.endswith('ий'): return word[:-2] + 'ия'
        if low.endswith('ия'): return word[:-2] + 'ии'
        if low.endswith('ья'): return word[:-2] + 'ьи'
        if low.endswith('ич'): return word + 'а'
        if low.endswith('енко') or low.endswith('ко') or low.endswith('но') or low.endswith('го'): return word
        if low.endswith('а') and len(word) >= 2 and word[-2] in consonants: return word[:-1] + 'ы'
        if low.endswith('я') and len(word) >= 2 and word[-2] in consonants: return word[:-1] + 'и'
        if low.endswith('ь'): return word[:-1] + 'я'
        if low.endswith('ец') and len(word) >= 4: return word[:-2] + 'ца'
        if word[-1] in consonants: return word + 'а'
        return word

    return ' '.join(_word_gen(p) for p in parts)

def get_manager_poa(manager_name: str) -> dict:
    """Запрашивает из БД данные доверенности менеджера по его имени."""
    if not manager_name or not manager_name.strip():
        return {}
    try:
        schema = os.environ.get('MAIN_DB_SCHEMA', 'public')
        name = manager_name.strip()
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute(
                f'''SELECT full_name, poa_number, poa_date FROM {schema}.users
                   WHERE LOWER(full_name) = LOWER(%s) OR LOWER(login) = LOWER(%s)
                   LIMIT 1''',
                (name, name)
            )
            row = cur.fetchone()
        if not row:
            return {}
        return {
            'full_name': row[0] or '',
            'poa_number': str(row[1]) if row[1] else '',
            'poa_date': str(row[2])[:10] if row[2] else '',
        }
    except Exception:
        return {}

def full_name_genitive(c):
    """ФИО в родительном падеже (для «гр. Иванова Петра Сергеевича»)."""
    ln = (c.get('last_name') or '').strip()
    fn = (c.get('first_name') or '').strip()
    mn = (c.get('middle_name') or '').strip()
    if not ln: return '___________'

    def decline_last(word, fn, mn):
        if not word: return word
        # Определяем пол по отчеству
        if mn.endswith('овна') or mn.endswith('евна') or mn.endswith('ична') or mn.endswith('инична'):
            gender = 'f'
        elif mn.endswith('ович') or mn.endswith('евич') or mn.endswith('ич'):
            gender = 'm'
        else:
            gender = 'm'
        w = word
        if gender == 'm':
            if w.endswith('ий'): return w[:-2] + 'ого'
            if w.endswith('ый'): return w[:-2] + 'ого'
            if w.endswith('ой'): return w[:-2] + 'ого'
            if w[-1] in 'аяь': return w
            if w[-1] not in 'ьъ': return w + 'а'
        else:
            if w.endswith('ая'): return w[:-2] + 'ой'
            if w.endswith('яя'): return w[:-2] + 'ей'
            if w.endswith('а'): return w[:-1] + 'ой'
            if w.endswith('я'): return w[:-1] + 'и'
        return w

    def decline_first(word, gender='m'):
        if not word: return word
        w = word
        if gender == 'm':
            if w.endswith('й'): return w[:-1] + 'я'
            if w.endswith('ь'): return w[:-1] + 'я'
            if w[-1] not in 'аяь': return w + 'а'
            return w
        else:
            if w.endswith('а'): return w[:-1] + 'ы'
            if w.endswith('я'): return w[:-1] + 'и'
            return w

    def decline_middle(word):
        if not word: return word
        w = word
        if w.endswith('ович') or w.endswith('евич'): return w + 'а'
        if w.endswith('ич'): return w + 'а'
        if w.endswith('овна') or w.endswith('евна'): return w[:-1] + 'ы'
        if w.endswith('ична') or w.endswith('инична'): return w[:-1] + 'ы'
        return w

    if mn.endswith('овна') or mn.endswith('евна') or mn.endswith('ична') or mn.endswith('инична'):
        gender = 'f'
    else:
        gender = 'm'

    ln_g = decline_last(ln, fn, mn)
    fn_g = decline_first(fn, gender)
    mn_g = decline_middle(mn)
    return ' '.join(filter(None, [ln_g, fn_g, mn_g]))

def delivery_addr(c):
    parts = [c.get('delivery_city',''), c.get('delivery_street',''), c.get('delivery_house','')]
    apt = c.get('delivery_apt','')
    if apt: parts.append(f'кв. {apt}')
    return ', '.join(p for p in parts if p) or '___________'

def get_company(user_id):
    try:
        schema = os.environ.get('MAIN_DB_SCHEMA', 'public')
        with get_db() as conn:
            cur = conn.cursor()
            # Пробуем поле state (основное), затем settings (старое)
            cur.execute(f'SELECT state FROM {schema}.app_state WHERE user_id = %s LIMIT 1', (user_id,))
            row = cur.fetchone()
            if row and row[0]:
                s = row[0] if isinstance(row[0], dict) else json.loads(row[0])
                # Данные компании могут лежать в state.settings.company или state.company
                settings = s.get('settings') or s
                company_data = settings.get('company') or {}
                logger.info(f'get_company(state): user_id={user_id}, keys={list(company_data.keys())}')
                if company_data:
                    return company_data
    except Exception as e:
        logger.error(f'get_company error: {e}')
    return {}

def co(company, key, default=''):
    return str(company.get(key) or '').strip() or default

def num_to_words(n):
    n = int(round(float(n or 0)))
    if n == 0: return 'ноль рублей'
    ones = ['','один','два','три','четыре','пять','шесть','семь','восемь','девять','десять','одиннадцать','двенадцать','тринадцать','четырнадцать','пятнадцать','шестнадцать','семнадцать','восемнадцать','девятнадцать']
    ones_f = ['','одна','две','три','четыре','пять','шесть','семь','восемь','девять','десять','одиннадцать','двенадцать','тринадцать','четырнадцать','пятнадцать','шестнадцать','семнадцать','восемнадцать','девятнадцать']
    tens = ['','','двадцать','тридцать','сорок','пятьдесят','шестьдесят','семьдесят','восемьдесят','девяносто']
    hundreds = ['','сто','двести','триста','четыреста','пятьсот','шестьсот','семьсот','восемьсот','девятьсот']
    def chunk(num, fem=False):
        parts=[]
        h=num//100; t=(num%100)//10; o=num%10
        if h: parts.append(hundreds[h])
        if t==1: parts.append((ones_f if fem else ones)[num%100])
        else:
            if t: parts.append(tens[t])
            if o: parts.append((ones_f if fem else ones)[o])
        return parts
    result=[]; millions=n//1_000_000; thousands=(n%1_000_000)//1_000; remainder=n%1_000
    if millions:
        p=chunk(millions); o2=millions%10; t2=(millions%100)//10
        s='миллионов' if (t2==1 or o2==0 or o2>=5) else ('миллион' if o2==1 else 'миллиона')
        result.extend(p); result.append(s)
    if thousands:
        p=chunk(thousands,True); o2=thousands%10; t2=(thousands%100)//10
        s='тысяч' if (t2==1 or o2==0 or o2>=5) else ('тысяча' if o2==1 else 'тысячи')
        result.extend(p); result.append(s)
    if remainder: result.extend(chunk(remainder))
    o2=n%10; t2=(n%100)//10
    rub='рублей' if (t2==1 or o2==0 or o2>=5) else ('рубль' if o2==1 else 'рубля')
    return ' '.join(result)+' '+rub

def passport_str(c):
    s=c.get('passport_series',''); n=c.get('passport_number','')
    if s and n: return f'{s} {n}'
    if n: return n
    return '___________'

def get_products(c):
    p = c.get('products')
    if not p: return []
    if isinstance(p, str):
        try: p = json.loads(p)
        except: return []
    if isinstance(p, list): return [x for x in p if x.get('name','').strip()]
    return []

def apply_vars(text: str, c: dict, company: dict) -> str:
    """Заменяет {{переменные}} шаблона реальными данными клиента и компании."""
    import re
    _total    = float(c.get('total_amount') or 0)
    _prepaid  = float(c.get('prepaid_amount') or 0)
    _balance  = float(c.get('balance_due') or 0) or max(0.0, _total - _prepaid)
    _delivery = float(c.get('delivery_cost') or 0)
    _assembly = float(c.get('assembly_cost') or 0)
    _fmt = lambda n: f"{int(n):,}".replace(',', ' ')

    # Адрес регистрации — может лежать в одном поле или собираться из частей
    _reg_addr = (str(c.get('registration_address') or '').strip()
                 or ', '.join(filter(None, [
                     str(c.get('reg_city') or '').strip(),
                     str(c.get('reg_street') or '').strip(),
                     str(c.get('reg_house') or '').strip(),
                     ('кв. ' + str(c.get('reg_apt'))) if c.get('reg_apt') else '',
                 ])) or '___________')

    # Адрес доставки расширенный
    _del_parts = [
        str(c.get('delivery_city') or '').strip(),
        str(c.get('delivery_street') or '').strip(),
        str(c.get('delivery_house') or '').strip(),
    ]
    if c.get('delivery_apt'):    _del_parts.append(f'кв. {c["delivery_apt"]}')
    if c.get('delivery_floor'):  _del_parts.append(f'эт. {c["delivery_floor"]}')
    if c.get('delivery_entrance'): _del_parts.append(f'подъезд {c["delivery_entrance"]}')
    _del_addr = ', '.join(p for p in _del_parts if p) or '___________'

    vals = {
        # ── Клиент ──────────────────────────────────────────────────────────
        'имя_клиента':           full_name(c),
        'имя_клиента_рп':        full_name_genitive(c),
        'фамилия':               str(c.get('last_name') or '').strip() or '___________',
        'имя':                   str(c.get('first_name') or '').strip() or '___________',
        'отчество':              str(c.get('middle_name') or '').strip() or '',
        'телефон_клиента':       str(c.get('phone') or '').strip() or '___________',
        'телефон':               str(c.get('phone') or '').strip() or '___________',
        'телефон2_клиента':      str(c.get('phone2') or '').strip() or '',
        'телефон2':              str(c.get('phone2') or '').strip() or '',
        'email_клиента':         str(c.get('email') or '').strip() or '',
        'email':                 str(c.get('email') or '').strip() or '',
        'мессенджер':            str(c.get('messenger') or '').strip() or '',
        # ── Паспорт ─────────────────────────────────────────────────────────
        'паспорт':               passport_str(c),
        'паспорт_серия':         str(c.get('passport_series') or '').strip() or '',
        'паспорт_номер':         str(c.get('passport_number') or '').strip() or '',
        'паспорт_выдан':         str(c.get('passport_issued_by') or '').strip() or '___________',
        'паспорт_дата':          fmt_date(c.get('passport_date') or c.get('passport_issued_date') or ''),
        'паспорт_код':           str(c.get('passport_code') or c.get('passport_dept_code') or '').strip() or '',
        # ── Адреса ──────────────────────────────────────────────────────────
        'адрес_регистрации':     _reg_addr,
        'город_клиента':         str(c.get('delivery_city') or c.get('reg_city') or c.get('city') or '').strip() or '___________',
        'адрес_доставки':        _del_addr,
        'город_доставки':        str(c.get('delivery_city') or '').strip() or '___________',
        'улица_доставки':        str(c.get('delivery_street') or '').strip() or '___________',
        'дом_доставки':          str(c.get('delivery_house') or '').strip() or '___________',
        'квартира_доставки':     str(c.get('delivery_apt') or '').strip() or '',
        'этаж_доставки':         str(c.get('delivery_floor') or '').strip() or '',
        'подъезд_доставки':      str(c.get('delivery_entrance') or '').strip() or '',
        'примечание_доставки':   str(c.get('delivery_note') or '').strip() or '',
        # ── Договор ─────────────────────────────────────────────────────────
        'номер_договора':        str(c.get('contract_number') or '___'),
        'дата_договора':         fmt_date_full(c.get('contract_date') or ''),
        'дата_договора_кратко':  fmt_date(c.get('contract_date') or ''),
        'сумма':                 _fmt(_total),
        'сумма_прописью':        num_to_words(_total),
        'аванс':                 _fmt(_prepaid),
        'аванс_прописью':        num_to_words(_prepaid),
        'остаток':               _fmt(_balance),
        'остаток_прописью':      num_to_words(_balance),
        'тип_оплаты':            str(c.get('payment_type') or '').strip() or '',
        'схема_оплаты':          str(c.get('custom_payment_scheme') or '').strip() or '',
        'стоимость_доставки':    _fmt(_delivery),
        'доставка_прописью':     num_to_words(_delivery),
        'стоимость_сборки':      _fmt(_assembly),
        'сборка_прописью':       num_to_words(_assembly),
        # ── Сроки ───────────────────────────────────────────────────────────
        'срок_изготовления':     str(c.get('production_days') or ''),
        'срок_сборки':           str(c.get('assembly_days') or ''),
        'дата_доставки':         fmt_date(c.get('delivery_date') or ''),
        'дата_доставки_полная':  fmt_date_full(c.get('delivery_date') or ''),
        # ── Кредит / рассрочка ──────────────────────────────────────────────
        'номер_кредита':         str(c.get('credit_contract_number') or '').strip() or '',
        'дата_кредита':          fmt_date_full(c.get('credit_contract_date') or ''),
        'банк_кредита':          str(c.get('credit_bank') or '').strip() or '',
        'аванс_кредит':          _fmt(float(c.get('credit_prepaid') or _prepaid)),
        'остаток_кредит':        _fmt(float(c.get('credit_balance') or _balance)),
        # ── Техпроект ───────────────────────────────────────────────────────
        'корпус':                str(c.get('tech_korpus') or '').strip() or '',
        'корпус2':               str(c.get('tech_korpus2') or '').strip() or '',
        'фасад':                 str(c.get('tech_fasad1') or '').strip() or '',
        'фасад2':                str(c.get('tech_fasad2') or '').strip() or '',
        'столешница':            str(c.get('tech_stoleshniza') or '').strip() or '',
        'стеновая':              str(c.get('tech_stenovaya') or '').strip() or '',
        'подсветка':             str(c.get('tech_podsvetka_type') or '').strip() or '',
        'цвет_подсветки':        str(c.get('tech_podsvetka_svet') or '').strip() or '',
        'фрезеровка':            str(c.get('tech_frezerovka') or '').strip() or '',
        # ── Компания ────────────────────────────────────────────────────────
        'компания':              str(company.get('name') or '').strip() or '___________',
        'город':                 str(company.get('city') or c.get('city') or '').strip() or '___________',
        'инн':                   str(company.get('inn') or '').strip() or '',
        'огрн':                  str(company.get('ogrn') or '').strip() or '',
        'кпп':                   str(company.get('kpp') or '').strip() or '',
        'инн_кпп':               (f"{company.get('inn')}/{company.get('kpp')}" if company.get('kpp') else str(company.get('inn') or '')).strip() or '',
        'адрес_компании':        str(company.get('address') or '').strip() or '',
        'телефон_компании':      str(company.get('phone') or '').strip() or '',
        'email_компании':        str(company.get('email') or '').strip() or '',
        'сайт_компании':         str(company.get('website') or '').strip() or '',
        'директор':              str(company.get('director') or '').strip() or '',
        'должность_директора':   str(company.get('directorPosition') or '').strip() or '',
        'банк':                  str(company.get('bank') or '').strip() or '',
        'бик':                   str(company.get('bik') or '').strip() or '',
        'расчётный_счёт':        str(company.get('rs') or '').strip() or '',
        'корр_счёт':             str(company.get('ks') or '').strip() or '',
        # ── Ответственные ───────────────────────────────────────────────────
        'менеджер':              str(c.get('manager_name') or '').strip() or '___________',
        'дизайнер':              str(c.get('designer_name') or c.get('designer') or '').strip() or '___________',
        'замерщик':              str(c.get('measurer') or '').strip() or '',
        # ── Доверенность менеджера ──────────────────────────────────────────
        'номер_доверенности':    str(c.get('manager_poa_number') or '____').strip(),
        'дата_доверенности':     fmt_date(c.get('manager_poa_date') or ''),
        # ── Прописи сроков ──────────────────────────────────────────────────
        'срок_изготовления_прописью': num_to_words(int(c.get('production_days') or 0)).replace(' рублей','').strip(),
        'срок_сборки_прописью':       num_to_words(int(c.get('assembly_days') or 0)).replace(' рублей','').strip(),
    }
    def replace_var(m):
        key = m.group(1).strip()
        return vals.get(key, m.group(0))
    return re.sub(r'\{\{([^}]+)\}\}', replace_var, text)


def render_blocks_to_docx(blocks: list, doc, font_fn, _base_pt: float, _font_name: str, c: dict, company: dict, products: list):
    """Рендерит блоки из конструктора PDF в DOCX документ."""
    from docx.shared import Pt, Mm, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ALIGN_MAP = {
        'left':    WD_ALIGN_PARAGRAPH.LEFT,
        'center':  WD_ALIGN_PARAGRAPH.CENTER,
        'right':   WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    for block in blocks:
        if not block.get('enabled', True):
            continue

        btype   = block.get('type', 'paragraph')
        content = apply_vars(block.get('content', ''), c, company)
        align   = ALIGN_MAP.get(block.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)
        fsize   = float(block.get('fontSize') or _base_pt)
        bold    = bool(block.get('bold', False))
        italic  = bool(block.get('italic', False))
        under   = bool(block.get('underline', False))
        mt      = block.get('marginTop')
        mb      = block.get('marginBottom')

        def make_para(text_content, para_align=None, para_bold=False, para_italic=False, para_under=False, para_size=None):
            p = doc.add_paragraph()
            p.alignment = para_align if para_align is not None else align
            p.paragraph_format.space_before = Pt(float(mt) * 2.835) if mt else Pt(0)
            p.paragraph_format.space_after  = Pt(float(mb) * 2.835) if mb else Pt(1)
            p.paragraph_format.line_spacing = Pt((_base_pt) * 1.25)
            lines = text_content.split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    p.add_run().add_break()
                r = p.add_run(line)
                r._r.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                r.font.name  = _font_name
                r.font.size  = Pt(para_size or fsize)
                r.bold       = para_bold or bold
                r.italic     = para_italic or italic
                r.underline  = para_under or under
            return p

        if btype == 'header':
            make_para(content, WD_ALIGN_PARAGRAPH.CENTER, para_bold=True, para_size=fsize or (_base_pt + 1))

        elif btype == 'section':
            make_para(content, WD_ALIGN_PARAGRAPH.CENTER, para_bold=True)

        elif btype == 'paragraph':
            make_para(content)

        elif btype == 'divider':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(float(mt) * 2.835) if mt else Pt(4)
            p.paragraph_format.space_after  = Pt(float(mb) * 2.835) if mb else Pt(4)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif btype == 'spacer':
            height_px = int(content or 20)
            height_pt = max(2, int(height_px * 0.75))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(height_pt)

        elif btype == 'lines':
            count = int(content or 6) if content.strip().isdigit() else 6
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            for _ in range(count):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(14)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                pBdr.append(bottom)
                pPr.append(pBdr)

        elif btype == 'table':
            rows_raw = [r for r in content.split('\n') if r.strip()]
            if not rows_raw:
                continue
            header_cols = rows_raw[0].split(';')
            body_rows   = [r.split(';') for r in rows_raw[1:]]
            col_widths  = block.get('colWidths') or []
            t = doc.add_table(rows=1 + len(body_rows), cols=len(header_cols))
            t.style = 'Table Grid'
            if col_widths and len(col_widths) == len(header_cols):
                total_width_cm = 16.0
                for ci, w in enumerate(col_widths):
                    col_w = Cm(total_width_cm * w / 100)
                    for row in t.rows:
                        row.cells[ci].width = col_w
            for ci, h in enumerate(header_cols):
                cell = t.cell(0, ci)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(h.strip())
                r.bold = True
                r.font.name = _font_name
                r.font.size = Pt(fsize)
                r._r.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            for ri, row_data in enumerate(body_rows):
                for ci in range(len(header_cols)):
                    val = row_data[ci].strip() if ci < len(row_data) else ''
                    cell = t.cell(ri + 1, ci)
                    p = cell.paragraphs[0]
                    r = p.add_run(val)
                    r.font.name = _font_name
                    r.font.size = Pt(fsize)
                    r._r.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        elif btype == 'calc_table':
            if not products:
                make_para('[Таблица из расчёта — нет данных]', WD_ALIGN_PARAGRAPH.CENTER)
                continue
            cts = block.get('calcTableSettings') or {}
            cols = cts.get('columns') or ['name', 'qty', 'unit', 'total']
            col_labels = {'name':'Наименование','qty':'Кол-во','unit':'Ед.','price':'Цена','total':'Сумма','article':'Артикул','manufacturer':'Производитель'}
            headers = [col_labels.get(col, col) for col in cols]
            rows_data = []
            grand_total = 0
            for prod in products:
                row = []
                for col in cols:
                    if col == 'name':         row.append(str(prod.get('name','') or ''))
                    elif col == 'qty':        row.append(str(prod.get('qty', prod.get('quantity','')) or ''))
                    elif col == 'unit':       row.append(str(prod.get('unit','шт') or 'шт'))
                    elif col == 'price':      row.append(f"{int(float(prod.get('price',0) or 0)):,}".replace(',', ' '))
                    elif col == 'total':      row.append(f"{int(float(prod.get('total', prod.get('amount',0)) or 0)):,}".replace(',', ' '))
                    elif col == 'article':    row.append(str(prod.get('article','') or ''))
                    elif col == 'manufacturer': row.append(str(prod.get('manufacturer','') or ''))
                rows_data.append(row)
                grand_total += float(prod.get('total', prod.get('amount', 0)) or 0)
            t = doc.add_table(rows=1 + len(rows_data), cols=len(cols))
            t.style = 'Table Grid'
            for ci, h in enumerate(headers):
                cell = t.cell(0, ci)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(h)
                r.bold = True
                r.font.name = _font_name
                r.font.size = Pt(fsize)
            for ri, row_data in enumerate(rows_data):
                for ci, val in enumerate(row_data):
                    cell = t.cell(ri + 1, ci)
                    p = cell.paragraphs[0]
                    r = p.add_run(val)
                    r.font.name = _font_name
                    r.font.size = Pt(fsize)
            if cts.get('showTotal', True) and len(cols) >= 2:
                total_row = doc.add_paragraph()
                total_row.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = total_row.add_run(f'Итого: {int(grand_total):,} руб.'.replace(',', ' '))
                r.bold = True
                r.font.name = _font_name
                r.font.size = Pt(fsize)

        elif btype == 'image':
            url = content.strip()
            if url:
                try:
                    import urllib.request, tempfile, os as _os
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tf:
                        urllib.request.urlretrieve(url, tf.name)
                        tmp_path = tf.name
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(tmp_path, width=Cm(14))
                    _os.unlink(tmp_path)
                except Exception as img_err:
                    make_para(f'[Фото: {url}]', WD_ALIGN_PARAGRAPH.CENTER)
            else:
                make_para('[Фото технического проекта]', WD_ALIGN_PARAGRAPH.CENTER)


# ── DOCX builder ─────────────────────────────────────────────────────────────

def build_docx(c: dict, doc_type: str, company: dict, template_settings: dict | None = None, template_blocks: list | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Mm, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Данные компании
    co_name    = co(company,'name','ООО «Интерьерные решения»')
    co_city    = co(company,'city','Саратов')
    co_inn     = co(company,'inn','6450106826')
    co_ogrn    = co(company,'ogrn','1196451012251')
    co_kpp     = co(company,'kpp','')
    co_addr    = co(company,'address','410018, г. Саратов, ул. Усть-Курдюмская, д. 3, пом. 1')
    co_bank    = co(company,'bank','')
    co_bik     = co(company,'bik','')
    co_rs      = co(company,'rs','')
    co_ks      = co(company,'ks','')
    co_phone   = co(company,'phone','')
    co_email   = co(company,'email','')

    # Данные клиента
    fname          = full_name(c)
    fname_gen      = full_name_genitive(c)
    contract_num   = c.get('contract_number') or '___'
    contract_date  = fmt_date_full(c.get('contract_date') or '')
    total          = float(c.get('total_amount') or 0)
    total_words    = num_to_words(total)
    prod_days      = int(c.get('production_days') or 45)
    prepaid        = float(c.get('prepaid_amount') or 0)
    balance        = float(c.get('balance_due') or 0) or max(0.0, total - prepaid)
    ptype          = c.get('payment_type','100% предоплата')
    manager        = (c.get('manager_name') or '').strip()
    manager_line   = manager or ('_' * 30)
    manager_gen    = genitive_name(manager) if manager else ('_' * 30)
    daddr          = delivery_addr(c)
    delivery_date  = fmt_date(c.get('delivery_date') or '')
    delivery_cost  = float(c.get('delivery_cost') or 0)
    assembly_cost  = float(c.get('assembly_cost') or 0)
    assembly_days  = int(c.get('assembly_days') or prod_days)
    products       = get_products(c)

    # ── Если переданы блоки шаблона — рендерим их вместо хардкода ─────────────
    if template_blocks:
        # Создаём документ
        from docx import Document as _Doc
        from docx.shared import Pt as _Pt, Mm as _Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
        _doc = _Doc()
        _sec = _doc.sections[0]
        _ts = template_settings or {}
        _landscape = _ts.get('orientation') == 'landscape'
        if _landscape:
            _sec.page_width  = _Mm(297); _sec.page_height = _Mm(210)
        else:
            _sec.page_width  = _Mm(210); _sec.page_height = _Mm(297)
        _fallback = _ts.get('marginMm')
        if doc_type == 'rules':   _dl,_dr,_dt,_db = 15,10,10,10
        elif doc_type == 'contract': _dl,_dr,_dt,_db = 20,10,10,10
        else: _dl,_dr,_dt,_db = 20,15,15,15
        _sec.left_margin   = _Mm(float(_ts['marginLeft'])   if 'marginLeft'   in _ts else (_fallback or _dl))
        _sec.right_margin  = _Mm(float(_ts['marginRight'])  if 'marginRight'  in _ts else (_fallback or _dr))
        _sec.top_margin    = _Mm(float(_ts['marginTop'])    if 'marginTop'    in _ts else (_fallback or _dt))
        _sec.bottom_margin = _Mm(float(_ts['marginBottom']) if 'marginBottom' in _ts else (_fallback or _db))
        _base = float(_ts.get('fontSize') or (10.0 if doc_type=='contract' else (9.5 if doc_type=='rules' else 11)))
        _fname = _ts.get('fontFamily') or 'Times New Roman'
        _style = _doc.styles['Normal']
        _style.font.name = _fname
        _style.font.size = _Pt(_base)
        _style.paragraph_format.line_spacing = _Pt(_base * 1.25)
        _style.paragraph_format.space_after  = _Pt(0)
        _style.paragraph_format.alignment    = _ALIGN.JUSTIFY
        render_blocks_to_docx(template_blocks, _doc, None, _base, _fname, c, company, products)
        _buf = BytesIO()
        _doc.save(_buf)
        return _buf.getvalue()

    # Кредитные данные
    cr_bank        = (c.get('credit_bank') or '').strip()
    cr_num         = (c.get('credit_contract_number') or '').strip()
    cr_date        = fmt_date_full(c.get('credit_contract_date') or '')
    cr_prepaid     = float(c.get('credit_prepaid') or prepaid)
    cr_balance     = float(c.get('credit_balance') or balance)
    # Доверенность менеджера
    mgr_poa_num    = (c.get('manager_poa_number') or '').strip() or '____'
    mgr_poa_date   = fmt_date(c.get('manager_poa_date') or '') if c.get('manager_poa_date') else '____________'

    # ── Создаём документ ─────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    ts = template_settings or {}
    landscape = ts.get('orientation') == 'landscape'
    if landscape:
        sec.page_width  = Mm(297); sec.page_height = Mm(210)
    else:
        sec.page_width  = Mm(210); sec.page_height = Mm(297)
    # Дефолтные поля по типу документа
    if doc_type == 'rules':
        def_l, def_r, def_t, def_b = 15, 10, 10, 10
    elif doc_type == 'contract':
        def_l, def_r, def_t, def_b = 20, 10, 10, 10
    else:
        def_l, def_r, def_t, def_b = 20, 15, 15, 15
    # Применяем настройки из шаблона (если заданы)
    fallback = ts.get('marginMm')
    sec.left_margin   = Mm(float(ts['marginLeft'])   if 'marginLeft'   in ts else (fallback or def_l))
    sec.right_margin  = Mm(float(ts['marginRight'])  if 'marginRight'  in ts else (fallback or def_r))
    sec.top_margin    = Mm(float(ts['marginTop'])    if 'marginTop'    in ts else (fallback or def_t))
    sec.bottom_margin = Mm(float(ts['marginBottom']) if 'marginBottom' in ts else (fallback or def_b))

    # Базовый размер шрифта — зависит от типа документа
    # Договор должен уместиться на 6 листов А4
    _base_pt = 10.0 if doc_type == 'contract' else (9.5 if doc_type == 'rules' else 11)
    _font_name = ts.get('fontFamily') or 'Times New Roman'

    style = doc.styles['Normal']
    style.font.name = _font_name
    style.font.size = Pt(_base_pt)
    style.paragraph_format.line_spacing = Pt(_base_pt * 1.25)
    style.paragraph_format.space_after  = Pt(0)
    style.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Вспомогательные функции ───────────────────────────────────────────────
    def font(run, size=None, bold=False):
        run.font.name = _font_name
        run.font.size = Pt(size if size is not None else _base_pt)
        run.bold = bold
        run._r.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    def heading(text, size=None):
        _sz = (size or (_base_pt + 2))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text.upper()); font(r, _sz, bold=True)
        return p

    def subheading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text); font(r)
        return p

    def section(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text); font(r, bold=True)
        p.paragraph_format.keep_with_next = True
        return p

    def para(text, indent=True, size=None, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.line_spacing = Pt(_base_pt * 1.25)
        if indent: p.paragraph_format.first_line_indent = Mm(7)
        r = p.add_run(text); font(r, size, bold)
        r.italic = italic
        return p

    def city_date(city, date_str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(3)
        from docx.oxml.ns import qn; from docx.oxml import OxmlElement
        r1 = p.add_run(f'г. {city}'); font(r1)
        p.add_run('\t\t\t\t\t')
        r2 = p.add_run(date_str); font(r2)
        return p

    def right_para(text, size=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text); font(r, size)
        return p

    def fill_cell(cell, text, bold=False, center=False, size=None):
        """Заполняет ячейку — каждая строка через \\n = отдельный параграф."""
        lines = str(text).split('\n')
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.line_spacing = Pt(_base_pt * 1.2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line); font(r, size, bold)

    def sig_table(left_header, left_body, right_header, right_body):
        """Таблица подписей: 2 колонки, заголовок жирный по центру, тело обычное."""
        # keep_with_next чтобы не отрываться от предыдущего текста
        anchor = doc.add_paragraph()
        anchor.paragraph_format.keep_with_next = True
        anchor.paragraph_format.space_before   = Pt(2)
        anchor.paragraph_format.space_after    = Pt(0)

        t = doc.add_table(rows=2, cols=2)
        t.style = 'Table Grid'
        # Заголовки
        fill_cell(t.cell(0,0), left_header,  bold=True, center=True)
        fill_cell(t.cell(0,1), right_header, bold=True, center=True)
        # Содержимое — данные сверху, подпись у нижнего края через распорку
        fill_cell(t.cell(1,0), left_body)
        fill_cell(t.cell(1,1), right_body)
        return t

    def simple_table(headers, rows, col_widths=None):
        """Простая таблица с заголовком и строками данных."""
        t = doc.add_table(rows=1+len(rows), cols=len(headers))
        t.style = 'Table Grid'
        if col_widths:
            for ci, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[ci].width = Cm(w)
        for ci, h in enumerate(headers):
            fill_cell(t.cell(0, ci), h, bold=True, center=True)
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                fill_cell(t.cell(ri+1, ci), str(val))
        return t

    # ══════════════════════════════════════════════════════════════════════════
    # ДОГОВОР БЫТОВОГО ПОДРЯДА
    # ══════════════════════════════════════════════════════════════════════════
    if doc_type == 'contract':
        def section_c(text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(text); font(r, bold=True)
            p.paragraph_format.keep_with_next = True
            return p

        heading('ДОГОВОР')
        subheading(f'бытового подряда на изготовление мебели № {contract_num}')
        city_date(co_city, contract_date)

        para(
            f'{co_name}, в лице менеджера {manager_gen}, '
            f'действующего на основании доверенности № {mgr_poa_num} от {mgr_poa_date}, '
            f'именуемый в дальнейшем «Подрядчик», и гр. {fname_gen}, '
            f'именуемый (ая) в дальнейшем «Заказчик», действующий (ая) как физическое лицо, '
            f'с одной стороны, отдельно именуемые – «Сторона», а совместно именуемые – «Стороны», '
            f'заключили настоящий Договор о нижеследующем:',
            indent=False
        )

        section_c('1. ПРЕДМЕТ ДОГОВОРА')
        para('1.1. Подрядчик обязуется выполнить работу по изготовлению мебели и передать результат работы Заказчику (мебель передается в разобранном виде), а Заказчик обязуется принять и оплатить результат работ.')
        para('1.2. Наименование, качественные характеристики, количество, дизайн мебели указываются в Техническом проекте, который является Приложением № 1 к настоящему Договору.')
        para('1.3. В случае необходимости выполнения Подрядчиком дополнительных работ, влекущих изменение объема, цены работ, включая изменение срока выполнения работ, стороны согласовывают данные изменения путем заключения дополнительного соглашения.')
        para(f'1.4. Срок выполнения работ составляет {prod_days} ({num_to_words(prod_days).replace(" рублей","").strip()}) рабочих дней, с момента согласования Технического проекта и получения Подрядчиком предварительной оплаты, в размере, указанном в разделе 3 Договора. Подрядчик вправе досрочно выполнить работу без получения предварительного согласия Заказчика.')
        para('1.4.1. В случае нарушения технологического процесса (поломка, остановка производственных линий, отсутствие энергоснабжения) по вине коммунальных и иных служб, нехватки сырья и (или) рабочей силы, если эти обстоятельства непосредственно повлияли на возможность надлежащего исполнения Подрядчиком своих обязательств по настоящему Договору, срок изготовления мебели переносится соразмерно времени, в течение которого действовали такие обстоятельства. В случае переноса сроков изготовления мебели по причине возникновения таких обстоятельств, Подрядчик уведомляет Заказчика о сроках переноса изготовления мебели в письменном виде по имеющимся средствам связи (социальные мессенджеры, электронная почта) с последующим заключением дополнительного соглашения об увеличении сроков по Договору.')
        para('1.4.2. Если Заказчик не обеспечил возможность проезда Подрядчика к месту разгрузки мебели, Подрядчик не несет ответственности за сроки доставки и разгрузки мебели и оставляет за собой право выставить счет за компенсацию дополнительных затрат, понесенных в связи с доставкой мебели, который Заказчик обязуется оплатить в течение 3 (трех) рабочих дней.')
        para('1.5. Работы, выполняемые Подрядчиком, предназначены удовлетворять бытовые или другие личные потребности Заказчика.')
        para('1.6. Заказчик проинформирован о том, что допускается незначительная разница оттенка цвета и текстуры покрытий мебели и (или) ее элементов, по сравнению с образцами, в пределах одного цветового тона, что не будет являться нарушением условий договора и основанием для предъявления соответствующих претензий Подрядчику.')
        para('1.7. Мебель изготавливается согласно индивидуальному заказу Заказчика (Технический проект). Подписанием данного договора Заказчик подтверждает, что им согласованы все характеристики мебели, включая, но не ограничиваясь: количество, размер, форма, габариты, материал, расцветка, комплектация, отделка, фурнитура, крепления и т.д. В связи с этим, на данный Договор не распространяются положения Закона РФ «О защите прав потребителей», в части возможности реализации потребителем права заменить товар или возвратить товар надлежащего качества.')

        section_c('2. РАЗРАБОТКА И СОГЛАСОВАНИЕ ТЕХНИЧЕСКОГО ПРОЕКТА')
        para('2.1. Технический проект разрабатывается Подрядчиком в течение 10 (десяти) рабочих дней с момента получения Подрядчиком предварительной оплаты, в размере, указанном в разделе 3 Договора.')
        para('2.2. Заказчик в течение 5 (пяти) рабочих дней с момента перечисления предварительной оплаты обеспечивает доступ Подрядчика в помещение, в котором будет установлена мебель для проведения замеров помещения.')
        para('2.3. Заказчик не позднее 5 (пяти) рабочих дней со дня получения от Подрядчика Технического проекта согласовывает, путем проставления личной подписи и даты согласования, или направляет мотивированный отказ от согласования Технического проекта в редакции Подрядчика с предложением своих замечаний/корректировок. В случае направления Заказчиком своих предложений/замечаний Подрядчик в течение 3 (трех) рабочих дней вносит изменения в Технический проект и направляет его на согласование Заказчику повторно.')
        para('Стороны согласовали допустимость не более 2 (двух) кругов правок в представленный Подрядчиком Технический проект. Правки Технического проекта более 2-х раз будут оценены Подрядчиком как дополнительные работы, подлежащие оплате на основании счета Подрядчика из расчета 5 000 (пять тысяч) рублей за каждый повторный Технический проект с изменениями и дополнениями Заказчика.')
        para('2.4. Все изменения или дополнения после подписания Сторонами Технического проекта по инициативе Заказчика недопустимы, за исключением изменений и дополнений, которые Подрядчик признает существенными (изменение параметров помещения по обстоятельствам, не зависящим от Заказчика и т.п.).')
        para('2.5. Заказчик обязуется в течение 3 (трех) рабочих дней с момента наступления обстоятельств, в силу которых возникла необходимость внесения изменений или дополнений в Технический проект, направить Подрядчику письменное уведомление с указанием перечня таких изменений или дополнений, а также обосновать причины внесения изменений или дополнений в Технический проект.')
        para('2.6. Подрядчик в течение 5 (пяти) рабочих дней с момента получения уведомления о внесении изменений или дополнений в Технический проект сообщает Заказчику о возможности или невозможности таких изменений или дополнений. В случае, если изменения или дополнения Технического проекта повлекут увеличение стоимости, сроков изготовления и иных условий, стороны внесут соответствующие изменения в договор путем заключения дополнительного соглашения.')

        section_c('3. СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЁТОВ')
        para(
            f'3.1. Общая стоимость работ, подлежащих выполнению по настоящему Договору складывается на основании '
            f'Калькуляции (Приложение №2 к настоящему договору) и составляет  {int(total):,} ({total_words}) '
            f'рублей, НДС не облагается на основании ст. 346.11 НК РФ. В стоимость работ включается стоимость материалов '
            f'Подрядчика, из которых производятся работы.'.replace(',', ' ')
        )
        para('3.2. Оплата работ осуществляется в следующем порядке:')

        if ptype == '100% предоплата':
            para(f'3.2.1. Предварительная оплата производится при заключении Договора в размере {int(total):,} ({total_words}) рублей.'.replace(',', ' '))
            para('3.2.2. Окончательный платёж за выполненные по Договору работы не предусмотрен.')
        elif ptype == '50% предоплата':
            para(f'3.2.1. Предварительная оплата производится при заключении Договора в размере {int(prepaid):,} ({num_to_words(prepaid)}) рублей.'.replace(',', ' '))
            para(f'3.2.2. Окончательный платёж за выполненные по Договору работы в размере {int(balance):,} ({num_to_words(balance)}) рублей осуществляется в течение 3 (трёх) дней с момента получения Заказчиком уведомления о готовности мебели, но не позднее дня доставки Заказчику, в случае если доставка осуществляется силами ООО «Интерьерные решения» на основании отдельно заключенного договора.'.replace(',', ' '))
        elif ptype == 'Рассрочка':
            para(f'3.2.1. Предварительная оплата производится при заключении Договора в размере {int(prepaid):,} ({num_to_words(prepaid)}) рублей.'.replace(',', ' '))
            para(f'3.2.2. Окончательный платёж за выполненные по Договору работы в размере {int(balance):,} ({num_to_words(balance)}) рублей осуществляется в течение 3 (трёх) дней с момента получения Заказчиком уведомления о готовности мебели, но не позднее дня доставки.'.replace(',', ' '))
        elif ptype == 'Кредит/рассрочка банка':
            bank_name = cr_bank or 'банка'
            para(f'3.2.1. Предварительная оплата производится при заключении Договора в размере {int(cr_prepaid):,} ({num_to_words(cr_prepaid)}) рублей, за счет заемных средств, при оформлении Заказчиком кредитного договора № {cr_num} от {cr_date}, заключенным между Заказчиком и {bank_name} на приобретение мебели и выплачивается Банком за Заказчика на условиях и в порядке, определенных кредитным договором.'.replace(',', ' '))
            para(f'3.2.2. Окончательный платёж за выполненные по Договору работы в размере {int(cr_balance):,} ({num_to_words(cr_balance)}) рублей, осуществляется за счет заемных средств, при оформлении Заказчиком кредитного договора № {cr_num} от {cr_date}, заключенным между Заказчиком и {bank_name} на приобретение мебели и выплачивается Банком за Заказчика на условиях и в порядке, определенных кредитным договором.'.replace(',', ' '))
        else:
            custom = (c.get('custom_payment_scheme') or '').strip()
            if custom:
                para(f'3.2.1. Порядок оплаты: {custom}')
            else:
                para(f'3.2.1. Предварительная оплата при заключении Договора: {int(prepaid):,} ({num_to_words(prepaid)}) рублей.'.replace(',', ' '))
                para(f'3.2.2. Окончательный платёж: {int(balance):,} ({num_to_words(balance)}) рублей в течение 3 дней с уведомления о готовности.'.replace(',', ' '))

        para('3.3. Оплата производится безналичным расчетом на счет Подрядчика либо наличными денежными средствами в кассу Подрядчика.')
        para('3.4. Обязательство Заказчика по безналичной оплате считается исполненным в момент зачисления денежных средств на счет Подрядчика, указанный в реквизитах.»')

        section_c('4. ПРАВА И ОБЯЗАННОСТИ СТОРОН')
        para('4.1. Подрядчик обязан:')
        para('4.1.1. Выполнить работу по Договору согласно Техническому проекту и передать Заказчику мебель в установленный срок.')
        para('4.1.2. Уведомить Заказчика о готовности мебели по электронной почте, путем обмена сообщениями (Telegram, WhatsUp, СМС и т.д.) на телефонный номер Сторон согласно реквизитам, указанным в Договоре.')
        para('4.1.3. Устранить недостатки, выявленные Заказчиком по результатам приемки работ.')
        para('4.1.4. Предоставить Заказчику необходимую и достоверную информацию о предлагаемой работе, ее видах и особенностях, о цене и форме оплаты, а также сообщить по просьбе его просьбе все относящиеся к Договору и соответствующей работе сведения. Подтверждением предоставления Заказчику указанной информации и его ознакомление с ней будет являться подписание Заказчиком настоящего Договора.')
        para('4.1.5. Предупредить Заказчика о возможных неблагоприятных для него последствиях выполнения его указаний о способе исполнения работы, а также иных не зависящих от Подрядчика обстоятельствах, которые грозят годности или прочности результатов выполняемой работы либо создают невозможность ее завершения в срок.')
        para('4.2. Заказчик обязан:')
        para('4.2.1. Согласовать и подписать Технический проект в срок, установленный Договором.')
        para('4.2.2. Оплатить стоимость работ в соответствии с условиями настоящего Договора.')
        para('4.2.3. Принять результат работ путем подписания Акта выполненных работ.')
        para('4.2.4. Обеспечить сохранность помещения, в котором должна быть установлена мебель в том виде, в котором помещение было на момент проведения замеров Подрядчиком.')
        para('4.2.5. Проводить ремонтно-отделочные работы в помещении, где будет размещена мебель, в строгом соответствии с Техническим проектом и только после получения письменного согласия Подрядчика на проведение таких работ.')
        para('4.2.6. Предоставлять Подрядчику информацию необходимую для выполнения работ по Договору и отражении ее в Техническом проекте, в том числе, но не ограничиваясь: систему крепления мебели, места установки мебели, где зазоры неприемлемы для Заказчика, свес кухонной столешницы у лицевой стороны и у стены, месторасположение и размеры техники (холодильник, вытяжка, варочная панель или газовая плита и т.д.), особенности отделки и (или) ремонта помещения, элементы декора, фотопечать, цветовые решения.')
        para('4.3. Подрядчик вправе:')
        para('4.3.1. Требовать подписания Заказчиком Акта выполненных работ в течение 5 (пяти) календарных дней с даты передачи мебели Заказчику.')
        para('4.3.2. Требовать своевременной оплаты работ в соответствии с п. 3.2 настоящего Договора.')
        para('4.3.3. Запрашивать у Заказчика информацию необходимую для надлежащего исполнения Договора.')
        para('4.3.4. Привлекать третьих лиц для исполнения обязательств по Договору. Подрядчик несет ответственность за действия/бездействие третьих лиц, выполняющих работу по настоящему Договору, как за свои собственные.')
        para('4.3.5. В случаях, когда исполнение работы по Договору стало невозможным вследствие действий или упущений Заказчика, Подрядчик сохраняет право на указанную в Договоре цену с учетом выполненной части работы.')
        para('4.3.6. Не приступать к работе, а начатую работу приостановить в случаях, когда нарушение Заказчиком обязательств, установленных в п. 4.2 Договора препятствует исполнению Договора, а также при наличии обстоятельств, очевидно свидетельствующих о том, что исполнение указанных обязанностей не будет произведено в установленный срок.')
        para('4.3.7. Досрочно выполнить работы и требовать от Заказчика принять результат работы и произвести его оплату.')
        para('4.3.8. В рекламных целях сделать фотоснимки изготовленной и установленной на месте мебели.')
        para('4.4. Заказчик вправе:')
        para('4.4.1. Выбрать модель мебели, цвет, компоновку, дизайн, материалы, фурнитуру из которых будет выполнена работа.')
        para('4.4.2. Отозвать согласие на обработку своих персональных данных Подрядчиком.')

        section_c('5. ГАРАНТИЯ И КАЧЕСТВО ВЫПОЛНЕННЫХ РАБОТ')
        para('5.1. Подтверждением качества мебели со стороны Подрядчика являются сертификаты соответствия, паспорта на товар, инструкции по эксплуатации и иная документация, подтверждающая качество мебели на момент доставки в соответствии с условиями настоящего Договора.')
        para('Гарантийный срок на мебельную фурнитуру (петли, выдвижные механизмы, выдвижные корзины, подъемные механизмы, сушки для посуды и т.п.) составляет 6 месяцев, гарантийный срок на мебель (корпус, фасады, столешницы) составляет 24 месяца, которые исчисляются с даты подписания Сторонами акта сдачи-приемки выполненных Работ, при условии надлежащей эксплуатации мебели. Срок службы изделий – 5 лет. Гарантийный срок эксплуатации Мебели и срок службы Мебели не распространяются на светильники, таймеры и другие встроенные бытовые приборы. Гарантийный срок таких приборов и срок службы указаны в паспорте на эти приборы.')
        para('5.2. Мебель, поставляемая в соответствии с настоящим Договором, должна быть осмотрена Заказчиком (уполномоченным представителем Заказчика) на предмет внешних повреждений без нарушения упаковки непосредственно при получении мебели от Подрядчика на складе последнего, либо при доставке по адресу ее установки.')
        para('5.3. Приемка Мебели по количеству и по качеству (внешние недостатки) осуществляется:')
        para('5.3.1. При самовывозе: Заказчиком на складе Подрядчика в момент приемки мебели и подписания акта выполненных работ;')
        para('5.3.2. При доставке: Подрядчиком (грузоперевозчиком Подрядчика): Заказчиком (уполномоченным им грузоперевозчиком или грузополучателем) в момент приемки мебели от Подрядчика/уполномоченного им грузоперевозчика и подписания акта об оказании услуг.')
        para('5.3.3. Если в будущем сборка мебели будет осуществляться силами ООО «Интерьерные решения» на основании отдельно заключенного договора, то Заказчик обязан вскрыть упаковку в присутствии уполномоченного лица ООО «Интерьерные решения», осуществляющего сборку мебели.')
        para('5.3.4. При обнаружении нарушения целостности упаковки, несоответствии мебели в момент приемки по качеству и/или количеству, Заказчик делает отметку в акте выполненных работ/оказания услуг. Мебель, в отношении которой у Заказчика имеются замечания, Заказчик принимает на ответственное хранение, о чем делается отметка в акте выполненных работ/оказания услуг. Плата за хранение в этом случае не взимается. Заказчик не вправе использовать (продавать, производить монтаж и т.д.) такую мебель.')
        para('5.4. Претензия по несоответствию мебели количеству, качеству (внешние дефекты), ассортименту направляется Заказчиком в адрес Подрядчика не позднее 3-х рабочих дней с даты поставки мебели с приложением фотографий, фиксирующих факт дефекта мебели либо факт некомплектности, а также фотографии упаковочного ярлыка (бирки), содержащего номер заказа.')
        para('5.4.1. Претензия должна содержать информацию о номере договора, дате приемки мебели, дате выявления недостатков, наименование и количество мебели с недостатками, описание недостатков, местонахождение мебели с недостатками, пожелания по урегулированию претензии. К претензии Заказчик обязан приложить фотографии мебели с недостатками, на которых отчетливо различимы все выявленные недостатки. Несоблюдение Заказчиком вышеуказанных требований, является основанием для отклонения претензии Подрядчиком.')
        para('5.5. В случае несоответствия количества, ассортимента, качества (внешние недостатки) мебели, возникших по вине Подрядчика, Подрядчик обязуется допоставить или обеспечить замену некачественной мебели в следующей поставке или в срок, согласованный Сторонами.')
        para('5.6. Заказчик лишается права ссылаться на недостатки мебели, если: мебель была принята без проверки на предмет внешних повреждений, на предмет соответствия ассортимента; мебель была принята без проверки на предмет соответствия количества.')
        para('5.7. В ходе рассмотрения претензии Подрядчик вправе запросить у Заказчика, а Заказчик обязан в течение 3 (трех) дней с момента получения запроса предоставить Подрядчику дополнительные фотографии мебели с недостатками, образцы мебели с недостатками. Кроме того, Подрядчик вправе выехать к месту нахождения мебели с недостатками для проведения совместной приемки мебели.')
        para('5.8. Подрядчик должен направить письменный ответ на претензию в течение 10 (Десяти) календарных дней с момента получения претензии.')
        para('5.9. В случае если претензия признается обоснованной, Подрядчик обязуется соразмерно уменьшить цену некачественной мебели или заменить некачественную мебель или ее часть (комплектующие) на качественную в срок, согласованный с Заказчиком дополнительно.')
        para('5.10. Гарантия не распространяется на недостатки мебели, которые возникли после передачи мебели Заказчику вследствие нарушения Заказчиком правил пользования мебелью, ее ненадлежащей эксплуатации, ненадлежащего или небрежного обслуживания, чрезмерной нагрузки на мебель, повреждения, нанесения на поверхность изделия едких веществ и/или жидкостей, использования мебели не по назначению, недостаточного и/или неправильного монтажа, произведенного третьими лицами, но не Заказчиком, а также на неисправности, возникшие из-за несоблюдения технических инструкций производителя, касающихся порядка и условий использования соответствующего вида мебели, на дефекты, возникшие вследствие естественного износа мебели. Гарантия также не распространяется на недостатки мебели, которые возникли после передачи мебели Заказчику вследствие действий третьих лиц, либо обстоятельств непреодолимой силы (форс-мажора).')

        section_c('6. ПОРЯДОК ПРИЁМКИ ВЫПОЛНЕННЫХ РАБОТ')
        para('6.1. По факту выполнения работ Подрядчик представляет Заказчику Акта выполненных работ в двух экземплярах по форме, согласованной в Приложении № 3 к Договору. При необоснованном отказе одной из сторон от подписания акта в нем делается отметка об этом и акт подписывается другой стороной в одностороннем порядке.')
        para('6.2. Передача мебели производится по Акту выполненных работ, который составляется в момент передачи мебели.')
        para('6.3. Приемка изготовленной мебели производится:')
        para('- в случае самовывоза Заказчиком по адресу склада Подрядчика;', indent=False)
        para('- в случае доставки мебели Подрядчиком в рамках заключенного между Заказчиком и Подрядчиком дополнительного договора доставки мебели по адресу Заказчика.', indent=False)
        para('6.4. В случае самовывоза Заказчиком мебели Подрядчик обеспечивает погрузку мебели на своем складе в транспортное средство Заказчика (перевозчика Заказчика), а Заказчик – их транспортировку и выгрузку.')
        para('6.5. Заказчик обязан в течение 3 (трех) календарных дней с даты получения уведомления от Подрядчика о готовности мебели к поставке, прибыть на склад Подрядчика для выборки мебели, в установленные в уведомлении дату и время.')
        para('6.6. После погрузки мебели на транспортное средство Заказчика Подрядчик не несет ответственности за повреждения мебели, произошедшие вследствие нарушения Заказчиком правил транспортировки, выгрузки, условий хранения и эксплуатации.')
        para('6.7. В случае доставки мебели Подрядчиком Заказчик обязуется подготовить помещение для приемки, подъезд и проход к нему, создать условия для сохранности мебели при приемке, обеспечить бесплатной парковкой автомобиль Подрядчика на период отгрузки мебели по адресу Заказчика.')
        para('6.8. При уклонении Заказчика от получения мебели, Подрядчик бесплатно хранит мебель в течение 7(семи) дней. С восьмого дня Подрядчик принимает мебель на ответственное хранение, о чем уведомляет Заказчика и вправе начислять за хранение в сумме 500 рублей за каждый день. Момент окончания ответственного хранения – передача мебели по акту приема-передачи. Дополнительных документов о принятии на ответственное хранение, кроме одностороннего уведомления, между сторонами не составляется.')
        para('6.9. При наступлении обстоятельств, объективно препятствующих поставке мебели в согласованные сторонами сроки, Подрядчик вправе в одностороннем порядке изменить срок отгрузки мебели. Указанными обстоятельствами могут быть действия третьих лиц, которые способствуют исполнению договора, погодные условия и иные обстоятельства.')
        para('6.10. При получении мебели Заказчик обязан осмотреть ее, проверить соответствие качества, количества и комплектности условиям Договора.')
        para('6.11. Риск случайной гибели или повреждения мебели, а также право собственности переходит от Подрядчика к Заказчику в момент подписания соответствующих документов о приемке (товарная накладная, универсальный передаточный документ и т.д.). В случае если доставку мебели осуществляет транспортная компания, то риск случайной гибели или повреждения мебели, а также право собственности на нее переходит к Заказчику с момента сдачи Подрядчиком мебели транспортной компании.')
        para('6.12. При наличии заказа на услуги монтажа мебели, который оформляется дополнительным договором между Подрядчиком и Заказчиком мебель монтируется Подрядчиком.')
        para('6.13. При отсутствии заказа на монтаж мебели, монтаж (установка) мебели осуществляется Заказчиком самостоятельно и за свой счет. В этом случае Подрядчик не несет ответственности за качество монтажных работ и возможные недостатки мебели, возникшие в результате самостоятельного монтажа Заказчиком.')

        section_c('7. ОТВЕТСТВЕННОСТЬ СТОРОН')
        para('7.1. Стороны несут ответственность за неисполнение или ненадлежащее исполнение своих обязательств по Договору в соответствии с Законом Российской Федерации от 7 февраля 1992 г. №2300-1 «О защите прав потребителей» и иными правовыми актами, принятыми в соответствии с ним.')
        para('7.2. За нарушение сроков оплаты услуг Подрядчик вправе потребовать с Заказчика уплаты неустойки (пени) за каждый день просрочки в размере 0,1 % от суммы задолженности.')
        para('7.3. Подрядчик не несет ответственности за невыполнение обязательств по Договору, если оно вызвано неисполнением соответствующих обязанностей Заказчика.')
        para('7.4. В случае отказа Заказчика от мебели надлежащего качества, изготовленной по индивидуальному заказу (эскизу), которая имеет нестандартные размеры, цвет, форму, и ее дальнейшая реализация невозможна, Стороны признают, что размер убытков Подрядчика составляет стоимость соответствующей (не вывезенной) партии мебели.')
        para('7.5. При разгрузке мебели силами Заказчика материальную ответственность за сохранность груза и автотранспортного средства несет Заказчик. В случае нанесения ущерба автотранспортному средству Подрядчика (Перевозчика Подрядчика), представитель Заказчика и Подрядчика (или представитель транспортной организации) составляют акт о причиненных повреждениях и причинах возникновения повреждений. При отказе представителя Заказчика от подписи в акте, акт составляется в одностороннем порядке Подрядчиком (представителем транспортной организации) и является основанием для возмещения причиненного ущерба.')

        section_c('8. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ')
        para('8.1. Все споры или разногласия, возникающие между Сторонами по настоящему Договору или в связи с ним, разрешаются путем переговоров между ними.')
        para('8.2. До передачи спора на рассмотрение суда Стороны предусматривают обязательный претензионный порядок урегулирования разногласий. Направление претензии возможно по электронной почте, путем обмена сообщениями (Telegram, WhatsUp, СМС и т.д.) на телефонный номер Сторон согласно реквизитам, указанным в Договоре. Срок ответа на претензию Сторон устанавливается в 10 (Десять) рабочих дней с момента получения претензии заинтересованной Стороны.')
        para('8.3. В случае невозможности разрешения разногласий путем переговоров споры или разногласия, возникающие между Сторонами, решаются в судебном порядке по месту нахождения Подрядчика.')

        section_c('9. СРОК ДЕЙСТВИЯ ДОГОВОРА')
        para('9.1. Договор вступает в силу с момента его подписания обеими Сторонами и действует до полного исполнения Сторонами принятых на себя обязательств.')
        para('9.2. Изменения и дополнения к Договору принимаются по обоюдному соглашению Сторон, путем подписания Дополнительного соглашения к Договору.')
        para('9.3. Договор может быть расторгнут досрочно по письменному соглашению Сторон, в одностороннем порядке в случаях, предусмотренных действующим законодательством Российской Федерации.')
        para('9.4. В случае одностороннего отказа от Договора Заказчик обязуется оплатить Подрядчику фактически понесенные им расходы, связанных с исполнением обязательств по данному Договору.')

        section_c('10. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ')
        para('10.1. При заключении Договора и подписании других документов в рамках его исполнения, Стороны согласны на факсимильное воспроизведение подписи («факсимиле»), уполномоченных лиц, а также на обмен электронными копиями документов. Стороны признают юридически значимой переписку между сторонами по электронной почте, а также переписку по номеру телефона, указанному в настоящем договоре в любых приложениях (Telegram, WhatsUp, СМС и т.д.). Направленные таким образом документы считаются подписанными простой электронной подписью и признаются Сторонами равнозначными бумажным, подписанным собственноручной подписью Сторон.')
        para('10.2. Стороны признают юридическую силу всех документов, требований, уведомлений, претензий, извещений, оформленных должным образом и направленных друг другу в электронном виде по указанным адресам электронной почты и телефонов во исполнение настоящего Договора. Стороны договорились, что при такой форме коммуникации считается доставленным тогда, когда приложение, через которое производится коммуникация подтверждает факт направления сообщения через свой внутренний интерфейс. При этом и Заказчик, и Подрядчик обязуются регулярно просматривать указанные источники коммуникации. Стороны договорились, что сообщение считается доставленным и в тех случаях, если оно поступило лицу, которому оно направлено, но по обстоятельствам, зависящим от него, не было ему вручено или адресат не ознакомился с ним.')
        para('10.3. Каждая Сторона обязуется обеспечивать конфиденциальность полученной ею в связи с заключением или исполнением Договора от другой Стороны информации ограниченного доступа.')
        para(
            f'10.4. Заказчик дает свое согласие на обработку своих персональных данных, а именно: на действия, совершаемые с использованием средств автоматизации или без использования таких средств, включая сбор, запись, систематизацию, накопление, хранение, уточнение (обновление, изменение), извлечение, использование, передачу (распространение, предоставление, доступ), обезличивание, блокирование, удаление, уничтожение его персональных данных, Обществу с ограниченной ответственностью «ИНТЕРЬЕРНЫЕ РЕШЕНИЯ» (ОГРН: {co_ogrn}, ИНН: {co_inn}), расположенному по адресу: {co_addr}.')
        para('Цель обработки персональных данных: исполнение настоящего договора. Заказчик дает свое согласие на использование следующих персональных данных: фамилия, имя, отчество, паспортные данные, адрес места жительства, фотографии изделий на объекте Заказчика; номер телефона, адрес электронной почты. Согласие предоставляется на срок действия настоящего договора, а после прекращения договора – в течение 12 месяцев с даты подписания Сторонами акта сдачи-приемки выполненных Работ.')
        para('Настоящее условие договора может быть изменено Заказчиком – субъектом персональных данных, в любой момент в одностороннем порядке путем отзыва согласия на обработку персональных данных. Отзыв согласия на обработку персональных данных осуществляется посредством составления письменного документа, который может быть направлен в адрес Подрядчика почтовым отправлением с уведомлением о вручении, либо вручен лично под расписку представителю Подрядчика.')

        # Строка подписи — одна строка: линия (подпись) ... линия (расшифровка)
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlEl
        from docx.shared import Cm as _Cm

        # Устанавливаем таб-стопы: центр листа для правой подписи
        def _add_tab_stop(para, pos_cm, align='left'):
            pPr = para._p.get_or_add_pPr()
            tabs_el = pPr.find(_qn('w:tabs'))
            if tabs_el is None:
                tabs_el = _OxmlEl('w:tabs')
                pPr.append(tabs_el)
            tab = _OxmlEl('w:tab')
            tab.set(_qn('w:val'), align)
            tab.set(_qn('w:pos'), str(int(pos_cm * 567)))  # 1cm = 567 twips
            tabs_el.append(tab)

        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(10)
        p_line.paragraph_format.space_after  = Pt(0)
        _add_tab_stop(p_line, 9.5)
        r1 = p_line.add_run('_______________________  /'); font(r1)
        p_line.add_run('\t')
        r2 = p_line.add_run('_______________________  /'); font(r2)

        p_lbl = doc.add_paragraph()
        p_lbl.paragraph_format.space_before = Pt(1)
        p_lbl.paragraph_format.space_after  = Pt(6)
        _add_tab_stop(p_lbl, 9.5)
        r3 = p_lbl.add_run('(подпись)'); font(r3, size=8)
        p_lbl.add_run('\t')
        r4 = p_lbl.add_run('(расшифровка подписи от руки)'); font(r4, size=8)

        para('10.5. К настоящему Договору прилагаются и являются неотъемлемой частью следующие приложения:')
        para('1. Технический проект.', indent=False)
        para('2. Калькуляция работ.', indent=False)
        para('3. Правила эксплуатации корпусной мебели.', indent=False)
        para('4. Образец Акта выполненных работ.', indent=False)

        # Раздел реквизитов — keep_with_next, без принудительного переноса страницы
        p_sec11 = doc.add_paragraph()
        p_sec11.paragraph_format.space_before   = Pt(4)
        p_sec11.paragraph_format.space_after    = Pt(1)
        p_sec11.paragraph_format.keep_with_next = True
        r11 = p_sec11.add_run('11. РЕКВИЗИТЫ СТОРОН'); font(r11, bold=True)

        # Реквизиты компании — все поля из настроек
        inn_kpp = f'ИНН/КПП: {co_inn}/{co_kpp}' if co_kpp else f'ИНН: {co_inn}'
        left_body = f'{co_name}\n{inn_kpp}\nОГРН: {co_ogrn}\nЮридический и фактический адрес: {co_addr}'
        if co_bank or co_bik or co_rs or co_ks:
            left_body += '\nБанковские реквизиты:'
            if co_bank: left_body += f'\n{co_bank}'
            if co_bik:  left_body += f'\nБИК: {co_bik}'
            if co_rs:   left_body += f'\nРас/с: {co_rs}'
            if co_ks:   left_body += f'\nКор/с: {co_ks}'
        if co_phone: left_body += f'\nТелефон: {co_phone}'
        if co_email: left_body += f'\nE-mail: {co_email}'
        left_body += f'\n\nМенеджер:\n{manager_line}\n\n\n_________________________\n(подпись) М.П.'

        # Реквизиты заказчика — паспорт, адрес, телефон (без канала связи)
        reg_city  = c.get('reg_city','') or ''
        reg_str   = c.get('reg_street','') or ''
        reg_house = c.get('reg_house','') or ''
        reg_apt   = c.get('reg_apt','') or ''
        reg_addr  = ', '.join(filter(None, [reg_city, reg_str, reg_house]))
        if reg_apt: reg_addr += f', кв. {reg_apt}'

        passport_issued_date = fmt_date(c.get('passport_issued_date') or '')
        passport_dept = c.get('passport_dept_code','') or ''
        passport_date_dept = ', '.join(filter(None, [passport_issued_date, passport_dept]))

        right_body = f'{fname}\nПаспорт. Серия, номер: {passport_str(c)}'
        right_body += f'\nКем выдан: {c.get("passport_issued_by") or "___________"}'
        if passport_date_dept: right_body += f'\nДата выдачи: {passport_date_dept}'
        right_body += f'\n\nАдрес прописки:\n{reg_addr or "___________"}'
        right_body += f'\n\nТелефон: {c.get("phone") or "___________"}'
        right_body += '\n\n\n\n\n\n\n\n_________________________\n(подпись)'
        sig_table('Подрядчик:', left_body, 'Заказчик:', right_body)

    # ══════════════════════════════════════════════════════════════════════════
    # АКТ ВЫПОЛНЕННЫХ РАБОТ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'act':
        # Шапка по центру
        p_app = doc.add_paragraph()
        p_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_app.paragraph_format.space_before = Pt(0); p_app.paragraph_format.space_after = Pt(2)
        font(p_app.add_run(f'Приложение № 4 к договору бытового подряда на изготовление мебели № {contract_num} от {contract_date}'), 9)

        heading('«АКТ ВЫПОЛНЕННЫХ РАБОТ»')

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date.paragraph_format.space_before = Pt(0); p_date.paragraph_format.space_after = Pt(6)
        font(p_date.add_run('от «____» ________________ 20____ г.'), _base_pt)

        manager_gen_act = genitive_name(manager) if manager else '_' * 30
        fname_gen_act = full_name_genitive(c)
        para(f'{co_name}, в лице менеджера {manager_gen_act}, действующего на основании доверенности № {mgr_poa_num} от {mgr_poa_date}, именуемый в дальнейшем «Подрядчик», и гр. {fname_gen_act}, именуемый (ая) в дальнейшем «Заказчик», подписали настоящий Акт выполненных работ о нижеследующем:', indent=False)

        p1 = doc.add_paragraph(); p1.paragraph_format.space_before = Pt(10); p1.paragraph_format.space_after = Pt(2)
        font(p1.add_run(f'1. Подрядчик изготовил для Заказчика мебель по договору бытового подряда № {contract_num} от {contract_date}:'), _base_pt)

        if products:
            rows = [(i+1, p.get('name','Кухонный гарнитур'), 'шт.', p.get('qty',1), '') for i, p in enumerate(products)]
        else:
            rows = [(1, 'Кухонный гарнитур', 'шт.', 1, f'{int(total):,}'.replace(',', ' '))]
        rows.append(('', '', '', 'ИТОГО:', f'{int(total):,} ({total_words})'.replace(',', ' ')))
        simple_table(['№','Наименование мебели, включая ее элементы','Ед. изм.','Кол-во изделий','Стоимость в руб.'], rows, [1,7,2,2,3])

        p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(10); p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run('2. Комплектность, количество, вид, характеристики мебели соответствуют условиям договора. Визуальный осмотр мебели на предмет повреждений, царапин, сколов, трещин и других недостатков произведен Заказчиком. Фурнитура (петли, выдвижные механизмы, подъемные механизмы и т.д.) работает исправно. Заказчик претензий по объему, качеству, результату и срокам выполнения работ: ')
        font(r2, _base_pt)
        r2b = p2.add_run('не имеет / имеет'); font(r2b, _base_pt, bold=True)
        r2c = p2.add_run(' (ненужное зачеркнуть).'); font(r2c, _base_pt)
        p2.paragraph_format.first_line_indent = None

        # 6 линий для записей замечаний
        from docx.oxml.ns import qn as _aqn; from docx.oxml import OxmlElement as _aEl
        for _ in range(6):
            pl = doc.add_paragraph()
            pl.paragraph_format.space_before = Pt(14)
            pl.paragraph_format.space_after  = Pt(0)
            pl.paragraph_format.line_spacing = Pt(14)
            pBdr = _aEl('w:pBdr'); bot = _aEl('w:bottom')
            bot.set(_aqn('w:val'), 'single'); bot.set(_aqn('w:sz'), '4'); bot.set(_aqn('w:space'), '1'); bot.set(_aqn('w:color'), '000000')
            pBdr.append(bot)
            pl._p.get_or_add_pPr().append(pBdr)
            font(pl.add_run(' '), _base_pt)

        p3 = doc.add_paragraph(); p3.paragraph_format.space_before = Pt(10); p3.paragraph_format.space_after = Pt(2)
        font(p3.add_run('3. В случае наличия замечаний Заказчик, после подписания акта, вправе требовать устранения замечаний, отражённых в данном акте.'), _base_pt)
        p3.paragraph_format.first_line_indent = None
        p4 = doc.add_paragraph(); p4.paragraph_format.space_before = Pt(2); p4.paragraph_format.space_after = Pt(2)
        font(p4.add_run('4. Настоящий акт подписан в 2 (двух) экземплярах по одному для каждой из Сторон.'), _base_pt)
        p4.paragraph_format.first_line_indent = None

        # Подписи без таблицы
        from docx.oxml.ns import qn as _sqn2; from docx.oxml import OxmlElement as _sEl2
        p_sig = doc.add_paragraph(); p_sig.paragraph_format.space_before = Pt(12)
        def _tab3(para, pos_cm):
            pPr = para._p.get_or_add_pPr()
            tabs = pPr.find(_sqn2('w:tabs'))
            if tabs is None: tabs = _sEl2('w:tabs'); pPr.append(tabs)
            tab = _sEl2('w:tab'); tab.set(_sqn2('w:val'), 'left'); tab.set(_sqn2('w:pos'), str(int(pos_cm * 567))); tabs.append(tab)
        _tab3(p_sig, 10)
        font(p_sig.add_run(f'Подрядчик: {co_name}'), _base_pt, bold=True)
        p_sig.add_run('\t')
        font(p_sig.add_run(f'Заказчик: {fname}'), _base_pt, bold=True)

        p_lines3 = doc.add_paragraph(); p_lines3.paragraph_format.space_before = Pt(14)
        _tab3(p_lines3, 10)
        font(p_lines3.add_run('______________________________'), _base_pt)
        p_lines3.add_run('\t')
        font(p_lines3.add_run('______________________________'), _base_pt)

        p_mp2 = doc.add_paragraph(); p_mp2.paragraph_format.space_before = Pt(1)
        font(p_mp2.add_run('М.П.'), _base_pt)

    # ══════════════════════════════════════════════════════════════════════════
    # ТЕХНИЧЕСКИЙ ПРОЕКТ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'tech':
        # ── Альбомная страница A4: ширина=297, высота=210
        from docx.enum.section import WD_ORIENT
        sec.orientation   = WD_ORIENT.LANDSCAPE
        sec.page_width    = Mm(297); sec.page_height   = Mm(210)
        sec.left_margin   = Mm(5);   sec.right_margin  = Mm(5)
        sec.top_margin    = Mm(5);   sec.bottom_margin = Mm(5)

        # Доступная ширина: 297 - 5 - 5 = 287мм, минус отступы ячеек ~4мм = 283мм
        CONTENT_W = Mm(283)

        # ── Заголовок
        p_hdr = doc.add_paragraph()
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_hdr.paragraph_format.space_before = Pt(0)
        p_hdr.paragraph_format.space_after  = Pt(0)
        r_hdr = p_hdr.add_run(
            f'Приложение № 1 к договору бытового подряда на изготовление мебели № {contract_num} от  {contract_date}'
        ); font(r_hdr, 9)

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after  = Pt(2)
        r_title = p_title.add_run('«Технический проект»'); font(r_title, 10, bold=True)

        # ── Данные материалов
        korpus    = c.get('tech_korpus','') or ''
        korpus2   = c.get('tech_korpus2','') or ''
        fasad1    = c.get('tech_fasad1','') or ''
        fasad2    = c.get('tech_fasad2','') or ''
        stoleshn  = c.get('tech_stoleshniza','') or ''
        stenovaya = c.get('tech_stenovaya','') or ''
        pod_type  = c.get('tech_podsvetka_type','') or ''
        pod_svet  = c.get('tech_podsvetka_svet','') or ''
        frezerovka= c.get('tech_frezerovka','') or ''

        # ── Таблица характеристик: 5 строк × 6 колонок
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlEl

        tbl = doc.add_table(rows=5, cols=6)
        tbl.style = 'Table Grid'
        tbl.autofit = False
        tbl.allow_autofit = False

        # Убираем отступы ячеек таблицы чтобы ширины были точными
        _tblCellMar = _OxmlEl('w:tblCellMar')
        for _s in ('top', 'left', 'bottom', 'right'):
            _m = _OxmlEl(f'w:{_s}')
            _m.set(_qn('w:w'), '0'); _m.set(_qn('w:type'), 'dxa')
            _tblCellMar.append(_m)
        tbl._tbl.tblPr.append(_tblCellMar)

        # Ширины колонок: сумма = CONTENT_W = 283мм
        col_w = [Mm(22), Mm(74), Mm(33), Mm(74), Mm(17), Mm(63)]
        for row in tbl.rows:
            for ci, w in enumerate(col_w):
                row.cells[ci].width = w

        def tc(ri, ci, text, bold=False, size=9):
            cell = tbl.cell(ri, ci)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.line_spacing = Pt(11)
            r = p.add_run(text); font(r, size, bold)

        # Строка 0: Корпус 1 | val | Столешница | val
        tc(0,0,'Корпус 1:', True);     tc(0,1, korpus)
        tc(0,2,'Столешница:', True);   tc(0,3, stoleshn)
        tbl.cell(0,4).merge(tbl.cell(0,5))

        # Строка 1: Корпус 2 | val | Стеновая панель | val
        tc(1,0,'Корпус 2:', True);     tc(1,1, korpus2)
        tc(1,2,'Стеновая панель:', True); tc(1,3, stenovaya)
        tbl.cell(1,4).merge(tbl.cell(1,5))

        # Строка 2: Фасад 1 | val | (пусто)
        tc(2,0,'Фасад 1:', True);      tc(2,1, fasad1)
        tbl.cell(2,2).merge(tbl.cell(2,3)).merge(tbl.cell(2,4)).merge(tbl.cell(2,5))

        # Строка 3: Фасад 2 | val | Подсветка | Тип: val | Свет: val
        tc(3,0,'Фасад 2:', True);      tc(3,1, fasad2)
        tc(3,2,'Подсветка', True)
        tc(3,3,'Тип:', True);          tc(3,4, pod_type)
        tc(3,5,'Свет:  ' + pod_svet)

        # Строка 4: Фрезеровка | объединённые ячейки
        tc(4,0,'Фрезеровка:', True)
        merged_frez = tbl.cell(4,1).merge(tbl.cell(4,2)).merge(tbl.cell(4,3)).merge(tbl.cell(4,4)).merge(tbl.cell(4,5))
        p_frez = merged_frez.paragraphs[0]
        p_frez.paragraph_format.space_before = Pt(1); p_frez.paragraph_format.space_after = Pt(1)
        r_frez = p_frez.add_run(frezerovka); font(r_frez, 9)

        # ── Изображение проекта в таблице с фиксированной высотой строки
        import urllib.request, io as _io, traceback as _tb
        from PIL import Image as _PILImage
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlEl

        tech_img = str(c.get('tech_image_url') or '').strip()

        # 210 - 5 - 5(поля) - ~20(заголовок) - ~22(таблица хар-к) - ~22(дисклеймер+подписи) = ~136мм
        IMG_H_MM = 130
        IMG_W = CONTENT_W
        IMG_H = Mm(IMG_H_MM)

        # Таблица 1×1 — контейнер для картинки
        it = doc.add_table(rows=1, cols=1)
        it.style = 'Table Grid'
        ic = it.cell(0, 0)
        ic.width = IMG_W

        # Минимальная высота строки (atLeast — растянется если картинка больше)
        _tr = it.rows[0]._tr
        _trPr = _tr.get_or_add_trPr()
        _trH = _OxmlEl('w:trHeight')
        _trH.set(_qn('w:val'), str(int(IMG_H_MM * 56.69)))
        _trH.set(_qn('w:hRule'), 'atLeast')
        _trPr.append(_trH)

        # Убираем внутренние отступы ячейки
        _tcPr = ic._tc.get_or_add_tcPr()
        _tcMar = _OxmlEl('w:tcMar')
        for _s in ('top', 'left', 'bottom', 'right'):
            _m = _OxmlEl(f'w:{_s}')
            _m.set(_qn('w:w'), '0'); _m.set(_qn('w:type'), 'dxa')
            _tcMar.append(_m)
        _tcPr.append(_tcMar)

        # 20 пустых строк перед картинкой чтобы она была по центру ячейки
        from docx.oxml.ns import qn as _qn2
        for _ in range(30):
            pe = ic.add_paragraph()
            pe.paragraph_format.space_before = Pt(0)
            pe.paragraph_format.space_after  = Pt(0)
            pe.paragraph_format.line_spacing = Pt(12)

        p_img = ic.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after  = Pt(0)

        print(f'[TECH] tech_img={tech_img[:80] if tech_img else "EMPTY"}')
        if tech_img:
            try:
                req = urllib.request.Request(tech_img, headers={'User-Agent': 'Mozilla/5.0'})
                img_bytes = urllib.request.urlopen(req, timeout=20).read()
                print(f'[TECH] loaded {len(img_bytes)} bytes')

                pil_img = _PILImage.open(_io.BytesIO(img_bytes))
                iw, ih = pil_img.size
                print(f'[TECH] size {iw}x{ih}')

                buf = _io.BytesIO()
                pil_img.convert('RGB').save(buf, format='JPEG', quality=90)
                buf.seek(0)

                if iw > 0 and ih > 0:
                    projected_h = int((IMG_W / iw) * ih)
                    add_kw = {'width': IMG_W} if projected_h <= IMG_H else {'height': IMG_H}
                else:
                    add_kw = {'width': IMG_W}

                print(f'[TECH] add_picture kw={add_kw}')
                p_img.add_run().add_picture(buf, **add_kw)
                print('[TECH] OK')
            except Exception as ex:
                print(f'[TECH] FAILED: {ex}\n{_tb.format_exc()}')
                font(p_img.add_run(f'[ ОШИБКА: {ex} ]'), 8)
        else:
            font(p_img.add_run('[ нет фото ]'), 10)

        # ── Дисклеймер
        p_disc = doc.add_paragraph()
        p_disc.paragraph_format.space_before = Pt(2)
        p_disc.paragraph_format.space_after  = Pt(2)
        r_disc = p_disc.add_run(
            'Подписывая Технический проект, Заказчик подтверждает, что ознакомлен с наименованием, '
            'качественными характеристиками, количеством, дизайном мебели и ему полностью понятны '
            'выполняемые Подрядчиком работы. Стороны согласовали, что мебель изготовлена специально '
            'для Заказчика по его индивидуальным параметрам. Приложение: бланк замера.'
        )
        font(r_disc, 8); r_disc.italic = True

        # ── Подписи через таблицу 2 колонки (надёжнее чем табуляция)
        from docx.oxml.ns import qn as _sqn
        from docx.oxml import OxmlElement as _sEl
        sig = doc.add_table(rows=3, cols=2)
        sig.style = 'Table Grid'
        sig.autofit = False
        sig.allow_autofit = False
        for row in sig.rows:
            row.cells[0].width = Mm(120)
            row.cells[1].width = Mm(163)
        _sb = _sEl('w:tblBorders')
        for _side in ('top','left','bottom','right','insideH','insideV'):
            _e = _sEl(f'w:{_side}'); _e.set(_sqn('w:val'),'none'); _sb.append(_e)
        sig._tbl.tblPr.append(_sb)

        def _sc(row, col, text, bold=False, size=8, sb=3, align=WD_ALIGN_PARAGRAPH.LEFT):
            p = sig.cell(row, col).paragraphs[0]
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after  = Pt(0)
            p.alignment = align
            font(p.add_run(text), size, bold)

        _sc(0, 0, f'Подрядчик: {co_name}', bold=True)
        _sc(0, 1, f'Заказчик:  {fname}', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _sc(1, 0, '______________________________', size=9, sb=8)
        _sc(1, 1, '______________________________', size=9, sb=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _sc(2, 0, 'М.П.', sb=1)

        # Убираем лишний параграф который Word добавляет после таблицы
        p_last = doc.add_paragraph()
        p_last.paragraph_format.space_before = Pt(0)
        p_last.paragraph_format.space_after  = Pt(0)
        p_last.paragraph_format.line_spacing = Pt(1)

    # ══════════════════════════════════════════════════════════════════════════
    # ПРАВИЛА ЭКСПЛУАТАЦИИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'rules':
        def section_c(text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(text); font(r, bold=True)
            p.paragraph_format.keep_with_next = True
            return p

        p_app = doc.add_paragraph()
        p_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_app.paragraph_format.space_before = Pt(0); p_app.paragraph_format.space_after = Pt(2)
        font(p_app.add_run(f'Приложение № 3 к договору бытового подряда на изготовление мебели № {contract_num} от {contract_date}'), 9)
        heading('«ПРАВИЛА ЭКСПЛУАТАЦИИ КОРПУСНОЙ МЕБЕЛИ»')

        section_c('1. ОБЩИЕ РЕКОМЕНДАЦИИ')
        para('1.1. Срок службы Мебели и сохранение его потребительских свойств напрямую зависят от соблюдения Заказчиком правил, изложенных в настоящем приложении.')
        para('1.2. Климатические условия и воздействия окружающей среды (свет, влажность, температура) напрямую влияют на состояние мебели.')
        para('1.3. Подрядчик рекомендует соблюдать оптимальные климатические условия в помещении, где установлена мебель: температура воздуха от +18°C до +25°C, относительная влажность воздуха 45% - 70%.')
        para('1.4. Следует оберегать мебель от длительного воздействия прямых солнечных лучей, источников тепла (батареи, обогреватели, духовые шкафы, плиты на расстоянии менее 0,5 м), а также от резких перепадов температуры и влажности.')
        para('1.5. Запрещается воздействие на Мебель агрессивных жидкостей (кислот, щелочей, растворителей), абразивных чистящих средств и материалов, способных повредить покрытие.')
        para('1.6. Мебель предназначена для эксплуатации в жилых или общественных помещениях в соответствии с ее функциональным назначением. Подрядчик не несет ответственности за повреждения, вызванные несоблюдением рекомендуемых условий.')
        para('1.7. Подрядчик гарантирует соответствие Мебели обязательным требованиям нормативных документов, действующих на территории РФ, в том числе:\n* ГОСТ 19917-2014 «Мебель для сидения и лежания. Общие технические условия»;\n* ГОСТ 32289-2013 «Плиты древесно-стружечные, облицованные пленками на основе термореактивных полимеров. Технические условия»;\n* ГОСТ 16371-2014 «Мебель. Общие технические условия»;\n* ТР ТС 025/2012 «О безопасности мебельной продукции».')

        section_c('2. УСЛОВИЯ ХРАНЕНИЯ')
        para('2.1. Хранение Мебели до ее установки должно осуществляться в сухих, проветриваемых, отапливаемых помещениях, защищенных от атмосферных осадков, прямых солнечных лучей и источников тепла. Запрещается хранить изделия в помещениях с повышенной влажностью (санузлы, бани, подвалы без отопления).')
        para('2.2. Рекомендуемый диапазон температуры воздуха от +10°C до +25°C. Запрещена эксплуатация при температурах ниже -20°C и выше +40°C, а также резкие перепады температур.')
        para('2.3. Рекомендуемая относительная влажность воздуха 45% - 70%. Длительное воздействие повышенной влажности (свыше 80%) или чрезмерной сухости (ниже 40%) недопустимо и приводит к деформации элементов (разбуханию, рассыханию, расслоению).')
        para('2.4. Мебель должна храниться в заводской упаковке на ровной поверхности. Запрещается хранить изделия в вертикальном положении, прислонив к стене.')
        para('2.5. Перед установкой Мебели, доставленной или хранившейся при отрицательных температурах, необходимо выдержать ее в оригинальной упаковке в условиях помещения не менее 72 часов (3 суток) для адаптации к комнатной температуре и влажности.')

        section_c('3. ПРАВИЛА ЭКСПЛУАТАЦИИ И УХОДА ЗА МЕБЕЛЬЮ')
        para('3.1. Общие правила:', bold=True, indent=False)
        para('3.1.1. Не превышайте максимально допустимые нагрузки на элементы Мебели: полки - до 15 кг, выдвижные ящики из ЛДСП - до 5 кг, ящики на системе «метабокс» - до 18 кг, ящики на системе «тандембокс» - до 35 кг.')
        para('3.1.2. Общая нагрузка на подвесной шкаф не должна превышать 70 кг. Высокие конструкции (пеналы, стеллажи) должны быть больше нагружены в нижней части.')
        para('3.1.3. Равномерно распределяйте нагрузку внутри шкафов и ящиков: тяжелые предметы размещайте ближе к краям и опорам, легкие - в центре.')
        para('3.1.4. Исключите попадание воды на незащищенные торцы деталей, места стыков и торцов изделия.')
        para('3.1.5. Запрещается воздействовать на поверхности абразивными, кислотными, щелочными средствами, ацетоном, растворителями.')
        para('3.1.6. Не допускается заслонять вентиляционные решетки и воздухозаборные отверстия бытовых приборов, встроенных в мебель.')
        para('3.1.7. При установке Мебели в бревенчатых домах, вследствие усадки таких домов, возможна деформация: перекашивание фасадов, опускание как верхних, так и нижних модулей, это не является признаком некачественной работы Подрядчика.')
        para('3.2. Корпус ЛДСП:', bold=True, indent=False)
        para('3.2.1. Главное правило! ЛДСП боится длительного контакта с водой. Попадание жидкости на кромки и тем более в места стыков недопустимо – плита разбухнет и деформируется. Все пролитые жидкости следует немедленно вытирать насухо.')
        para('3.2.2. Запрещается мыть ЛДСП большим количеством воды, использовать пароочистители. Не оставляйте не отжатые тряпки, мокрые полотенца на торцах деталей.')
        para('3.2.3. Запрещается применять абразивные и едкие химические средства.')
        para('3.2.4. Протирайте сухой или слегка влажной хорошо отжатой тканью. Для удаления загрязнений использовать мягкие мыльные растворы.')
        para('3.2.5. Берегите кромки и торцы от сколов и ударов. Именно через поврежденные кромки влага легче всего проникает внутрь плиты.')
        para('3.2.6. Не превышайте максимально допустимые нагрузки на полки и перегородки. Равномерно распределяйте вес.')
        para('3.2.7. Избегайте расположения Мебели вплотную к отопительным приборам (батареям, обогревателям), это может привести к расслоению ламинированного слоя и деформации плиты.')
        para('3.3. Фасады МДФ:', bold=True, indent=False)
        para('3.3.1. Избегайте резких перепадов температуры и попадания прямых солнечных лучей.')
        para('3.3.2. Не допускайте длительного контакта поверхности с сильно нагретыми предметами (сковородки, кастрюли, утюги).')
        para('3.3.3. Избегайте ударов, царапин острыми предметами, давления на выступающие элементы.')
        para('3.3.4. Исключите контакт с агрессивными химическими веществами: растворителями, ацетоном, сильными чистящими порошками, средствами для мытья стекол. Регулярно удаляйте пыль мягкой тканью.')
        para('3.3.5. Разрешается использовать нейтральные (pH-нейтральные) средства, а также мягкие мыльные растворы.')
        para('3.3.6. Запрещается мыть МДФ большим количеством воды, использовать пароочистители. Не оставляйте не отжатые тряпки, мокрые полотенца на торцах деталей.')
        para('3.3.7. Не рекомендуется снимать защитную пленку с фасадов до окончания процесса установки мебели.')
        para('3.3.8. При изготовлении дверных полотен высотой более 1600 мм, рекомендовано использовать систему выпрямления дверей, чтобы избежать возможной деформации изделия.')
        para('3.4. Столешницы и стеновые панели из пластика и искусственного (акрилового) камня:', bold=True, indent=False)
        para('3.4.1. Для очистки используйте мягкую ткань, мягкую губку, салфетки из микрофибры и средства для ухода за глянцевыми поверхностями.')
        para('3.4.2. Запрещается использовать жесткие и металлические губки, щетки и абразивные чистящие средства. Использовать при очистке кислоты, щелочи, соли, растворители.')
        para('3.4.5. Запрещается использовать поверхность в качестве разделочной доски. Воздействовать на поверхность столешницы острыми предметами. Передвигать по поверхности посуду с металлическим дном.')
        para('3.4.6. На пластиковой столешнице не оставляйте лужи воды, не отжатые тряпки, мокрые полотенца, особенно на стыках и возле мойки. Сразу удаляйте воду с поверхностей столешницы и со стыковочных швов.')
        para('3.4.7. Не допускается ставить на поверхность столешницы горячие (>60°С) предметы. Используйте специальные термоизоляционные подставки под горячее, разделочные доски, салфетки, коврики.')
        para('3.4.8. Не допускается размораживать продукты, или оставлять на длительное время на поверхности столешницы сильно охлаждённые (<0°C) предметы.')
        para('3.6. Стеклянные и зеркальные поверхности:', bold=True, indent=False)
        para('3.6.1. Очищайте специальными средствами для стекол и зеркал, нанося состав на мягкую ткань, а не прямо на поверхность.')
        para('3.6.2. Избегайте абразивных средств и жестких губок. Не допускайте ударных и чрезмерных нагрузок на стеклянные полки.')
        para('3.7. Фурнитура и механизмы:', bold=True, indent=False)
        para('3.7.1. Протирание сухой мягкой тканью. Для удаления загрязнений допускается использование слабого мыльного раствора.')
        para('3.7.2. Запрещается прилагать чрезмерные усилия для открывания/закрывания дверей и ящиков. Угол открывания распашных дверей, как правило, не превышает 90°.')
        para('3.7.3. Запрещается открывать фасады с системой «Push to Open» любым способом, кроме нажатия на фасад.')
        para('3.7.4. Выдвигайте ящики полностью только при необходимости, держась за ручки или фасад.')
        para('3.7.5. Регулярно удаляйте пыль, крошки и грязь с направляющих и других движущихся частей.')
        para('3.7.6. Регулярная регулировка и смазка механизмов (парафином или специальными средствами) являются обязанностью Заказчика и не покрываются гарантией.')
        para('3.7.7. Лицевую фурнитуру следует чистить мягкими тканями с применением хозяйственного мыла, после чего вытирать насухо. Не использовать средства, содержащие абразивные материалы (наждачную бумагу, соду и др.).')

        p_italic = doc.add_paragraph()
        p_italic.paragraph_format.space_before = Pt(6)
        r_it = p_italic.add_run('Соблюдая эти несложные правила, вы сохраните безупречный вид и функциональность вашей мебели на долгие годы. Помните: что несоблюдение правил эксплуатации может привести к сокращению сроков службы и преждевременному выходу из строя элементов кухонной мебели.')
        font(r_it, 9); r_it.italic = True
        p_warn = doc.add_paragraph()
        p_warn.paragraph_format.space_before = Pt(4)
        r_w = p_warn.add_run('ВНИМАНИЕ! Подрядчик не несет ответственность за последствия от несоблюдения установленных норм и правил по уходу и эксплуатации корпусной мебели.')
        font(r_w, 9, bold=True); r_w.italic = True

        # Подписи без рамки
        from docx.oxml.ns import qn as _sqn2
        from docx.oxml import OxmlElement as _sEl2
        p_sig = doc.add_paragraph()
        p_sig.paragraph_format.space_before = Pt(10)
        p_sig.paragraph_format.space_after  = Pt(0)
        def _tab2(para, pos_cm):
            pPr = para._p.get_or_add_pPr()
            tabs = pPr.find(_sqn2('w:tabs'))
            if tabs is None:
                tabs = _sEl2('w:tabs'); pPr.append(tabs)
            tab = _sEl2('w:tab')
            tab.set(_sqn2('w:val'), 'left')
            tab.set(_sqn2('w:pos'), str(int(pos_cm * 567)))
            tabs.append(tab)
        _tab2(p_sig, 10)
        font(p_sig.add_run(f'Подрядчик: {co_name}'), 9, bold=True)
        p_sig.add_run('\t')
        font(p_sig.add_run(f'Заказчик: {fname}'), 9, bold=True)

        p_lines2 = doc.add_paragraph()
        p_lines2.paragraph_format.space_before = Pt(12)
        p_lines2.paragraph_format.space_after  = Pt(0)
        _tab2(p_lines2, 10)
        font(p_lines2.add_run('______________________________'), 9)
        p_lines2.add_run('\t')
        font(p_lines2.add_run('______________________________'), 9)

        p_mp = doc.add_paragraph()
        p_mp.paragraph_format.space_before = Pt(1)
        font(p_mp.add_run('М.П.'), 9)

    # ══════════════════════════════════════════════════════════════════════════
    # ДОГОВОР ДОСТАВКИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'delivery':
        heading('ДОГОВОР')
        subheading('на оказание услуг по доставке мебели')
        city_date(co_city, contract_date)
        para(f'{co_name}, в лице менеджера {manager_line}, именуемый «Исполнитель», и гр. {fname}, именуемый «Заказчик», заключили настоящий Договор:', indent=False)

        section('1. ПРЕДМЕТ ДОГОВОРА')
        para(f'1.1. Исполнитель обязуется доставить мебель по договору подряда № {contract_num} от {contract_date}.')
        para(f'1.2. Адрес доставки: {daddr}. Этаж: {c.get("delivery_floor") or "___"}. Лифт: {c.get("delivery_elevator") or "нет"}.')
        para(f'1.3. Дата доставки: {delivery_date}. Конкретное время согласовывается дополнительно.')
        para('1.4. Доставка осуществляется в разобранном виде.')

        section('2. ПРАВА И ОБЯЗАННОСТИ СТОРОН')
        para('2.1. Исполнитель обязан: доставить в срок по указанному адресу; обеспечить сохранность при транспортировке; уведомить об изменении времени не менее чем за 2 часа.')
        para('2.2. Заказчик обязан: обеспечить доступ к месту доставки; оплатить услуги; при обнаружении повреждений сообщить до подписания акта.')
        para('2.3. Исполнитель вправе однократно предоставить скидку в размере стоимости доставки (8 000 руб.) при доставке в пределах г. Саратова и г. Энгельса.')

        section('3. СТОИМОСТЬ И ОПЛАТА')
        para('3.1. Стоимость услуг определяется Приложением № 1.')
        para('3.2. Оплата производится в день доставки до начала разгрузки.')

        section('4. ОТВЕТСТВЕННОСТЬ')
        para('4.1. При повреждении мебели по вине Исполнителя при транспортировке — Исполнитель возмещает ущерб.')
        para('4.2. При отказе от доставки менее чем за 24 часа Исполнитель вправе удержать фактически понесённые расходы.')

        section('5. ПРОЧИЕ УСЛОВИЯ')
        para('5.1. Договор вступает в силу с момента подписания. Составлен в двух экземплярах.')
        inn_kpp = f'ИНН/КПП: {co_inn}/{co_kpp}' if co_kpp else f'ИНН: {co_inn}'
        left_body = f'{co_name}\n{co_addr}\n{inn_kpp}'
        if co_rs: left_body += f'\nр/с: {co_rs}'
        if co_bik: left_body += f'\nБИК: {co_bik}'
        left_body += f'\n\nМенеджер: ______________________________\nМ.П.'
        sig_table(
            'Исполнитель', left_body,
            'Заказчик', f'{fname}\nПаспорт: {passport_str(c)}\nАдрес доставки: {daddr}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # АКТ ПРИЁМА ДОСТАВКИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'act_delivery':
        dcost_words = num_to_words(delivery_cost) if delivery_cost else '___________'
        right_para(f'к договору на оказание услуг по доставке мебели от {contract_date}')
        heading('«АКТ ПРИЁМА-ПЕРЕДАЧИ ДОСТАВКИ МЕБЕЛИ»')
        city_date(co_city, contract_date)
        para(f'{co_name}, именуемый «Исполнитель», и гр. {fname}, именуемый «Заказчик», составили настоящий Акт:', indent=False)
        para(f'1. Исполнитель доставил Заказчику мебель по адресу: {daddr}.')
        para(f'2. Дата доставки: {delivery_date}.')
        para('3. Мебель доставлена в полном объёме, внешних механических повреждений не выявлено.')
        para('4. Заказчик произвёл осмотр мебели в момент приёмки. Претензий нет.')
        para(f'5. Стоимость услуг по доставке: {delivery_cost:,.0f} ({dcost_words}) рублей. Оплата произведена.')
        para('6. Услуги по доставке выполнены в полном объёме.')
        sig_table(
            'Исполнитель', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\nПаспорт: {passport_str(c)}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # ДОГОВОР МОНТАЖА
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'assembly':
        heading('ДОГОВОР')
        subheading('на оказание услуг по сборке и монтажу мебели')
        city_date(co_city, contract_date)
        para(f'{co_name}, именуемый «Исполнитель», и гр. {fname}, именуемый «Заказчик», заключили настоящий Договор:', indent=False)

        section('1. ПРЕДМЕТ ДОГОВОРА')
        para(f'1.1. Исполнитель обязуется выполнить работы по сборке и монтажу мебели по договору подряда № {contract_num} от {contract_date}.')
        para(f'1.2. Адрес выполнения работ: {daddr}.')
        para(f'1.3. Ориентировочная дата начала работ: {delivery_date}. Срок выполнения: {assembly_days} рабочих дней.')
        para('1.4. В объём работ входит: сборка корпусных элементов, установка фасадов и фурнитуры, регулировка петель, монтаж столешницы и стеновых панелей, подключение подсветки (при наличии).')

        section('2. ПРАВА И ОБЯЗАННОСТИ СТОРОН')
        para('2.1. Исполнитель обязан: выполнить монтаж качественно; убрать строительный мусор; уведомить о дефектах мебели или помещения.')
        para('2.2. Заказчик обязан: обеспечить доступ и электроснабжение; принять работы; оплатить согласно условиям.')

        section('3. СТОИМОСТЬ И ОПЛАТА')
        acost_str = f'{assembly_cost:,.0f} ({num_to_words(assembly_cost)}) рублей' if assembly_cost else '______________________________'
        para(f'3.1. Стоимость работ по сборке и монтажу составляет {acost_str}.')
        para('3.2. Оплата производится в день завершения монтажных работ до подписания Акта приёмки.')

        section('4. ОТВЕТСТВЕННОСТЬ')
        para('4.1. Исполнитель несёт ответственность за качество монтажных работ в течение 12 месяцев.')
        para('4.2. Гарантия не распространяется на дефекты от нарушения правил эксплуатации или механических повреждений.')

        section('5. ПРОЧИЕ УСЛОВИЯ')
        para('5.1. Приёмка работ оформляется подписанием Акта приёмки. 5.2. Договор в двух экземплярах.')
        sig_table(
            'Исполнитель', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\nПаспорт: {passport_str(c)}\nАдрес монтажа: {daddr}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # АКТ ПРИЁМА СБОРКИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'act_assembly':
        acost_words = num_to_words(assembly_cost) if assembly_cost else '___________'
        right_para(f'к договору на оказание услуг по сборке и монтажу мебели от {contract_date}')
        heading('«АКТ ПРИЁМА-ПЕРЕДАЧИ ВЫПОЛНЕННЫХ РАБОТ ПО СБОРКЕ И МОНТАЖУ МЕБЕЛИ»')
        city_date(co_city, contract_date)
        para(f'{co_name}, именуемый «Исполнитель», и гр. {fname}, именуемый «Заказчик», составили настоящий Акт:', indent=False)
        para(f'1. Исполнитель выполнил работы по сборке и монтажу мебели по адресу: {daddr}.')

        if products:
            rows = [(i+1, p.get('name','Кухонный гарнитур'), 'шт.', p.get('qty',1)) for i, p in enumerate(products)]
        else:
            rows = [(1, 'Кухонный гарнитур', 'шт.', 1)]
        simple_table(['№','Наименование мебели','Ед. изм.','Кол-во'], rows, [1,10,2,2])

        para('2. Объём работ: сборка корпусных элементов, установка фасадов, регулировка петель, монтаж столешницы.')
        para('3. Заказчик произвёл проверку. Фурнитура проверена. Претензий нет.')
        para(f'4. Стоимость работ: {assembly_cost:,.0f} ({acost_words}) рублей. Оплата произведена.')
        para('5. Гарантийный срок на монтажные работы — 12 месяцев.')
        sig_table(
            'Исполнитель', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\nПаспорт: {passport_str(c)}\nТелефон: {c.get("phone") or "___________"}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # КАЛЬКУЛЯЦИЯ ДОСТАВКИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'delivery_calc':
        right_para('Приложение № 1 к договору на оказание услуг по доставке мебели')
        heading('«КАЛЬКУЛЯЦИЯ НА ВЫПОЛНЕНИЕ УСЛУГ ПО ДОСТАВКЕ МЕБЕЛИ»')
        simple_table(
            ['Наименование работ и услуг','Ед. изм.','Кол-во','Цена, руб.','Стоимость, руб.'],
            [
                ('Доставка в пределах г. Саратова и г. Энгельса *','1 услуга','1','8 000','8 000'),
                ('Доставка за пределы г. Саратова и г. Энгельса **','1 км','','70',''),
                ('','','Итого:','','8 000'),
                ('','','Скидка ***:','','8 000'),
                ('','','Итого со скидкой:','','0'),
            ],
            [8,2,2,2,2]
        )
        para('* Исполнитель вправе однократно предоставить скидку в размере 8 000 руб.', indent=False)
        para('** Километраж от склада Исполнителя до адреса Заказчика по Яндекс.Картам.', indent=False)
        para('*** Размер скидки определяется Исполнителем индивидуально.', indent=False)
        sig_table(
            'Подрядчик', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # ПРАЙС ПОДЪЁМ МЕБЕЛИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'delivery_lift':
        right_para('Приложение № 2 к договору на оказание услуг по доставке мебели')
        heading('«ПРАЙС НА ВЫПОЛНЕНИЕ УСЛУГ ПО ПОДЪЁМУ И ЗАНОСУ МЕБЕЛИ»')
        para('Подъём мебели при отсутствии лифта и занос при невозможности парковки вплотную к подъезду.', indent=False, bold=True)
        simple_table(
            ['Наименование','Ед. изм.','Кол-во (этаж)','Цена, руб.'],
            [
                ('Квадратура корпуса до 20 кв.м','руб./этаж','','550'),
                ('Квадратура корпуса 20–25 кв.м','руб./этаж','','650'),
                ('Квадратура корпуса более 25 кв.м','руб./этаж','','750'),
                ('Перемещение вручную при невозможности подъезда','1 м','','30'),
                ('Подъём столешницы','1 уп./1 этаж','','350'),
                ('Подъём стеновой панели','1 уп./1 этаж','','250'),
                ('Подъём крупных частей корпуса','1 уп./1 этаж','','250'),
                ('Подъём дверей-купе','1 дверь/1 этаж','','150'),
            ],
            [9,3,3,2]
        )
        para('1. Подъём на лифте — бесплатно. 2. Занос на 1 этаж — бесплатно при парковке вплотную.', indent=False)
        para('* Услуги рассчитываются по факту оказания.', indent=False)
        sig_table(
            'Подрядчик', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # КАЛЬКУЛЯЦИЯ СБОРКИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'assembly_calc':
        acost_val = assembly_cost or 17500
        right_para('Приложение № 1 к договору на выполнение работ по монтажу и сборке мебели')
        heading('«КАЛЬКУЛЯЦИЯ НА ВЫПОЛНЕНИЕ РАБОТ ПО СБОРКЕ МЕБЕЛИ»')
        simple_table(
            ['Наименование работ и услуг','Ед. изм.','Кол-во','Цена, руб.','Стоимость, руб.'],
            [
                (f'Сборка и монтаж мебели по договору № {contract_num} от {contract_date} *','работа','1',f'{acost_val:,.0f}',f'{acost_val:,.0f}'),
                ('Выезд сборщика за пределы г. Саратова и г. Энгельса **','1 км','40','','0'),
                ('','','Итого:','',f'{acost_val:,.0f}'),
                ('','','Скидка ***:','',f'{acost_val:,.0f}'),
                ('','','Итого со скидкой:','','0'),
            ],
            [9,2,2,2,2]
        )
        para('* Стоимость сборки. Повторный выезд по вине Заказчика — по данной калькуляции.', indent=False)
        para('** Километраж от склада Подрядчика по Яндекс.Картам.', indent=False)
        para('*** Размер скидки определяется Подрядчиком индивидуально.', indent=False)
        sig_table(
            'Подрядчик', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # ПРАЙС ДОП. РАБОТ СБОРКИ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'assembly_extra':
        right_para('Приложение № 2 к договору на выполнение работ по монтажу и сборке мебели')
        heading('«ПРАЙС НА ДОПОЛНИТЕЛЬНЫЕ РАБОТЫ»')
        simple_table(
            ['Наименование работ','Ед. изм.','Кол-во','Цена, руб.'],
            [
                ('Врезка мойки с герметизацией (без подвода воды)','шт.','','1 250'),
                ('Врезка варочной поверхности (без пуска газа)','шт.','','1 250'),
                ('Установка вытяжки без принудительного воздуховода','шт.','','1 500'),
                ('Установка духового шкафа','шт.','','750'),
                ('Установка встраиваемой СВЧ-печи','шт.','','650'),
                ('Установка встраиваемого холодильника','шт.','','2 400'),
                ('Установка посудомоечной машины (без подвода воды)','шт.','','1 500'),
                ('Установка светодиодной ленты','п.м.','','1 250'),
                ('Установка столешницы в подоконник','шт.','','3 000'),
                ('Установка ручек (куплены клиентом)','шт.','','70'),
            ],
            [10,2,2,3]
        )
        para('Стоимость работ, не включённых в прайс, обсуждается индивидуально.', indent=False)
        para('* Работы рассчитываются по факту выполнения.', indent=False)
        sig_table(
            'Подрядчик', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'addendum':
        heading('ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ')
        subheading(f'к договору бытового подряда № {contract_num} от {contract_date}')
        city_date(co_city, contract_date)
        para(f'{co_name}, именуемый «Подрядчик», и гр. {fname}, именуемый «Заказчик», заключили настоящее соглашение:', indent=False)

        section('1. ПРЕДМЕТ СОГЛАШЕНИЯ')
        para('1.1. Стороны договорились внести следующие изменения в Договор:')
        for _ in range(6):
            p_bl = doc.add_paragraph('_' * 85)
            p_bl.paragraph_format.space_after = Pt(4)

        section('2. ПРОЧИЕ УСЛОВИЯ')
        para('2.1. Настоящее соглашение является неотъемлемой частью Договора и вступает в силу с момента подписания.')
        para('2.2. В остальной части условия Договора остаются без изменений.')
        para('2.3. Соглашение составлено в двух экземплярах.')
        inn_kpp = f'ИНН/КПП: {co_inn}/{co_kpp}' if co_kpp else f'ИНН: {co_inn}'
        left_body = f'{co_name}\nОГРН: {co_ogrn}, {inn_kpp}\n{co_addr}\n\nМенеджер: {manager_line}\nПодпись: ______________________________\nМ.П.'
        sig_table(
            'Подрядчик', left_body,
            'Заказчик', f'{fname}\nПаспорт: {passport_str(c)}\nТел.: {c.get("phone") or "___________"}\n\nПодпись: ______________________________'
        )

    # ══════════════════════════════════════════════════════════════════════════
    # СПЕЦИФИКАЦИЯ НА ТЕХНИКУ
    # ══════════════════════════════════════════════════════════════════════════
    elif doc_type == 'tech_spec':
        right_para(f'Приложение к договору бытового подряда № {contract_num} от {contract_date}')
        heading('«СПЕЦИФИКАЦИЯ НА ПОСТАВКУ ТЕХНИКИ»')
        rows = [('','шт.','','','') for _ in range(8)]
        rows.append(('','','','ИТОГО:',''))
        simple_table(['Наименование','Ед. изм.','Кол-во','Цена, руб.','Стоимость, руб.'], rows, [8,2,2,3,3])
        sig_table(
            'Подрядчик', f'{co_name}\n\nМенеджер: ______________________________\nМ.П.',
            'Заказчик', f'{fname}\n\nПодпись: ______________________________'
        )

    else:
        # Неизвестный тип — пустой документ с заголовком
        heading(f'Документ: {doc_type}')
        para('Тип документа не поддерживается.')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HANDLER ───────────────────────────────────────────────────────────────────

def handler(event: dict, context) -> dict:
    """Генератор DOCX-документов. Принимает client_id и doc_type, возвращает base64 DOCX."""
    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': CORS, 'body': ''}

    payload = verify_token(event)
    if not payload:
        return err('Unauthorized', 401)

    qs = event.get('queryStringParameters') or {}
    action   = qs.get('action', '')
    cid      = qs.get('client_id', '')
    doc_type = qs.get('doc', 'contract')

    if action not in ('doc_docx', 'doc_zip'):
        return err('Неизвестное действие', 404)

    if not cid:
        return err('Нет client_id')

    user_id = payload.get('sub') or payload.get('user_id') or payload.get('id')
    schema = os.environ.get('MAIN_DB_SCHEMA', 'public')

    # Тестовый клиент для превью шаблона
    if cid == 'preview':
        from datetime import date
        client = {
            'last_name': 'Иванов', 'first_name': 'Иван', 'middle_name': 'Иванович',
            'phone': '+7 (999) 123-45-67', 'phone2': '+7 (999) 765-43-21', 'email': 'ivanov@mail.ru',
            'city': 'Саратов',
            'passport_series': '4520', 'passport_number': '123456',
            'passport_issued_by': 'ОУФМС по г. Москве', 'passport_date': '2015-03-15', 'passport_code': '770-001',
            'registration_address': 'г. Москва, ул. Ленина, д. 5, кв. 12',
            'delivery_city': 'г. Москва', 'delivery_street': 'ул. Садовая', 'delivery_house': 'д. 3', 'delivery_apt': '8',
            'contract_number': '877', 'contract_date': str(date.today()),
            'total_amount': 350000, 'prepaid_amount': 175000, 'balance_due': 175000,
            'payment_type': 'наличные',
            'delivery_cost': 3000, 'assembly_cost': 8000,
            'production_days': 30, 'assembly_days': 2,
            'delivery_date': str(date.today()),
            'manager_name': 'Сазонов Василий Николаевич',
            'designer_name': 'Петрова Анна Сергеевна',
            'products': [], 'tech_image_url': '',
        }
    else:
        # Загружаем клиента из БД
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute(f'SELECT * FROM {schema}.clients WHERE id = %s', (cid,))
            row = cur.fetchone()
            if not row:
                return err('Клиент не найден', 404)
            cols = [d[0] for d in cur.description]
            client = dict(zip(cols, row))

            # Для tech — подгружаем фото категории render если нет tech_image_url
            if doc_type in ('tech',) or action == 'doc_zip':
                if not client.get('tech_image_url'):
                    cur.execute(f"SELECT url FROM {schema}.client_photos WHERE client_id = %s AND category = 'render' ORDER BY uploaded_at DESC LIMIT 1", (cid,))
                    photo_row = cur.fetchone()
                    if photo_row:
                        client['tech_image_url'] = photo_row[0]
    company = get_company(user_id)

    # Подгружаем доверенность менеджера из БД
    mgr_name = client.get('manager_name', '') or ''
    if mgr_name:
        poa = get_manager_poa(mgr_name)
        if poa.get('poa_number') or poa.get('poa_date'):
            client['manager_poa_number'] = poa.get('poa_number', '')
            client['manager_poa_date']   = poa.get('poa_date', '')

    # Загружаем настройки и блоки шаблона
    template_settings = {}
    template_blocks = None
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute(
            f"SELECT settings, blocks FROM {schema}.doc_templates WHERE user_id = %s AND doc_type = %s AND is_default = true LIMIT 1",
            (str(user_id), doc_type)
        )
        trow = cur.fetchone()
        if trow:
            template_settings = trow[0] or {}
            raw_blocks = trow[1]
            if raw_blocks:
                if isinstance(raw_blocks, str):
                    try: raw_blocks = json.loads(raw_blocks)
                    except: raw_blocks = None
                if isinstance(raw_blocks, list) and raw_blocks:
                    template_blocks = raw_blocks

    if action == 'doc_docx':
        logger.info(f'doc_docx: {doc_type} for client {cid}, tech_img={str(client.get("tech_image_url",""))[:80]}')
        docx_bytes = build_docx(client, doc_type, company, template_settings, template_blocks)
        logger.info(f'doc_docx: generated {len(docx_bytes)} bytes')
        return {
            'statusCode': 200,
            'headers': {**CORS, 'Content-Type': 'application/json'},
            'body': json.dumps({'data': base64.b64encode(docx_bytes).decode()}),
        }

    elif action == 'doc_zip':
        import zipfile, io as _io
        DOCS_ZIP = {
            'contract':       '01. Договор бытового подряда',
            'tech':           '02. Технический проект (Прил.1)',
            'act':            '03. Акт выполненных работ (Прил.4)',
            'rules':          '04. Правила эксплуатации (Прил.3)',
            'delivery':       '05. Договор доставки',
            'act_delivery':   '06. Акт приёма доставки',
            'assembly':       '07. Договор монтажа',
            'act_assembly':   '08. Акт приёма сборки',
            'delivery_calc':  '09. Калькуляция доставки (Прил.1)',
            'delivery_lift':  '10. Прайс подъём мебели (Прил.2)',
            'assembly_calc':  '11. Калькуляция сборки (Прил.1)',
            'assembly_extra': '12. Прайс доп. работ (Прил.2)',
            'addendum':       '13. Дополнительное соглашение',
            'tech_spec':      '14. Спецификация на технику',
        }
        fname_client = full_name(client).replace(' ', '_')[:25]
        zip_buf = _io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for dt, name in DOCS_ZIP.items():
                try:
                    # Загружаем блоки для каждого типа документа отдельно
                    _tblocks = None
                    _tsettings = template_settings.copy()
                    with get_db() as _conn:
                        _cur = _conn.cursor()
                        _cur.execute(
                            f"SELECT settings, blocks FROM {schema}.doc_templates WHERE user_id = %s AND doc_type = %s AND is_default = true LIMIT 1",
                            (str(user_id), dt)
                        )
                        _tr = _cur.fetchone()
                        if _tr:
                            _tsettings = _tr[0] or {}
                            _rb = _tr[1]
                            if _rb:
                                if isinstance(_rb, str):
                                    try: _rb = json.loads(_rb)
                                    except: _rb = None
                                if isinstance(_rb, list) and _rb:
                                    _tblocks = _rb
                    b = build_docx(client, dt, company, _tsettings, _tblocks)
                    zf.writestr(f'{name} — {fname_client}.docx', b)
                    logger.info(f'zip: added {dt} ({len(b)} bytes)')
                except Exception as ex:
                    logger.error(f'zip: failed {dt}: {ex}')
        fname_zip = f'Документы — {fname_client}.zip'
        return {
            'statusCode': 200,
            'headers': {**CORS, 'Content-Type': 'application/json'},
            'body': json.dumps({'data': base64.b64encode(zip_buf.getvalue()).decode(), 'filename': fname_zip}),
        }